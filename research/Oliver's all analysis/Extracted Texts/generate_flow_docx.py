"""Generate NETS Homework Session Script documentation."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    if level == 1:
        h.font.size = Pt(20)
        h.font.bold = True
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(12)
    elif level == 2:
        h.font.size = Pt(15)
        h.font.bold = True
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    else:
        h.font.size = Pt(12)
        h.font.bold = True
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

# ============================================================
# COVER
# ============================================================
for _ in range(5):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('NETS Homework Session')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0xf5, 0x9e, 0x0b)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Complete Session Script')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run(
    'A step-by-step walkthrough of one complete homework session\n'
    'From the moment a student opens their assignment to the moment results are delivered\n\n'
    'Optimized Flow — 6 Phases + Engagement Hook\n'
    'Total Duration: 25 minutes 30 seconds'
)
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(
    'NETS — National Education Transformation System\n'
    'Version 1.0 | April 3, 2026\n'
    'Republic of Uzbekistan, Ministry of Education\n'
    'Proprietary'
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

doc.add_page_break()

# ============================================================
# THE CAST
# ============================================================
doc.add_heading('The Cast', level=1)

doc.add_paragraph('This script follows one student through one complete homework session.')

cast = [
    ('Sardor Karimov', 'Student — Grade 5-A, School 14, Tashkent. Reading PISA: 1.4. Streak: 4 days.'),
    ('Ms. Barno', 'Teacher — History, Grade 5. Assigned this homework with default settings.'),
    ('The System', 'NETS Homework Engine — handles assembly, adaptation, scoring, and reporting.'),
    ('Mr. Karimov', 'Parent — receives a plain-language summary after the session.'),
]

for name, role in cast:
    p = doc.add_paragraph()
    run = p.add_run(f'{name} ')
    run.bold = True
    p.add_run(f'— {role}')

doc.add_paragraph()

doc.add_paragraph('Assignment: History Grade 5, Chapter III, Section 10 — "Yozuv — insoniyatning buyuk kashfiyoti" (Writing — Humanity\'s Great Discovery). Textbook pages 30-33.')

doc.add_page_break()

# ============================================================
# ACT 1: PRE-SESSION
# ============================================================
doc.add_heading('Act 1: Behind the Screen', level=1)
doc.add_paragraph('Time: 0:00 — Duration: <1 second')

doc.add_paragraph(
    'Sardor clicks on his homework assignment. He sees a loading screen with the NETS logo and a progress bar. '
    'He does not know it yet, but in the next half-second, the system is doing four things.'
)

doc.add_heading('Step 1: Load Sardor\'s Profile', level=2)
doc.add_paragraph('The system pulls up Sardor\'s data:')

profile_data = [
    'Reading PISA: 1.4 (target: 2.0)',
    'Domain breakdown: access=1.6, interpret=1.3, evaluate=1.1',
    'Transition skills mastered: "Locate explicit information in short text"',
    'Transition skills in progress: "Identify main idea," "Compare info across simple sources"',
    'Mastery map: 12 standards tracked. 3 are below 0.6 (needs review).',
    'Spaced repetition: 8 items due for review from last week\'s session.',
    'Current streak: 4 days. One more day unlocks the 7-day streak bonus (+50% XP).',
]

for item in profile_data:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Step 2: Check Content Pool', level=2)
doc.add_paragraph(
    'The system checks: does this topic have enough approved content? It scans the database for items tagged to '
    'Chapter III, Section 10, pages 30-33. It finds 72 approved items — above the minimum of 65. Content check passes.'
)

doc.add_heading('Step 3: Assemble the Session', level=2)
doc.add_paragraph('The system selects specific items for each phase:')

assembly = [
    'Phase 1 (Activation): 6 recall items from the previous chapter ("Tarixiy manbalar" — Historical Sources), '
    'prioritized by spaced repetition. 2 items Sardor got wrong last time, 3 items not seen in 7+ days, 1 random.',
    'Phase 2 (Story Mode): 2 narrative segments from pages 30-33, each with a checkpoint question.',
    'Phase 3 (Practice Lab): Tile Match (reinforce — current level) + Why Chain (stretch — one level above, '
    'targets Sardor\'s weakest skill: "compare info across sources").',
    'Phase 4 (Real-Life): 1 transfer scenario — "Museum visit: identifying a Sogdian writing clay fragment."',
    'Phase 5 (Boss): 4 questions — 2 easy (PISA L2), 1 medium (PISA L3), 1 hard (PISA L4).',
    'Phase 6 (Reflection): 3 prompt templates loaded; one will be selected based on performance.',
]

for i, item in enumerate(assembly):
    doc.add_paragraph(f'{i+1}. {item}')

doc.add_heading('Step 4: Apply Teacher Overrides', level=2)
doc.add_paragraph(
    'Ms. Barno did not change any settings. The session runs with defaults: Standard mode (25 min), '
    'Boss HP at 100, hints enabled, Real-Life mandatory, auto-selected games.'
)

doc.add_paragraph(
    'The session plan is locked in. The loading screen fades. Sardor sees his first screen.'
)

doc.add_page_break()

# ============================================================
# ACT 2: ENGAGEMENT HOOK
# ============================================================
doc.add_heading('Act 2: The Warm-Up', level=1)
doc.add_paragraph('Time: 0:01 — Duration: 30 seconds')

doc.add_paragraph(
    'Before the real homework starts, Sardor sees a bright screen:'
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Quick Warm-Up! 30 seconds. How many can you get?')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0xf5, 0x9e, 0x0b)

doc.add_paragraph('A timer starts counting down from 30. Questions appear one at a time:')

hook_qs = [
    'An image of Egyptian hieroglyphs appears. "What is this?" Options: [Writing] [Drawing] [Art] [Code]. '
    'Sardor taps "Writing." Correct. +10 XP.',
    'A photo of a clay tablet with wedge marks. "How old is this?" Options: [100 years] [1,000 years] '
    '[5,000 years] [10,000 years]. Sardor taps "5,000 years." Correct. +10 XP.',
    'The Phoenician alphabet shown. "How many letters?" Options: [10] [22] [26] [50]. Sardor taps "22." '
    'Correct. +10 XP.',
    'A modern Uzbek textbook cover. "What writing system?" Options: [Cyrillic] [Latin] [Arabic] [Hieroglyphs]. '
    'Sardor taps "Latin." Correct. +10 XP.',
]

for q in hook_qs:
    doc.add_paragraph(q, style='List Bullet')

doc.add_paragraph(
    'The timer hits 0. The screen changes:'
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Great warm-up! 4 out of 4. Let\'s start the real session!')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)

doc.add_paragraph(
    'Sardor earned 40 XP in 30 seconds. His brain is now primed on the topic of writing systems. '
    'He feels good — four wins in a row. He is ready.'
)

doc.add_page_break()

# ============================================================
# ACT 3: ACTIVATION
# ============================================================
doc.add_heading('Act 3: Memory Sprint', level=1)
doc.add_paragraph('Time: 0:31 — Duration: 2 minutes')

doc.add_paragraph('The screen changes to a darker, more focused theme. A timer appears: 1:00.')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Xotira Sprinti — Oldingi bilimlarni faollashtirish')
run.font.size = Pt(14)
run.font.bold = True

doc.add_paragraph('Question 1 of 6 appears. No multiple choice — Sardor must type his answer.')

sprint_qs = [
    ('"Tarixiy manbalar necha turga bo\'linadi? Nomlang."', 'Sardor types: "2 ta — moddiy va yozma." '
     'Correct. +100 XP + 10 streak. Response time: 8 seconds.'),
    ('"Arxeologiya so\'zining ma\'nosi nima?"', 'Sardor types: "Qadimshunoslik." Correct. +100 XP + 10 streak. '
     'Response time: 6 seconds.'),
    ('"Grigoriy kalendari qaysi yilda qabul qilindi?"', 'Sardor types: "1582." Correct. +100 XP + 10 streak. '
     'Response time: 5 seconds. Speed bonus: +10 XP.'),
    ('"Xronologiya so\'zi qaysi tildan olingan?"', 'Sardor types: "Yunoncha." Correct. +100 XP + 10 streak. '
     'Response time: 7 seconds.'),
    ('"Avesto qayerda yaratilgan?"', 'Sardor types: "Buxoro." Incorrect. The correct answer is "Xorazm." '
     'Streak resets. No XP. Response time: 12 seconds.'),
    ('"Dastlabki globusni kim yaratgan?"', 'Sardor types: "Krates." Correct. +100 XP. Response time: 9 seconds.'),
]

for q, result in sprint_qs:
    p = doc.add_paragraph()
    run = p.add_run(f'Q: {q}')
    run.bold = True
    doc.add_paragraph(result, style='List Bullet')

doc.add_paragraph('Timer runs out. Results:')

results = [
    '5 out of 6 correct = 83% accuracy',
    'Total XP earned: 540',
    'System verdict: "Ajoyib! Yangi kontentga tayyorsiz." (Excellent! You are ready for new content.)',
    'Flow state indicator shows green: 78% — in the flow zone.',
]

for r in results:
    doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph(
    'Sardor clicks "Faza 2 ga o\'tish" (Proceed to Phase 2).'
)

doc.add_page_break()

# ============================================================
# ACT 4: STORY MODE
# ============================================================
doc.add_heading('Act 4: Story Mode', level=1)
doc.add_paragraph('Time: 2:31 — Duration: 5 minutes')

doc.add_heading('Segment 1: The First Writing', level=2)
doc.add_paragraph('Sardor reads:')

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
p.paragraph_format.right_indent = Cm(1)
run = p.add_run(
    '"Sardor, a 5th-grade student from Samarkand, went on a school trip to an archaeological dig site. '
    'The archaeologist, Akmal aka, pulled an ancient clay tablet from the ground. \'This is one of the first '
    'forms of writing,\' Akmal aka said. \'Thousands of years ago, people started drawing symbols on clay to '
    'remember things — how much grain they had, how many animals they owned. These were not letters like we use '
    'today. They were pictures and marks that meant something.\' Sardor looked closely at the tablet. The marks '
    'looked like tiny triangles pressed into wet clay. \'Who made this?\' he asked. \'The Sumerians,\' Akmal aka '
    'replied. \'They lived in Mesopotamia, and they invented what we call cuneiform — wedge-shaped writing.\'"'
)
run.italic = True
run.font.size = Pt(11)

doc.add_paragraph('A checkpoint question appears:')

p = doc.add_paragraph()
run = p.add_run('Checkpoint: ')
run.bold = True
p.add_run('"Nima uchun odamlar dastlab loyga belgi qo\'yishni boshladilar?" (Why did people first start putting marks on clay?)')

doc.add_paragraph('Sardor types: "Narsalarni eslab qolish uchun — don, hayvonlar soni." (To remember things — grain, animal counts.)')

doc.add_paragraph('Correct. The system shows: "To\'g\'ri! Keyingi segmentga o\'tamiz." (Correct! Moving to the next segment.)')

doc.add_heading('Segment 2: The Alphabet Revolution', level=2)
doc.add_paragraph('Sardor reads:')

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
p.paragraph_format.right_indent = Cm(1)
run = p.add_run(
    '"The next day, Sardor\'s teacher, Barno opa, showed the class a picture of the Phoenician alphabet. '
    '\'This changed everything,\' she said. \'Before this, writing systems had thousands of symbols. You had to '
    'memorize them all. But the Phoenicians created an alphabet with only 22 letters. With just 22 symbols, you '
    'could write any word. This made writing accessible to ordinary people, not just priests and scribes.\' '
    'Sardor thought about his own alphabet — the Uzbek Latin script he uses every day. It has 29 letters. '
    'Not so different from 22, he realized."'
)
run.italic = True
run.font.size = Pt(11)

doc.add_paragraph('Checkpoint question:')

p = doc.add_paragraph()
run = p.add_run('Checkpoint: ')
run.bold = True
p.add_run('"Finikiy alifbosi nima uchun muhim edi? Nima uchun 22 ta harf minglab belgilardan yaxshiroq?" '
          '(Why was the Phoenician alphabet important? Why were 22 letters better than thousands of symbols?)')

doc.add_paragraph('Sardor types: "Chunki 22 ta harfni o\'rganish oson. Hammayoq yozishni o\'rgana oldi, faqat ruhoniylar emas." '
                  '(Because 22 letters are easy to learn. Everyone could learn to write, not just priests.)')

doc.add_paragraph('Correct. Both segments completed.')

doc.add_heading('Memory Palace Lock-In (1 minute)', level=2)
doc.add_paragraph(
    'Before moving on, Sardor sees a map of Registan Square with 5 locations marked. He drags 5 key concepts '
    'from today\'s lesson onto the map:'
)

palace_items = [
    'Main Gate → "Cuneiform — wedge-shaped writing on clay"',
    'Left Madrasah → "Phoenician alphabet — 22 letters"',
    'Right Madrasah → "Sumerians — first writers"',
    'Central Fountain → "Alphabet revolution — writing for everyone"',
    'The Pool → "Uzbek alphabet — 29 letters, from Phoenician roots"',
]

for item in palace_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'A quick recall test: "What was at the Left Madrasah?" Sardor answers correctly. 4 out of 5 correct. '
    '+200 XP. The concepts are now locked into his memory.'
)

doc.add_paragraph('Sardor clicks "Faza 3 ga o\'tish."')

doc.add_page_break()

# ============================================================
# ACT 5: PRACTICE LAB
# ============================================================
doc.add_heading('Act 5: Practice Lab', level=1)
doc.add_paragraph('Time: 8:31 — Duration: 6 minutes')

doc.add_heading('Game 1: Tile Match (Reinforce — 3 minutes)', level=2)
doc.add_paragraph(
    'Sardor sees a grid of 12 face-down tiles. He flips two at a time, trying to match writing systems with '
    'their descriptions.'
)

matches = [
    '"Cuneiform" ↔ "Wedge-shaped marks on clay" — Match! +100 XP',
    '"Hieroglyphs" ↔ "Egyptian picture writing" — Match! +100 XP',
    '"Phoenician" ↔ "First alphabet, 22 letters" — Match! +100 XP',
    '"Sogdian" ↔ "Uzbek ancestor\'s writing, 23 letters" — Match! +100 XP',
    '"Avesta" ↔ "Book created in Khorezm" — Match! +100 XP',
    '"Clay tokens" ↔ "Earliest counting system" — Match! +100 XP',
]

for m in matches:
    doc.add_paragraph(m, style='List Bullet')

doc.add_paragraph('All 6 pairs matched in 2:15. Perfect score. +100 bonus XP. Total from this game: 700 XP.')

doc.add_heading('Game 2: Why Chain (Stretch — 3 minutes)', level=2)
doc.add_paragraph(
    'This game targets Sardor\'s weakest skill: "compare information across sources." It is a dialogue — the system '
    'asks "why?" and Sardor types his reasoning.'
)

dialogue = [
    ('System:', '"Nima uchun yozuvning ixtiro qilinishi insoniyat uchun juda muhim bo\'lgan?" '
                '(Why was the invention of writing so important for humanity?)'),
    ('Sardor:', '"Chunki bilimlarni saqlab qolish imkoni paydo bo\'ldi." (Because it became possible to preserve knowledge.)'),
    ('System:', '"Juda yaxshi! Lekin NIMA UCHUN bilimlarni saqlab qolish zarur edi?" '
                '(Very good! But WHY was it necessary to preserve knowledge?)'),
    ('Sardor:', '"Chunki inson xotirasi cheklangan. Bilimlar yo\'qolib ketishi mumkin edi." '
                '(Because human memory is limited. Knowledge could be lost.)'),
    ('System:', '"Ajoyib! Bilimlarni avlodlarga uzatish... Lekin yozuvsiz avlodlar bilimni QANDAY uzatgan?" '
                '(Wonderful! Passing knowledge to generations... But WITHOUT writing, HOW did generations pass knowledge?)'),
    ('Sardor:', '"Faqat og\'zaki — lekin bu noaniq va yo\'qolib ketishi mumkin." '
                '(Only orally — but that is imprecise and can be lost.)'),
    ('System:', '"Mukammal! Siz yozuvning nima uchun \'insoniyatning buyuk kashfiyoti\' deb atalishini '
                'tushuntirib berdingiz!" (Perfect! You explained why writing is called \'humanity\'s great discovery!\')'),
]

for speaker, text in dialogue:
    p = doc.add_paragraph()
    run = p.add_run(f'{speaker} ')
    run.bold = True
    p.add_run(text)

doc.add_paragraph('Sardor completed all 4 levels. +600 XP + 250 bonus for completing the chain. Total: 850 XP.')

doc.add_paragraph('Sardor clicks "Faza 4 ga o\'tish."')

doc.add_page_break()

# ============================================================
# ACT 6: REAL-LIFE CHALLENGE
# ============================================================
doc.add_heading('Act 6: Real-Life Challenge', level=1)
doc.add_paragraph('Time: 14:31 — Duration: 5 minutes')

doc.add_paragraph('The screen shows a scenario Sardor has never seen before:')

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
p.paragraph_format.right_indent = Cm(1)
run = p.add_run(
    '"You are visiting the State Museum of History of Uzbekistan in Tashkent with your family. '
    'There are over 250,000 exhibits. In one display case, you see a clay fragment with ancient Sogdian writing on it. '
    'Next to it, there is a photograph of a cave painting from prehistoric times. Your younger sister asks: '
    '\'Which one is older? Which one tells us more about the people who made it?\'"'
)
run.italic = True

doc.add_paragraph('Two questions follow:')

p = doc.add_paragraph()
run = p.add_run('Question 1: ')
run.bold = True
p.add_run('"Which source is likely older — the clay fragment or the cave painting? Explain your reasoning."')

doc.add_paragraph(
    'Sardor types: "The cave painting is probably older because people drew pictures before they invented writing. '
    'The clay fragment has writing on it, and writing was invented later. Cave paintings are from prehistoric times, '
    'before any writing system existed."'
)

p = doc.add_paragraph()
run = p.add_run('Question 2: ')
run.bold = True
p.add_run('"Which source tells us more about the people who made it? Why?"')

doc.add_paragraph(
    'Sardor types: "The clay fragment tells us more because writing can give specific information — names, dates, '
    'what they traded, what they believed. Cave paintings show us what animals looked like and that people could draw, '
    'but they cannot tell us exact details like writing can. However, cave paintings tell us about their art and '
    'daily life in a way that writing cannot."'
)

doc.add_paragraph('The system scores both responses using keyword matching against the rubric:')

scoring = [
    'Question 1: Correct reasoning (cave paintings predate writing). Full marks. +200 XP.',
    'Question 2: Nuanced answer — recognizes both sources have value, explains why writing is more specific '
    'but art has its own value. Full marks. +200 XP.',
]

for s in scoring:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph('Total from Real-Life: 400 XP. Sardor clicks "Faza 5 ga o\'tish."')

doc.add_page_break()

# ============================================================
# ACT 7: BOSS FIGHT
# ============================================================
doc.add_heading('Act 7: Boss Fight', level=1)
doc.add_paragraph('Time: 19:31 — Duration: 5 minutes')

doc.add_paragraph('The screen darkens. A boss HP bar appears at the top: 100/100.')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Yakuniy Boss — Boss YENGILISHI SHART')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0xef, 0x44, 0x44)

doc.add_heading('Question 1 (Easy — PISA L2, -10 HP)', level=2)
doc.add_paragraph(
    '"Qadimgi yozuv turlarini to\'g\'ri ketma-ketlikda joylashtiring: eslatuvchi belgilar, rasmli yozuv, '
    'so\'z yozuvi, alifboli yozuv. Bu ketma-ketlikda asosiy rivojlanish omili nima?" '
    '(Place ancient writing types in correct order: reminder marks, picture writing, word writing, alphabetic writing. '
    'What was the main driving force in this sequence?)'
)

doc.add_paragraph(
    'Sardor types: "Odamlarga ko\'proq ma\'lumot yozish kerak bo\'ldi va oddiy belgilar yetarli bo\'lmay qoldi. '
    'Hayot murakkablashdi." (People needed to write more information and simple marks were not enough. Life became more complex.)'
)

doc.add_paragraph('Correct. Boss takes 10 damage. HP: 90/100. Combo: 1. +10 XP.')

doc.add_heading('Question 2 (Easy — PISA L2, -10 HP)', level=2)
doc.add_paragraph(
    '"Misr iyerogliflari va Shumer mixxati \'so\'z yozuvi\' turiga kiradi. Bu yozuvlarning asosiy kamchiligi nima edi?" '
    '(Egyptian hieroglyphs and Sumerian cuneiform are \'word writing.\' What was their main disadvantage?)'
)

doc.add_paragraph(
    'Sardor types: "Juda ko\'p belgilarni yodda saqlash kerak edi. Har bir so\'z uchun alohida belgi bor edi, '
    'shuning uchun o\'rganish qiyin edi." (Too many symbols had to be memorized. Each word had its own symbol, '
    'so it was hard to learn.)'
)

doc.add_paragraph('Correct. Boss takes 10 damage. HP: 80/100. Combo: 2. +10 XP.')

doc.add_heading('Question 3 (Medium — PISA L3, -20 HP)', level=2)
doc.add_paragraph(
    '"Finikiy alifbosi 22 ta harfdan iborat bo\'lgan. Bu nima uchun \'inqilob\' hisoblanadi?" '
    '(The Phoenician alphabet had 22 letters. Why is this considered a \'revolution\'?)'
)

doc.add_paragraph(
    'Sardor types: "Chunki 22 ta harf bilan har qanday so\'zni yozish mumkin edi. Minglab belgilarni yodlash '
    'shart emas edi. Ko\'proq odamlar yozishni o\'rgana oldi." (Because with 22 letters you could write any word. '
    'You did not need to memorize thousands of symbols. More people could learn to write.)'
)

doc.add_paragraph('Correct. Boss takes 20 damage. HP: 60/100. Combo: 3 → COMBO ACTIVATED! Next correct = 2x damage. +20 XP.')

doc.add_heading('Question 4 (Hard — PISA L4, -30 HP × 2 = -60 HP)', level=2)
doc.add_paragraph(
    '"Tasavvur qiling: yozuv hech ixtiro qilinmaganida, bugungi dunyodagi ilm-fan, texnologiya va madaniyat '
    'qanday bo\'lardi? Asoslang." (Imagine: if writing had never been invented, what would science, technology, '
    'and culture be like today? Justify.)'
)

doc.add_paragraph(
    'Sardor types: "Bilimlar faqat og\'zaki uzatilardi. Ko\'p narsa yo\'qolib ketardi. Har bir avlod noldan '
    'boshlashi kerak bo\'lardi. Ilm-fan sekin rivojlanardi, chunki hech kim oldingi olimlarning ishlarini yozib '
    'qoldirmagan bo\'lardi. Texnologiya bo\'lmas edi, chunki ko\'rsatmalarni yozib bo\'lmasdi." '
    '(Knowledge would only be passed orally. Much would be lost. Each generation would have to start from scratch. '
    'Science would develop slowly because no one could write down previous scientists\' work. There would be no '
    'technology because instructions could not be written down.)'
)

doc.add_paragraph(
    'Correct. Combo multiplier activates: 30 × 2 = 60 damage. Boss HP: 0/100. BOSS DEFEATED!'
)

doc.add_paragraph('Star calculation:')

stars = [
    'Attempt 1: Yes',
    'Hints used: 0',
    'HP remaining when defeated: 0 (but all questions answered correctly)',
    'Result: 3 stars — Boss defeated on first attempt, no hints, all correct.',
]

for s in stars:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph('Boss XP: 1000 XP (3-star rating).')

doc.add_paragraph('Sardor clicks "Faza 6 ga o\'tish."')

doc.add_page_break()

# ============================================================
# ACT 8: REFLECTION
# ============================================================
doc.add_heading('Act 8: Reflection', level=1)
doc.add_paragraph('Time: 24:31 — Duration: 2 minutes')

doc.add_paragraph('The screen softens. A prompt appears, selected based on Sardor\'s 83% accuracy:')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('"Bugun qaysi strategiya sizga eng ko\'p yordam berdi? Nima uchun bu siz uchun ishadi?"')
run.font.size = Pt(13)
run.font.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('("What strategy helped you the most today? Why did it work for you?")')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.italic = True

doc.add_paragraph('Sardor types:')

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(
    '"Tile Match o\'yini menga eng ko\'p yordam berdi. Chunki men vizual o\'rganuvchiman — rasmlar va '
    'belgilarni ko\'rganda yaxshiroq eslab qolaman. Nima uchun ishadi? Chunki men juftliklarni topishda '
    'har bir yozuv turini ikki marta ko\'rdim — bir marta nomini, bir marta ta\'rifini. Bu ikki marta '
    'takrorlash bo\'ldi."'
)
run.italic = True

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(
    '("Tile Match helped me the most. Because I am a visual learner — I remember better when I see pictures '
    'and symbols. Why did it work? Because when matching pairs, I saw each writing type twice — once its name, '
    'once its definition. That was double repetition.")'
)
run.italic = True

doc.add_paragraph('Character count: 287 (minimum 10 met). Sardor clicks "Yuborish." +100 XP.')

doc.add_page_break()

# ============================================================
# ACT 9: RESULTS
# ============================================================
doc.add_heading('Act 9: Results Screen', level=1)
doc.add_paragraph('Time: 26:31')

doc.add_paragraph('Sardor sees his final results:')

results_display = [
    'Total XP earned: 4,290 XP (including streak multiplier approaching)',
    'Stars: 3 out of 3',
    'Accuracy: 83% overall',
    'PISA level change: Reading 1.4 → 1.5',
    'Streak: 5 days (2 more days until 7-day bonus)',
    'Mastery: 0.68 → 0.72',
]

for r in results_display:
    doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph('Sardor closes the app. He is done. But the system is not.')

doc.add_page_break()

# ============================================================
# ACT 10: BEHIND THE SCENES
# ============================================================
doc.add_heading('Act 10: What Happens After Sardor Closes the App', level=1)

doc.add_heading('1. PISA Profile Update', level=2)
doc.add_paragraph(
    'The system recalculates Sardor\'s PISA level. He scored 83% overall, with particularly strong performance '
    'on Apply and Analyze items. His Reading PISA moves from 1.4 to 1.5. His "interpret" domain improves from '
    '1.3 to 1.4. The transition skill "L1→L2: compare info across simple sources" is marked as mastered.'
)

doc.add_heading('2. Teacher Report Generated', level=2)
doc.add_paragraph('Ms. Barno receives this report in her dashboard:')

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(
    'Student: Sardor Karimov | Date: April 3, 2026\n'
    'Topic: History 5, Section 10 — "Yozuv tarixi"\n'
    'Overall accuracy: 83% | PISA change: +0.1 | Stars: 3/3\n\n'
    'Strong: UZ-HIST-5-WRITING-01 (writing types), UZ-HIST-5-WRITING-02 (chronology)\n'
    'Needs work: None identified this session.\n\n'
    'Note: Sardor is approaching his 7-day streak. Consider recognizing this in class.\n'
    'Class insight: 4 out of 28 students struggled with the Real-Life scenario on source comparison. '
    'Recommend a 5-minute class discussion on evaluating historical sources.'
)
run.font.size = Pt(10)

doc.add_heading('3. Parent Summary Sent', level=2)
doc.add_paragraph('Mr. Karimov receives this on his phone:')

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(
    '"Sardor bugungi tarix uy vazifasini a\'lo darajada tugatdi. 3 yulduz oldi. '
    'Qadimgi yozuv turlarini juda yaxshi o\'zlashtirdi. Uyda mashq: uyingizdagi turli yozuvlarni birga toping — '
    'kitobdagi lotin harflar, eski hujjatlardagi kirill yozuvi, pullar ustidagi raqamlar. '
    'Sardordan har birini tasniflashni so\'rang."'
)
run.font.size = Pt(10)
run.italic = True

doc.add_heading('4. Feed-Forward to Next Session', level=2)
doc.add_paragraph('The system updates Sardor\'s profile for next time:')

feed_forward = [
    'No weak standards identified — all standards scored above 0.8.',
    'Spaced repetition: All 6 Memory Sprint items marked as "correct" — next review in 7 days.',
    'Memory Palace items: 4 out of 5 correct — the missed item ("Uzbek alphabet — 29 letters") '
    'will appear in the next session\'s Memory Sprint.',
    'Next session topic announced: Section 11 — "Avesto — ajdodlarimiz yaratgan ilk yozma tarixiy manba."',
    'Streak counter: 5 days. 2 more days until the 7-day bonus (+50% XP multiplier).',
]

for item in feed_forward:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================
# SESSION SUMMARY TABLE
# ============================================================
doc.add_heading('Session Summary', level=1)

table = doc.add_table(rows=10, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Phase', 'Time', 'Duration', 'XP Earned']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)

data = [
    ['Pre-Session', '0:00', '<1 sec', '—'],
    ['Engagement Hook', '0:01', '30 sec', '40'],
    ['Activation', '0:31', '2:00', '540'],
    ['Story Mode', '2:31', '5:00', '350'],
    ['Practice Lab', '8:31', '6:00', '1,550'],
    ['Real-Life', '14:31', '5:00', '400'],
    ['Boss Fight', '19:31', '5:00', '1,000'],
    ['Reflection', '24:31', '2:00', '100'],
    ['TOTAL', '', '25:30', '4,290'],
]

for row_idx, row_data in enumerate(data):
    for col_idx, cell_text in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
                if row_idx == 8:
                    run.bold = True

doc.add_page_break()

# ============================================================
# DESIGN NOTES
# ============================================================
doc.add_heading('Design Notes', level=1)

doc.add_heading('Why This Flow Works', level=2)

notes = [
    ('The Engagement Hook is not assessed.',
     'It exists purely for psychological priming. Starting with easy wins puts the student in a positive '
     'mindset before cognitive work begins. Research: Csikszentmihalyi (1990) — flow state entry.'),
    ('Memory Sprint uses text input, not multiple choice.',
     'For Grade 5+, multiple choice allows guessing. Text input forces genuine recall. Research: '
     'Roediger & Karpicke (2006) — active recall is 3x more effective than recognition.'),
    ('Story Mode has 2 segments, not 3.',
     'Tighter pacing reduces cognitive load. The Memory Palace lock-in at the end captures retention at its '
     'peak — immediately after learning. Research: Maguire et al. (2003) — +300% retention.'),
    ('Practice Lab has 2 games, not 3.',
     'Two interleaved games deliver 90% of the benefit with 33% less time and less context switching. '
     'Research: Rohrer & Taylor (2007) — diminishing returns after 2 mixed sets.'),
    ('Real-Life is open response, not multiple choice.',
     'Transfer cannot be measured with multiple choice. If a student can answer by repeating a textbook '
     'definition, it is not transfer. Open response forces genuine reasoning.'),
    ('Boss has 4 questions, not 6.',
     'Item Response Theory shows 4 well-calibrated items at tiered difficulty produce the same measurement '
     'reliability as 6. Research: Lord (1980).'),
    ('Boss failure triggers Socratic tutoring, not punishment.',
     'Students who fail are guided through what they got wrong with pre-written hints, then given different '
     'questions to retry. Max 2 attempts, then partial credit. Research: Stigler & Hiebert (1999) — '
     'productive struggle, not productive punishment.'),
    ('Reflection is private.',
     'The student\'s response is visible only to them. The teacher sees aggregate themes. This encourages '
     'honest self-assessment. Research: Flavell (1979) — +20-30% next session performance.'),
]

for title_text, body_text in notes:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text} ')
    run.bold = True
    p.add_run(body_text)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Proprietary — Ministry of Education, Republic of Uzbekistan')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

# ============================================================
# SAVE
# ============================================================
output_path = r'C:\Users\Recruiter\Documents\Class A EDU\claude analysis\Flow Recommended.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
