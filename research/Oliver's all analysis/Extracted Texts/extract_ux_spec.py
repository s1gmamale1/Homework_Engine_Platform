"""Extract full text content from NETS-UI-UX-Design-Spec.docx."""
from docx import Document
import os

SOURCE = r"C:\Users\Recruiter\Documents\Class A EDU\All analysis\Universal Specs\NETS-UI-UX-Design-Spec.docx"
OUTPUT = r"C:\Users\Recruiter\Documents\Class A EDU\All analysis\Extracted Texts\_ui_ux_design_spec_extracted.txt"

def extract_text(filepath):
    doc = Document(filepath)
    lines = []

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        if "Heading" in style:
            level = style.replace("Heading ", "").replace("Heading", "1")
            try:
                level = int(level)
            except:
                level = 1
            prefix = "#" * level
            lines.append(f"{prefix} {text}")
        elif style == "Title":
            lines.append(f"# {text}")
        else:
            lines.append(text)

    for i, table in enumerate(doc.tables):
        lines.append(f"\n--- TABLE {i+1} ---")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " | ") for cell in row.cells]
            lines.append(" | ".join(cells))
        lines.append(f"--- END TABLE {i+1} ---\n")

    return "\n".join(lines)

size = os.path.getsize(SOURCE)
text = extract_text(SOURCE)
with open(OUTPUT, "w", encoding="utf-8") as f:
    f.write(text)
print(f"Extracted: {size} bytes source, {len(text)} chars -> {OUTPUT}")
