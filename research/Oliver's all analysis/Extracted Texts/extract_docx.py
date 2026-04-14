"""Extract full text content from 3 .docx files for comparison."""
from docx import Document
import os

FILES = {
    "THEIR_BLUEPRINT": r"C:\Users\Recruiter\Downloads\Telegram Desktop\NETS-Homework-Engine-Blueprint.docx",
    "THEIR_HISTORY": r"C:\Users\Recruiter\Downloads\Telegram Desktop\Grade5-History-S1-Homework-Design.docx",
    "OUR_BLUEPRINT": r"C:\Users\Recruiter\Documents\Class A EDU\claude analysis\NETS-Homework-Engine-Blueprint-Complete.docx",
}

OUTPUT_DIR = r"C:\Users\Recruiter\Documents\Class A EDU\claude analysis"

def extract_text(filepath):
    doc = Document(filepath)
    lines = []

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        if "Heading" in style:
            level = style.replace("Heading ", "").replace("Heading", "1")
            try:
                level = int(level)
            except:
                level = 1
            prefix = "#" * level
            lines.append(f"{prefix} {text}")
        elif style == "Title":
            lines.append(f"# {text}")
        else:
            lines.append(text)

    for i, table in enumerate(doc.tables):
        lines.append(f"\n--- TABLE {i+1} ---")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " | ") for cell in row.cells]
            lines.append(" | ".join(cells))
        lines.append(f"--- END TABLE {i+1} ---\n")

    return "\n".join(lines)


for label, path in FILES.items():
    size = os.path.getsize(path)
    text = extract_text(path)
    outpath = os.path.join(OUTPUT_DIR, f"_extracted_{label}.txt")
    with open(outpath, "w", encoding="utf-8") as f:
        f.write(text)
    print(f"{label}: {size} bytes source, {len(text)} chars extracted -> {outpath}")
