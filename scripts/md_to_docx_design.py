"""Convert NETS-UI-UX-Design-Spec.md to a docx with proper heading/table/code styling."""
import os, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = os.path.join(os.path.dirname(__file__), "..", "standards", "NETS-UI-UX-Design-Spec.md")
OUT = os.path.join(os.path.dirname(__file__), "..", "standards", "NETS-UI-UX-Design-Spec.docx")

with open(SRC, "r", encoding="utf-8") as f:
    lines = f.read().splitlines()

doc = Document()

# Tweak normal style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def add_code_block(code_lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    run = p.add_run("\n".join(code_lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_table(rows):
    # rows: list of list of cells (first row = header)
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = text.strip()
            if r_idx == 0:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True

def parse_inline(text):
    """Strip markdown bold/italic/code markers and return plain text segments with formatting hints."""
    segments = []  # list of (text, bold, italic, code)
    i = 0
    while i < len(text):
        if text.startswith("**", i):
            end = text.find("**", i + 2)
            if end != -1:
                segments.append((text[i + 2:end], True, False, False))
                i = end + 2
                continue
        if text.startswith("`", i):
            end = text.find("`", i + 1)
            if end != -1:
                segments.append((text[i + 1:end], False, False, True))
                i = end + 1
                continue
        # plain char — accumulate until next marker
        next_bold = text.find("**", i)
        next_code = text.find("`", i)
        candidates = [x for x in [next_bold, next_code] if x != -1]
        stop = min(candidates) if candidates else len(text)
        segments.append((text[i:stop], False, False, False))
        i = stop
    return segments

def add_formatted_paragraph(text, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    for seg_text, bold, italic, code in parse_inline(text):
        if not seg_text:
            continue
        run = p.add_run(seg_text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if code:
            run.font.name = "Consolas"
            run.font.size = Pt(10)
    return p

# ─── Walk lines, detect blocks ───
i = 0
n = len(lines)
pending_table = []

def flush_table():
    global pending_table
    if not pending_table:
        return
    # First row is header, second row is separator (|---|---|), rest are data
    rows = []
    for idx, line in enumerate(pending_table):
        if idx == 1:  # separator row
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    add_table(rows)
    pending_table = []

while i < n:
    line = lines[i]
    stripped = line.rstrip()

    # Table rows
    if stripped.startswith("|") and "|" in stripped[1:]:
        pending_table.append(stripped)
        i += 1
        continue
    else:
        if pending_table:
            flush_table()

    # Code block fenced
    if stripped.startswith("```"):
        i += 1
        buf = []
        while i < n and not lines[i].startswith("```"):
            buf.append(lines[i])
            i += 1
        add_code_block(buf)
        i += 1
        continue

    # Headings
    m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
    if m:
        level = len(m.group(1))
        text = m.group(2).strip()
        if level == 1:
            doc.add_heading(text, level=0)
        else:
            doc.add_heading(text, level=min(level - 1, 5))
        i += 1
        continue

    # Horizontal rule
    if stripped.strip() == "---":
        doc.add_paragraph("_" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue

    # Bullet list
    if re.match(r"^\s*[-*]\s+", stripped):
        text = re.sub(r"^\s*[-*]\s+", "", stripped)
        add_formatted_paragraph(text, style="List Bullet")
        i += 1
        continue

    # Numbered list
    if re.match(r"^\s*\d+\.\s+", stripped):
        text = re.sub(r"^\s*\d+\.\s+", "", stripped)
        add_formatted_paragraph(text, style="List Number")
        i += 1
        continue

    # Blockquote
    if stripped.startswith(">"):
        text = stripped.lstrip("> ").strip()
        p = add_formatted_paragraph(text)
        p.paragraph_format.left_indent = Pt(24)
        for run in p.runs:
            run.italic = True
        i += 1
        continue

    # Blank line
    if not stripped.strip():
        i += 1
        continue

    # Regular paragraph
    add_formatted_paragraph(stripped)
    i += 1

if pending_table:
    flush_table()

doc.save(OUT)
print(f"Saved: {OUT}")
