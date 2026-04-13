"""Generate NETS-Homework-Engine-Universal-Blueprint-Flow-Diagram.docx with embedded
whiteboard image, ASCII flow diagram, and reference tables."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

STD = os.path.join(os.path.dirname(__file__), "..", "standards")
IMG = os.path.join(STD, "Flow_diagram.jpg")
OUT = os.path.join(STD, "NETS-Homework-Engine-Universal-Blueprint-Flow-Diagram.docx")

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

doc.add_heading("NETS Homework Engine — Blueprint Flow Diagram", level=0)

p = doc.add_paragraph()
p.add_run("Date: ").bold = True
p.add_run("2026-04-07")
p = doc.add_paragraph()
p.add_run("Source diagram: ").bold = True
p.add_run("standards/Flow_diagram.jpg (whiteboard sketch)")
p = doc.add_paragraph()
p.add_run("Companion to: ").bold = True
p.add_run("NETS-Homework-Engine-Universal-Blueprint-Summary.docx")

doc.add_paragraph(
    "This document is the diagram-first version of the Blueprint summary. "
    "Each block in the flow corresponds to a phase or sub-phase in the student journey. "
    "Read top-to-bottom, then left-to-right where panels split."
)

# ─── Whiteboard image ───
doc.add_heading("Original Whiteboard Sketch", level=1)
doc.add_paragraph(
    "The hand-drawn flow diagram below was the source for this document. "
    "All blocks in the digitized flow that follows trace directly to elements on this whiteboard."
)
try:
    doc.add_picture(IMG, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    doc.add_paragraph(f"[Image could not be embedded: {e}]")

# ─── ASCII flow diagram ───
doc.add_heading("Digitized Flow Diagram", level=1)
doc.add_paragraph(
    "The same flow, formalized as a vertical block chart. Each block is named "
    "after its whiteboard counterpart and tagged with its phase code (P1-P7) where applicable."
)

ASCII = r"""
═══════════════════════════════════════════════════════════════════
                       PRE-HOMEWORK BLOCK
═══════════════════════════════════════════════════════════════════

   ┌──────────┐      ┌──────────┐      ┌──────────────┐
   │ OPENING  │ ───▶ │ LOADING  │ ───▶ │ "Are you OK? │
   │  (tap)   │      │ SCREEN   │      │  Learn how"  │
   └──────────┘      └──────────┘      └──────┬───────┘
                                              │
                                              ▼
   ┌─────────────────────────────────────────────────────────┐
   │            THEME PREVIEW  ─  8 PANELS                   │
   │            (swipe left/right · whoosh transition)       │
   ├─────────┬─────────┬─────────┬─────────┬─────────┬───────┤
   │ 1       │ 2       │ 3       │ 4       │ 5       │ ...   │
   │ Summary │ Refined │Examples │Research │  Refs   │       │
   │ of book │ Content │         │         │         │       │
   │ content │         │         │         │         │       │
   │   ◀ ▶   │   ◀ ▶   │   ◀ ▶   │   ◀ ▶   │   ◀ ▶   │       │
   ├─────────┼─────────┼─────────┴─────────┴─────────┴───────┤
   │ 6       │ 7       │ 8                                   │
   │  Why    │Industry │ Reference                           │
   │  this   │  apps   │   Links                             │
   │ matters │         │ (any lang)                          │
   │   ◀ ▶   │   ◀ ▶   │   ◀ ▶                               │
   └─────────┴─────────┴─────────────────────────────────────┘
                                              │
                                whoosh effect transition
                                              │
                                              ▼
                              ┌────────────────────┐
                              │    FLASH CARDS     │
                              │  ┌──┐ ┌──┐ ┌──┐    │
                              │  │  │ │  │ │  │    │
                              │  └──┘ └──┘ └──┘    │
                              │  3D circle carousel│
                              │  · tap to flip     │
                              │  · swipe to rotate │
                              └─────────┬──────────┘
                                        │
                                        ▼
                              ┌────────────────────┐
                              │ ▶ START MY HOMEWORK│
                              └─────────┬──────────┘
                                        │
                                        ▼
═══════════════════════════════════════════════════════════════════
                    HOMEWORK ENGINE  ·  7 PHASES
═══════════════════════════════════════════════════════════════════

                              ┌────────────────────┐
                              │   LOADING SCREEN   │
                              │  (quote · fact ·   │
                              │    corny line)     │
                              └─────────┬──────────┘
                                        │
                                        ▼
                              ┌────────────────────┐    P1
                              │   MEMORY SPRINT    │
                              │ ≤2 min · flexible  │
                              │ MC / Speed Match / │
                              │ Flash Sprint /     │
                              │ Fill-Blanks / Order│
                              └─────────┬──────────┘
                                        │
                              "did you know..."
                                        │
                                        ▼
                              ┌────────────────────┐    P2
                              │     STORY MODE     │
                              │  Problem ▶ Struggle│
                              │  ▶ Discovery ▶     │
                              │      Solution      │
                              │  3 segments, IELTS │
                              │  comprehension     │
                              │  Stranger Test gate│
                              └─────────┬──────────┘
                                        │
                                        ▼
                              ┌────────────────────┐    P3
                              │    GAME BREAKS     │
                              │  interleaved with  │
                              │       Story        │
                              ├────────────────────┤
                              │  ┌──────────────┐  │
                              │  │   GAME 1     │  │ ← reinforce
                              │  │ (Default 16) │  │
                              │  └──────┬───────┘  │
                              │         │          │
                              │  ┌──────▼───────┐  │
                              │  │   GAME 2     │  │ ← stretch
                              │  │ (Default 16) │  │
                              │  └──────┬───────┘  │
                              │         │          │
                              │  ┌──────▼───────┐  │
                              │  │   GAME 3     │  │ ← transition skill
                              │  │ (Interactive │  │
                              │  │  Catalog)    │  │
                              │  └──────────────┘  │
                              │ Rule: >=1 Catalog │
                              │ + >=2 Default 16  │
                              └─────────┬──────────┘
                                        │
                              "did you know..."
                                        │
                                        ▼
                              ┌────────────────────┐    P4
                              │  REAL-LIFE APP     │
                              │ "You are the       │
                              │   expert..."       │
                              │ first-person POV   │
                              │ tricky distractors │
                              │ explain your call  │
                              └─────────┬──────────┘
                                        │
                                        ▼
                              ┌────────────────────┐    P5
                              │   CONSOLIDATION    │
                              │   Memory Palace    │
                              │   mnemonic lock    │
                              │  "I'm ready"  ★    │
                              └─────────┬──────────┘
                                        │
                              "did you know..."
                                        │
                                        ▼
                              ┌────────────────────┐    P6
                              │   SOLO WITH AI     │
                              │   ──────────────   │
                              │   ⚔  BOSS FIGHT    │
                              ├────────────────────┤
                              │  Sub Boss   (per)  │
                              │  Big Boss   (week) │
                              │  Mythic    (<5%)   │
                              │                    │
                              │  HP · combo · star │
                              │  hint = boss heal  │
                              └─────────┬──────────┘
                                        │
                                        ▼
                              ┌────────────────────┐    P7
                              │     REFLECTION     │
                              │   metacognition    │
                              │   private journal  │
                              └─────────┬──────────┘
                                        │
                                        ▼
═══════════════════════════════════════════════════════════════════
                  POST-SESSION  ·  AI ANALYSIS BLOCK
                       (server-side, <10 sec)
═══════════════════════════════════════════════════════════════════

                              ┌────────────────────┐
                              │  DATA EXTRACTION   │
                              │  per phase / Bloom │
                              │  / standard /      │
                              │  transition skill  │
                              └─────────┬──────────┘
                                        │
                                        ▼
                              ┌────────────────────┐
                              │   DATA ANALYSIS    │
                              │  root-cause find   │
                              │  WHY not just WHAT │
                              └─────────┬──────────┘
                                        │
                                        ▼
                              ┌────────────────────┐
                              │  DATA COMPARISON   │
                              │  vs prior sessions │
                              │  trajectory check  │
                              └─────────┬──────────┘
                                        │
                                        ▼
                              ┌────────────────────┐
                              │    ACTION PLAN     │
                              │  next session map  │
                              │  weak-area queue   │
                              │  spaced repetition │
                              └─────────┬──────────┘
                                        │
                                        ▼
                              ┌────────────────────┐
                              │     AI UPDATE      │
                              │  PISA recalc (IRT) │
                              │  mastery map       │
                              └─────────┬──────────┘
                                        │
                                        ▼
                              ┌────────────────────┐
                              │      REVISION      │
                              │  Duolingo Mode if  │
                              │     score < 60%    │
                              │  Repass option if  │
                              │      60-79%        │
                              └─────────┬──────────┘
                                        │
                                        ▼
                              ┌────────────────────┐
                              │       MEMORY       │
                              │  persisted to      │
                              │  student profile + │
                              │  spaced repetition │
                              │  queue (SM-2)      │
                              └─────────┬──────────┘
                                        │
                                        ▼
                              ┌────────────────────┐
                              │  TEACHER + PARENT  │
                              │     REPORTS        │
                              │  generated · sent  │
                              └────────────────────┘
"""

p = doc.add_paragraph()
run = p.add_run(ASCII)
run.font.name = "Consolas"
run.font.size = Pt(8)

# ─── Block-by-block reference tables ───
doc.add_heading("Block-by-Block Reference", level=1)

doc.add_heading("Pre-Homework Block", level=2)
t = doc.add_table(rows=6, cols=4)
t.style = "Light Grid Accent 1"
hdr = ["Block", "What it does", "Duration", "Scoring"]
for i, h in enumerate(hdr):
    t.rows[0].cells[i].text = h
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
rows = [
    ("Opening", "Student taps the homework assignment", "< 1s", "None"),
    ("Loading Screen", "First quote/fact/corny-line card while session loads", "2-4s", "None"),
    ("Theme Preview", "8 swipeable panels with whoosh + letter-assembly transition", "2-3 min, paced", "None"),
    ("Flash Cards", "3D circular carousel of formulas/concepts/rules — tap to flip", "1-2 min, paced", "None"),
    ("Start My Homework", "Explicit gateway button — tapping it starts the session timer", "tap", "Timer starts here"),
]
for r, row in enumerate(rows):
    for c, val in enumerate(row):
        t.rows[r + 1].cells[c].text = val

doc.add_heading("Homework Engine Block (7 phases)", level=2)
t = doc.add_table(rows=8, cols=4)
t.style = "Light Grid Accent 1"
hdr = ["Phase", "Block", "Key Visual / Behavior", "Duration"]
for i, h in enumerate(hdr):
    t.rows[0].cells[i].text = h
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
rows = [
    ("P1", "Memory Sprint", "Flexible-format pool, fast pulse animations", "≤2 min"),
    ("P2", "Story Mode", "Cinematic parallax, story arc Problem→Struggle→Discovery→Solution", "5-7 min"),
    ("P3", "Game Breaks", "3 games interleaved with story segments (Game 1 → Story 2 → Game 2 → Story 3 → Game 3)", "6-9 min"),
    ("P4", "Real-Life Application", "Case file aesthetic, first-person expert POV, rubber-stamp verdict", "4-6 min"),
    ("P5", "Consolidation", "Isometric Memory Palace room", "2-3 min"),
    ("P6", "Solo with AI (Boss)", "Sub / Big / Mythical tier, HP system, combo, hint-tax", "6-12 min"),
    ("P7", "Reflection", "Quiet journal, paper texture, low ambient sound", "1-2 min"),
]
for r, row in enumerate(rows):
    for c, val in enumerate(row):
        t.rows[r + 1].cells[c].text = val

doc.add_paragraph(
    "Loading screens (Did You Know / quote / corny line) appear between every phase "
    "transition AND at one deliberate ~60% mid-session breath moment."
)

doc.add_heading("Post-Session Block", level=2)
t = doc.add_table(rows=9, cols=3)
t.style = "Light Grid Accent 1"
hdr = ["Block", "What it does", "Visible to Student?"]
for i, h in enumerate(hdr):
    t.rows[0].cells[i].text = h
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
rows = [
    ("Data Extraction", "Pulls per-phase, per-Bloom, per-standard, per-transition-skill performance", "No"),
    ("Data Analysis", "Root-cause WHY (not just WHAT was wrong)", "No"),
    ("Data Comparison", "vs prior sessions, trajectory check (3+ declines → Boost Mode)", "No"),
    ("Action Plan", "Next session content map, weak-area queue, SM-2 spaced repetition update", "No"),
    ("AI Update", "PISA level recalc (IRT + exponential decay), mastery map update", "No"),
    ("Revision", "Duolingo Mode loop if <60%, Repass offer if 60-79%", "YES (if triggered)"),
    ("Memory", "Persisted to student profile and spaced repetition queue", "No"),
    ("Teacher / Parent Reports", "Detailed report to teacher, simple-language report to parent", "YES (delivered async)"),
]
for r, row in enumerate(rows):
    for c, val in enumerate(row):
        t.rows[r + 1].cells[c].text = val

# ─── Cross-references ───
doc.add_heading("Cross-References", level=1)
for ref in [
    "Whoosh source: standards/Flow_diagram.jpg",
    "Full Blueprint: standards/NETS-Homework-Engine-Universal-Blueprint.docx",
    "Full UNIFIED spec: standards/NETS-Homework-Engine-UNIFIED.md",
    "UI/UX & Animation Spec: standards/NETS-UI-UX-Design-Spec.md (loading screens, panel transitions, carousel, per-phase animations)",
    "Improvements log: standards/IMPROVEMENTS_TO_THE_CURRENT_FRAMEWORKS.md",
]:
    doc.add_paragraph(ref, style="List Bullet")

doc.save(OUT)
print(f"Saved: {OUT}")
