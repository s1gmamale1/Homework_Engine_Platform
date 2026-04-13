"""
Update NETS-Homework-Engine-Blueprint.docx in place with the
modifications from IMPROVEMENTS_TO_THE_CURRENT_FRAMEWORKS.md (2026-04-07).

Strategy: load existing docx, locate section headings, insert new
paragraphs immediately BEFORE the next section heading so the new
content lands at the end of the section it belongs to.
"""

import os
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = os.path.join(
    os.path.dirname(__file__), "..", "standards",
    "NETS-Homework-Engine-Blueprint.docx",
)

doc = Document(PATH)
body = doc.element.body


# ───────────────────────── helpers ─────────────────────────

def find_paragraph_by_text(text_substring, style_prefix=None):
    """Return the first paragraph whose text contains text_substring."""
    for p in doc.paragraphs:
        if style_prefix and not p.style.name.startswith(style_prefix):
            continue
        if text_substring in p.text:
            return p
    return None


def insert_para_before(anchor_p, text, style_name="Normal", bold=False):
    """Insert a new paragraph immediately before anchor_p."""
    new_p = OxmlElement("w:p")
    anchor_p._p.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, anchor_p._parent)
    try:
        para.style = doc.styles[style_name]
    except KeyError:
        pass
    if text:
        run = para.add_run(text)
        if bold:
            run.bold = True
    return para


def insert_block_before(anchor_p, blocks):
    """blocks = list of (style_name, text) tuples. Inserts in order."""
    for style, text in blocks:
        insert_para_before(anchor_p, text, style_name=style)


# ─────────────── Build the new content blocks ───────────────

PRE_HOMEWORK_INTRO = [
    ("Heading 1", "PRE-HOMEWORK SESSIONS — Theme Preview + Flash Cards (NEW 2026-04-07)"),
    ("Normal", "Per IMPROVEMENTS_TO_THE_CURRENT_FRAMEWORKS.md, the student journey now begins BEFORE the 7-phase Homework Engine. Two new sessions sit between opening the assignment and pressing the \"Start my Homework\" button. Both are student-paced, optional but recommended, and do NOT contribute to scoring or XP."),
]

SECTION_0A = [
    ("Heading 1", "SECTION 0-A: THEME PREVIEW (NEW)"),
    ("Heading 2", "Position and Purpose"),
    ("Normal", "Position: First thing the student sees when opening a homework assignment, after the invisible Pre-Session load (Section 0)."),
    ("Normal", "Purpose: Hook the student. Connect the topic to their world. Explain why it matters. Equip the mental model BEFORE any practice begins. This is the bridge between reading the textbook and doing the gamified homework — the place where engagement was previously dying."),
    ("Normal", "Duration: 2-3 minutes, student-paced. No timer. No pressure. No XP."),
    ("Heading 2", "The 8 Required Components"),
    ("Normal", "1. Summary of Book Content — a cleaner reframing of the chapter, plain language, better visual hierarchy. Not a rewrite, a refocusing."),
    ("Normal", "2. Better Explanation — concepts the textbook assumes the student already understands, clarified. Fills the gap between textbook language and student comprehension."),
    ("Normal", "3. Examples — additional worked examples beyond the textbook, step-by-step."),
    ("Normal", "4. Real-Life Research — the origin story of the concept. Who discovered it? What problem did it solve? What changed in the world?"),
    ("Normal", "5. Personal Hook — first-person address connecting topic to the student's own life. \"Remember when you saw...?\" / \"Maybe you've wondered why...?\""),
    ("Normal", "6. Why This Matters — explicit statement of real-world relevance, not \"it's on the exam\"."),
    ("Normal", "7. Industry Application — how the concept is used in real industries, jobs, and technology."),
    ("Normal", "8. Additional Materials — curated external resources (videos, articles, simulators). Language does NOT matter. English, Russian, or Uzbek are all permitted. Student picks."),
    ("Heading 2", "Design Rules"),
    ("Normal", "No quiz, no pressure. Exploration mode. There are no answers to give, so there are no wrong answers."),
    ("Normal", "Student-paced. No timer. Student scrolls/clicks at their own speed and may skip components."),
    ("Normal", "Visual-first. Diagrams, photos, short video clips preferred over walls of text."),
    ("Normal", "First-person POV throughout. \"You\" not \"the student\". Direct address only, never third-person."),
    ("Normal", "Ends with explicit transition to Flash Cards: \"Before we start — here are the key ideas you'll need.\""),
]

SECTION_0B = [
    ("Heading 1", "SECTION 0-B: FLASH CARDS (NEW)"),
    ("Heading 2", "Position and Purpose"),
    ("Normal", "Position: Between Theme Preview (Section 0-A) and the 7-phase Homework Engine."),
    ("Normal", "Purpose: Surface every formula, concept, rule, and definition the student will need during the homework. The last reference moment before the challenge begins. Doubles as a quick-access reference deck DURING the homework."),
    ("Normal", "Duration: 1-2 minutes for the initial pass. Available throughout the homework on demand."),
    ("Heading 2", "What the Card Pool Must Contain"),
    ("Normal", "All Key Formulas — every formula, equation, or rule the student will encounter in the homework. One per card."),
    ("Normal", "All Key Concepts — definitions, terminology, classifications. \"What is an oxide?\" / \"What does valence mean?\""),
    ("Normal", "All Key Rules — conditions, exceptions, boundaries. \"Combustion requires 3 things: fuel, oxygen, heat. If any one is missing, fire stops.\""),
    ("Heading 2", "Design Rules"),
    ("Normal", "One concept per card. No cramming."),
    ("Normal", "Swipeable / scrollable. Student flips through at their own pace."),
    ("Normal", "No testing. Reference, not assessment. No right/wrong."),
    ("Normal", "Returnable during homework. A \"Flash Cards\" button is available throughout the session as a quick-access overlay. NOT a hint (no penalty), just a reference."),
    ("Normal", "Ends with the explicit \"Start my Homework\" button. Only after this press does the session timer start and XP accrue."),
]

MEMORY_SPRINT_UPDATE = [
    ("Heading 2", "1.4 Memory Sprint — Format Flexibility (UPDATED 2026-04-07)"),
    ("Normal", "The original spec said \"single-answer recall\" — too narrow. All four AI demos produced typed quiz lists. The new rule: Memory Sprint can be ANY fast-paced format that hits prior chapters in ≤2 minutes and delivers quick dopamine."),
    ("Normal", "Approved formats:"),
    ("Normal", "  • Quick MC / Binary — 5-8 tap-to-answer questions"),
    ("Normal", "  • Speed Match — drag-and-drop term ↔ definition / formula ↔ result"),
    ("Normal", "  • Flash Sprint — rapid-fire flashcards, swipe Know / Review"),
    ("Normal", "  • Fill-in-Blanks Race — complete sentences with missing keywords, timed"),
    ("Normal", "  • Order the Steps — arrange procedure steps in correct order"),
    ("Normal", "Content creator chooses the format that fits the subject. Mixing formats within a Sprint is allowed. All formats must have a Start button, instant per-item feedback, and a final score display."),
]

STORY_MODE_UPDATE = [
    ("Heading 2", "2.6 Story Mode — Story Arc Mandatory (UPDATED 2026-04-07)"),
    ("Normal", "All four demos produced textbook paragraphs and called them narratives. The new rule: every Story Mode segment must follow a complete story arc — Problem → Struggle → Discovery → Solution."),
    ("Normal", "Problem — opens with a real problem or mystery. Not \"Today we'll learn about...\" but \"In 1774, a priest named Priestley heated a red powder and something impossible happened...\""),
    ("Normal", "Struggle — character tries familiar approaches and they fail. Tension builds."),
    ("Normal", "Discovery — the new concept arrives as the missing tool. Textbook content lands HERE."),
    ("Normal", "Solution — character applies the concept, situation resolves. Student gets the answer but EARNED it by following the logic."),
    ("Normal", "Real-world examples preferred. If the concept has a real historical discovery or documented case, use it as the spine. Invented scenarios are a fallback."),
    ("Normal", "IELTS-style comprehension on checkpoints. Test main idea, inference, comparing across segments, vocabulary in context — NOT fact recall. Fact recall belongs in Memory Sprint."),
    ("Heading 3", "The Stranger Test (Quality Gate)"),
    ("Normal", "Before deployment, every Story Mode segment must pass: if a person who has never seen this textbook content can follow the segment and answer the checkpoint question using only the narrative logic (not prior knowledge), it passes. If it requires the reader to already know the content, it fails. Failing segments return to the content pipeline for rewrite."),
]

GAME_BREAKS_UPDATE = [
    ("Heading 3", "Game: Tic Tac Toe vs AI (NEW 2026-04-07)"),
    ("Normal", "A strategic game where the student plays Tic Tac Toe against the AI. Each move requires answering a subject-related question correctly."),
    ("Normal", "Game flow: AI asks a first question to determine who starts (correct → student first, wrong → AI first). On the student's turn: student picks a tile, AI asks a question, correct answer = X placed where chosen, wrong = X placed on a random tile (0.2% mercy chance it lands where intended). On AI's turn: AI moves strategically with no question."),
    ("Normal", "Difficulty starts calibrated to PISA level. After each wrong answer, question difficulty drops one tier. Floor: 60% of student's current PISA level. AI plays optimally — it will not deliberately let the student win."),
    ("Normal", "3 chances total (3 games). All 3 failed = task FAILED, score = 0, student must redo via Duolingo Mode. Win = standard XP + bonus for unused attempts."),
    ("Normal", "The system MUST track and display win rate for this game mode. Win rate feeds the mastery map and is visible on the student's profile."),
    ("Heading 3", "Mystery Box — Reworked as Case-Opening (UPDATED 2026-04-07)"),
    ("Normal", "The Mystery Box is no longer a static interleaving question. It is now a case-opening event — explicitly NOT gambling. Every attempt yields XP. The mystery is what question appears and what reward pool item is earned, not whether the student learns."),
    ("Normal", "Reward pool (drawn on correct answer): XP Boost 40%, Extra AI Usage 25%, Avatar Frame 15%, Customization Item 10%, Coupon 7%, Mythical Boss Event Ticket 3%. Wrong answer = small participation XP only, no bonus item. Partial answer = reduced XP + smaller bonus."),
    ("Normal", "Appearance: either random surprise pop-up during homework, or pre-placed by the content creator at a milestone."),
]

REAL_LIFE_UPDATE = [
    ("Heading 2", "4.6 Real-Life Challenge — First-Person Expert POV (UPDATED 2026-04-07)"),
    ("Normal", "All demos used third-person scenarios (\"Tomir's plant is dying\", \"Salimjon is in his garage\"). Students were observers, not decision-makers. Every Real-Life Challenge must now put the student in the role of THE EXPERT who must solve a real case."),
    ("Normal", "POV: direct address only. \"You are...\" / \"Your job is...\" / \"A client comes to you...\" Never third-person."),
    ("Normal", "Role: student is positioned as an expert — fire safety inspector, scientist, engineer, consultant, doctor, lawyer, historian — never as a student answering a textbook question."),
    ("Normal", "Case: real-world cases preferred (historical incidents, documented events, known engineering challenges). Realistic hypotheticals only as fallback."),
    ("Normal", "Tricky questions: multiple-path with plausible wrong answers. Forces real reasoning, not pattern matching."),
    ("Normal", "Explanation required: student must justify the choice. AI evaluates the reasoning, not just the choice."),
    ("Normal", "Example (Combustion): OLD — \"Salimjon's garage has petrol. Should he light the stove?\" NEW — \"You are a fire safety inspector called to a workshop. The owner stores gasoline, uses a wood-burning stove, and keeps an asbestos fire blanket in the corner. He asks: 'Is my setup safe?' You spot 3 problems. What are they, and what do you tell the owner?\""),
]

BOSS_UPDATE = [
    ("Heading 2", "6.6 Three-Tier AI Boss System (UPDATED 2026-04-07)"),
    ("Normal", "The Final Boss is no longer a single type. There are now three distinct AI Boss tiers, each serving a different pedagogical purpose. Sub Boss is the default; Big Boss replaces it weekly; Mythical Boss appears as a rare random event. The HP system, damage tiers, hint tax, Socratic tutoring on failure, and mastery star calculations from sections 6.1-6.5 above apply to all three tiers unless explicitly overridden."),
    ("Heading 3", "Sub Boss (every homework session — default)"),
    ("Normal", "Unchanged from the existing Phase 6 specification. Tests current lesson content at the student's PISA-calibrated difficulty. 3 attempts max. Socratic tutoring on failure. Regenerated questions on retry."),
    ("Heading 3", "Big Boss (weekly, replaces that week's final Sub Boss)"),
    ("Normal", "Auto-assigned once per week per student. Cross-subject scope: AI identifies the student's 3 weakest standards across ALL completed subjects and builds the question set around them."),
    ("Normal", "Difficulty: one tier above the student's current PISA level in each subject. If the student is Math 1.7, Big Boss Math questions target PISA L2-3."),
    ("Normal", "Format: same as Sub Boss — open reasoning, no MC for Grade 5+, Socratic hints available (cost HP)."),
    ("Normal", "Reward: 2× XP. Weekly leaderboard badge. Streak freeze token if score > 80%."),
    ("Normal", "Failure: weak standards persist into next week's Big Boss until mastered."),
    ("Heading 3", "Mythical Boss (random elite event, <5% chance)"),
    ("Normal", "Trigger: random appearance, less than 5% chance per homework session. The student cannot predict, force, or farm it. Appears as a surprise: \"⚠️ A Mythical Boss has appeared!\""),
    ("Normal", "Scope: ALL subjects passed up to current point. Questions span Math, Science, Biology, History, anything completed."),
    ("Normal", "Question type — one of: (1) Real-World Case Scenario requiring multi-subject reasoning, or (2) Advanced Mathematical Challenge with high-level equations or cross-disciplinary math."),
    ("Normal", "Difficulty: PISA L5-6 regardless of student level. Designed to overwhelm. Most students will fail — by design."),
    ("Normal", "Hints: ZERO. No Socratic tutoring. ONE attempt only."),
    ("Normal", "Reward if passed: 5× normal session XP, exclusive title (\"Mythical Slayer\", \"The Chosen One\"), permanent exclusive avatar frame, entry into Mythical Hall of Fame (opt-in, class-visible)."),
    ("Normal", "If failed: NO penalty. Small participation XP. \"The Mythical Boss retreats... for now.\""),
    ("Normal", "Psychological role: aspiration, social buzz, legend status. The rarity makes it legendary."),
]

DUOLINGO_MODE = [
    ("Heading 1", "SECTION 8-A: DUOLINGO MODE (NEW 2026-04-07)"),
    ("Normal", "Triggered when a student scores below 60% on a homework session. The student MUST redo the failed work until they reach 60%. No skip. No XP for the failed attempts."),
    ("Heading 2", "Step-by-Step Flow"),
    ("Normal", "1. Immediate feedback: \"You scored X%. You need 60% to pass. Let's work on what you missed.\""),
    ("Normal", "2. Task-by-task review: the system identifies which specific questions the student got wrong and re-presents them with additional scaffolding (hints, simplified versions)."),
    ("Normal", "3. Adjusted difficulty: re-attempt questions are one difficulty tier lower."),
    ("Normal", "4. Repeat until 60%: no upper limit on attempts, but the teacher is auto-flagged after 3 failed sessions on the same topic."),
    ("Normal", "5. No XP penalty: XP is only awarded once the student crosses 60%. Previous failed attempts earn 0 XP (not negative)."),
    ("Heading 2", "Why It Exists"),
    ("Normal", "The original framework had no clear pass/fail threshold and no remediation loop. A student who failed simply moved on. Duolingo Mode closes the loop by refusing to rubber-stamp failure while keeping the framing growth-mindset positive."),
]

CELEBRATION_TIERS = [
    ("Heading 1", "SECTION 8-B: ASSESSMENT CELEBRATION TIERS (NEW 2026-04-07)"),
    ("Normal", "Replaces the old binary pass/fail with a tiered emotional response system that matches what the student actually achieved. Every score range has a distinct UI response, reward calculation, and (where applicable) follow-up offer."),
    ("Heading 2", "Score-Tier Matrix"),
    ("Normal", "100% — PERFECT. On-screen celebration (confetti, animation, sound). Maximum XP. Temporary title: \"Perfectionist\"."),
    ("Normal", "99% — Almost Perfect. Tease: \"So close! One more point!\" Maximum XP minus 1%."),
    ("Normal", "90-98% — Excellent. Strong praise. Full XP. Streak fire animation."),
    ("Normal", "80-89% — Very Good. Bonus XP for hitting 80% on first attempt. Standard praise."),
    ("Normal", "60-79% — Passing. Standard XP. Repass option offered: \"You passed! Want to try again for a better score?\""),
    ("Normal", "Below 60% — Failed. Enters Duolingo Mode (Section 8-A). No XP for failed tasks."),
    ("Heading 2", "Repass Option (60-79%)"),
    ("Normal", "If accepted: homework restarts with regenerated questions (same standards, different numbers/scenarios)."),
    ("Normal", "If declined: session completes with current score and XP."),
    ("Normal", "Maximum 2 repass attempts per homework."),
    ("Heading 2", "Did You Know — Loading Screen Insights (NEW)"),
    ("Normal", "Between phase transitions, the system displays 1-2 sentence \"Did You Know...\" insights related to the current lesson theme. 1 per phase transition. 5-second readable. Conversational tone. No quiz follow-up. Random off-topic facts are rejected. Example: \"Did you know? The word 'oxygen' comes from Greek words meaning 'acid producer' — because scientists originally thought ALL acids contained oxygen. They were wrong.\""),
    ("Heading 2", "Answer Checking — Dual Pathway (NEW)"),
    ("Normal", "Hardcoded questions (MC, matching, numeric, formula input) are checked by rule-based exact match — no AI, instant feedback, <0.1s."),
    ("Normal", "Open-reasoning questions (typed responses, justifications, essays, case analyses) are checked by AI Tier 3 in real time. AI receives: question, rubric, student answer, PISA level. AI returns: score 0-100%, partial credit breakdown, one-sentence feedback, all within 3 seconds. Hard 5-second timeout falls back to keyword matching against rubric concepts. AI never reveals the correct answer during evaluation."),
]

# ─────────── Locate anchors and inject content ───────────

# 1. Insert PRE-HOMEWORK intro + Section 0-A + 0-B before Section 1
section1 = find_paragraph_by_text("SECTION 1: PHASE 1", style_prefix="Heading 1")
if section1 is not None:
    insert_block_before(section1, PRE_HOMEWORK_INTRO + SECTION_0A + SECTION_0B)
    print("Inserted Pre-Homework + 0-A + 0-B before SECTION 1")
else:
    print("WARN: SECTION 1 anchor not found")

# 2. Memory Sprint flexible format — insert before Section 2
section2 = find_paragraph_by_text("SECTION 2: PHASE 2", style_prefix="Heading 1")
if section2 is not None:
    insert_block_before(section2, MEMORY_SPRINT_UPDATE)
    print("Inserted Memory Sprint flexible-format block before SECTION 2")

# 3. Story Mode arc + Stranger Test — insert before Section 3
section3 = find_paragraph_by_text("SECTION 3: PHASE 3", style_prefix="Heading 1")
if section3 is not None:
    insert_block_before(section3, STORY_MODE_UPDATE)
    print("Inserted Story Mode arc block before SECTION 3")

# 4. Game Breaks: Tic Tac Toe + Mystery Box rework — insert before Section 4
section4 = find_paragraph_by_text("SECTION 4: PHASE 4", style_prefix="Heading 1")
if section4 is not None:
    insert_block_before(section4, GAME_BREAKS_UPDATE)
    print("Inserted Tic Tac Toe + Mystery Box rework before SECTION 4")

# 5. Real-Life Expert POV — insert before Section 5
section5 = find_paragraph_by_text("SECTION 5: PHASE 5", style_prefix="Heading 1")
if section5 is not None:
    insert_block_before(section5, REAL_LIFE_UPDATE)
    print("Inserted Real-Life expert POV block before SECTION 5")

# 6. Boss 3-tier system — insert before Section 7
section7 = find_paragraph_by_text("SECTION 7: POST-BOSS", style_prefix="Heading 1")
if section7 is not None:
    insert_block_before(section7, BOSS_UPDATE)
    print("Inserted 3-tier Boss block before SECTION 7")

# 7. Section 8-A Duolingo Mode + 8-B Celebration Tiers — insert before Section 9
section9 = find_paragraph_by_text("SECTION 9: CUSTOMIZATION", style_prefix="Heading 1")
if section9 is not None:
    insert_block_before(section9, DUOLINGO_MODE + CELEBRATION_TIERS)
    print("Inserted Section 8-A + 8-B before SECTION 9")

# 8. Update Appendix C heading title (duration label)
for p in doc.paragraphs:
    if "Appendix C" in p.text and "Session Timeline" in p.text:
        for run in p.runs:
            if "20" in run.text and "30" in run.text:
                run.text = run.text.replace("20–30 min", "30-45 min").replace("20-30 min", "30-45 min")
        # if no run-level match (text could be split), rewrite the whole paragraph text
        if "30-45" not in p.text:
            full = p.text.replace("20–30 min", "30-45 min").replace("20-30 min", "30-45 min")
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = full
            else:
                p.add_run(full)
        print(f"Updated Appendix C title: {p.text}")
        break

doc.save(PATH)
print(f"\nSaved updated Blueprint to: {PATH}")
