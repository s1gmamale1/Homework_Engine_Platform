"""Generate summary docx files for NETS Homework Engine Standards and Blueprint."""
import os
from docx import Document

OUT = os.path.join(os.path.dirname(__file__), "..", "standards")

# ─── FILE 1: STANDARDS SUMMARY ───

doc = Document()
doc.add_heading("NETS Homework Engine — Unified Specification Summary", level=0)

meta = [
    ("Source", "standards/NETS-Homework-Engine-UNIFIED.md (v2.0, 2,072 lines, 20 sections + 2 appendices)"),
    ("Status", "Single Source of Truth — supersedes all previous specs"),
    ("Target", "8.8M K-11 students, Uzbekistan"),
    ("PISA 2022 Baseline", "Math 364, Reading 336, Science 355, Creative Thinking 14/60"),
    ("Goal", "80%+ below-L2 students reach L2-3 within 2 years"),
]
for label, val in meta:
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(val)

# Core Principles
doc.add_heading("Core Principles (\u00a71)", level=1)
for title, desc in [
    ("Textbook is Source of Truth", "NETS wraps content in engagement, never alters facts"),
    ("Flow State is Mandatory", "70-80% success rate target (Csikszentmihalyi), adjust every 3-5 questions"),
    ("Intrinsic Motivation > Extrinsic", "Zero pay-to-win, all rewards celebrate learning"),
    ("No Busywork Rule", "Every task must scaffold ONE specific PISA transition skill"),
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{title} \u2014 ").bold = True
    p.add_run(desc)

# Content Architecture
doc.add_heading("Content Architecture (\u00a72)", level=1)
for b in [
    "Hierarchy: Textbook > Chapter > Topic > Learning Objective > Content Block > Assessment Item",
    "Learning Units: 1-3 per chapter section, each maps to one curriculum standard",
    "Each LU generates 10-15 tasks across Bloom\u2019s (4R, 3U, 3Ap, 2An, 1-2 E/C)",
    "Dual Standard Codes: Primary descriptive UZ-MATH-5-FRAC-01 + alias dotted MAT.5.3.4.1",
    "Mandatory tagging: textbook_ref, standard_ref, pisa_ref (with transition_skill), blooms_level, game_mechanic, prerequisites",
]:
    doc.add_paragraph(b, style="List Bullet")

# PISA Framework
doc.add_heading("PISA Framework (\u00a73)", level=1)
for b in [
    "7 levels (Below 1 through 6) defined for Math, Reading, Science",
    "Continuous scale tracking per student (e.g., 1.7) via IRT-weighted performance",
    "Domain breakdowns \u2014 Math: quantity/space/change/uncertainty; Reading: access/integrate/reflect; Science: identify/explain/interpret/evidence",
    "History uses dual-domain mapping (Reading + Creative Thinking)",
]:
    doc.add_paragraph(b, style="List Bullet")

# Session Engine
doc.add_heading("Session Engine (\u00a74)", level=1)
for b in [
    "Standard: 20-30 min (homework) | Extended: 40-50 min (in-class) | Recovery: 15-20 min",
    "Initialization: load profile \u2192 load assignment \u2192 calculate difficulty band \u2192 select content \u2192 verify transition skill coverage \u2192 build session plan",
]:
    doc.add_paragraph(b, style="List Bullet")

# 7 Phases
doc.add_heading("7 Phases (\u00a75)", level=1)
table = doc.add_table(rows=8, cols=5)
table.style = "Light Grid Accent 1"
for i, h in enumerate(["Phase", "Name", "Duration", "Bloom\u2019s", "Key Rules"]):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True

for r, row in enumerate([
    ["P1", "Memory Sprint", "2 min", "Remember", "5-8 recall items from PREVIOUS chapters, spaced repetition"],
    ["P2", "Story Mode", "5-7 min", "Understand", "3 segments + checkpoints, textbook narrative, CPA for math"],
    ["P3", "Game Breaks", "6-9 min", "Apply", "3 games interleaved with Story, 1 reinforce + 1 stretch + 1 transition skill"],
    ["P4", "Real-Life Challenge", "3-5 min", "Analyze/Evaluate", "Uzbek-context transfer, PISA L2+ required"],
    ["P5", "Consolidation", "2-3 min", "Remember", "Mnemonic Lock before boss, Memory Palace default"],
    ["P6", "Final Boss", "5-10 min", "Apply\u2192Create", "HP system (50/100/150 by grade), NO SKIPPING"],
    ["P7", "Reflection", "2-3 min", "Meta", "Metacognitive prompts, privacy-protected"],
]):
    for c, val in enumerate(row):
        table.rows[r + 1].cells[c].text = val

doc.add_paragraph()
for label, val in [
    ("Boss HP", "Grades 1-4: 50 | Grades 5-8: 100 | Grades 9-11: 150"),
    ("Boss Stars", "1\u2605 = defeated | 2\u2605 = \u22642 attempts, >50% HP | 3\u2605 = 1st attempt, no hints, >80% HP"),
    ("Boss Failure", "Socratic tutoring \u2192 regenerated questions \u2192 max 3 attempts"),
]:
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(val)

# 14 Game Mechanics
doc.add_heading("16 Game Mechanics (\u00a76, updated 2026-04-07)", level=1)
doc.add_paragraph("Memory Sprint, Spaced Flashcards, Tile Match, Sentence Fill, Memory Palace, Story Mode, Adaptive Quiz (IRT), Mystery Box (case-opening), Movement Breaks (Grades 1-4 required), Why Chain, Peer Teaching (Grades 5+), Real-Life Challenge, Reflection Journal, Final Boss (3-tier: Sub/Big/Mythical), Tic Tac Toe vs AI (NEW), Notebook Capture (NEW \u2014 hand-written/drawn tasks evaluated by AI vision)")
doc.add_paragraph("Plus Creative Lab (optional, Grades 5+, PISA L3+)", style="List Bullet")
doc.add_paragraph(
    "Complementary pool: the Interactive Knowledge-Gated Game Catalog "
    "(standards/NETS-Interactive-Game-Catalog.md, 12 games) is a sibling pool that "
    "Phase 3 draws from per the dual-catalog rule \u2014 \u22651 game from the Interactive "
    "Catalog and \u22652 from the Default 16 every session. See \u00a76.9.",
    style="List Bullet",
)

# Difficulty Adaptation
doc.add_heading("Difficulty Adaptation (\u00a79)", level=1)
for b in [
    "Real-time: <60% \u2192 tier down, >90% \u2192 tier up, 70-80% \u2192 flow zone",
    "Secondary: 3 consecutive correct \u2192 +0.5 PISA sub-level, 2 consecutive wrong \u2192 -0.5",
    "Emergency: <40% in game \u2192 insert Story Mode micro-segment for that standard",
    "Cross-session: update PISA profile, spaced repetition, mastery map, trajectory",
]:
    doc.add_paragraph(b, style="List Bullet")

# Content Pipeline
doc.add_heading("Content Pipeline (\u00a710)", level=1)
for b in [
    "4 phases: Ingestion (OCR) \u2192 Task Generation (3 AI tiers) \u2192 QA \u2192 Session Assembly",
    "Minimum per topic: ~85 content items, all tagged/reviewed/traceable",
    "Cost: ~$61 per textbook, ~$4,880 for Grades 5-8 pilot",
]:
    doc.add_paragraph(b, style="List Bullet")

# Gamification Economy
doc.add_heading("Gamification Economy (\u00a711)", level=1)
for b in [
    "XP: 50-1000 per activity, streak multipliers (up to 2x at 30+ days)",
    "Badges: Mentor, Streak Master, PISA Champion, Boss Slayer, etc.",
    "Leaderboard: class-level only, weekly reset, anti-competition design",
]:
    doc.add_paragraph(b, style="List Bullet")

# Edge Cases
doc.add_heading("Edge Cases (\u00a712)", level=1)
for b in [
    "Missed homework: 3-day grace, condensed recovery session",
    "Extended absence (7+ days): Catch-Up Mode with reduced boss HP",
    "Boost Mode: triggered by 3+ declining sessions",
    "Low engagement: reduce session length, increase game ratio",
]:
    doc.add_paragraph(b, style="List Bullet")

# Anti-Cheat
doc.add_heading("Anti-Cheat (\u00a713)", level=1)
doc.add_paragraph("Question regeneration, response time analysis, pattern analysis, Socratic verification, device monitoring, copy-paste detection")
p = doc.add_paragraph()
p.add_run("Escalation: ").bold = True
p.add_run("soft warning \u2192 parent \u2192 teacher \u2192 admin")

# Teacher Controls
doc.add_heading("Teacher Controls (\u00a714)", level=1)
doc.add_paragraph("Session mode, grace period, boss HP modifier, hint toggle, difficulty floor, game selection, recovery limits")

# Integrations
doc.add_heading("Integrations (\u00a715)", level=1)
for b in [
    "Kundalik/eMaktab: rosters, attendance, grades",
    "Parent portal: PISA progression, recommendations, weekly summary",
    "Admin analytics: per-school/region/national PISA levels",
]:
    doc.add_paragraph(b, style="List Bullet")

# JSON Schema
doc.add_heading("JSON Schema (\u00a716)", level=1)
for b in [
    "Task ID format: HW-{SUBJECT}{GRADE}-{TOPIC}-{SEQ}",
    "Required fields: task_id, textbook_ref, curriculum_standard, pisa_level (with transition_skill), blooms_level, game_mechanic, content",
    "AI tiers: 1=template, 2=mid-range (Haiku), 3=premium (Sonnet/Opus)",
]:
    doc.add_paragraph(b, style="List Bullet")

# Key Design Decisions
doc.add_heading("Key Design Decisions (Appendix B)", level=1)
for b in [
    "Boss HP Grades 5-8: 80\u2192100 (fixed math)",
    "Final Boss: Level 3-6 tiered (resolved contradiction)",
    "Session: 20-30 min default (50 min = Extended)",
    "Dual standard codes (descriptive primary)",
    "transition_skill field mandatory on every task",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.save(os.path.join(OUT, "NETS-Homework-Engine-Standards-Summary.docx"))
print("Standards summary saved.")


# ─── FILE 2: BLUEPRINT SUMMARY ───

doc = Document()
doc.add_heading("NETS Homework Engine \u2014 Blueprint Summary", level=0)

meta = [
    ("Source", "standards/NETS-Homework-Engine-Universal-Blueprint.docx"),
    ("Revision", "Incorporates IMPROVEMENTS_TO_THE_CURRENT_FRAMEWORKS.md (2026-04-07)"),
    ("Relationship", "UNIFIED = WHAT (content, standards, games). Blueprint = HOW (session lifecycle, engagement hooks, implementation). Neither sufficient alone."),
    ("Markers", "[SPEC \u00a7X.X] = cross-ref to UNIFIED spec. [HOOK] = engagement mechanic. [AI: Tier X] = which AI handles it. [NEW] / [UPDATED] = changes from this revision."),
    ("Duration (UPDATED)", "Standard: 30-45 min (was 20-30) | Extended: 45-60 min (was 40-50) | Recovery: unchanged"),
]
for label, val in meta:
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(val)

# Engagement Philosophy
doc.add_heading("Engagement Philosophy", level=1)
doc.add_paragraph("The Blueprint builds on four behavioral frameworks to turn the UNIFIED spec into an addictive (but ethical) learning loop:")
for b in [
    "White Hat gamification (epic meaning, accomplishment, empowerment) as primary motivator \u2014 chosen over Black Hat (scarcity, loss avoidance) which is used sparingly and only for re-engagement.",
    "Hook Model (Nir Eyal): every phase has Trigger \u2192 Action \u2192 Variable Reward \u2192 Investment. External triggers (push notifications) convert to internal triggers (streak fear, progress desire) over time.",
    "Octalysis (Yu-kai Chou): all 8 core drives are activated during Boss Fight; other phases use subsets. This was a deliberate choice to make the Boss the emotional peak.",
    "Duolingo mechanics adapted: streak freezes (500 gems, max 2/week), hearts system (boss hints cost HP), endowed progress (bar starts at ~15%).",
]:
    doc.add_paragraph(b, style="List Bullet")

# Section 0
doc.add_heading("Section 0: Pre-Session (<2 seconds, server-side, invisible)", level=1)
p = doc.add_paragraph()
p.add_run("Purpose: ").bold = True
p.add_run("Everything that happens before the student sees anything. Assembles the entire session plan server-side so the experience feels instant.")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("What it organizes:").bold = True
for b in [
    "Load student profile (PISA levels, mastery map, transition skills, streak, trajectory)",
    "Load assignment (textbook, chapter, objectives, deadline, mode)",
    "Hard block: content pool must have \u226585 tagged items or session CANNOT proceed \u2014 this prevents half-baked sessions",
    "Calculate difficulty band: floor/target/ceiling based on PISA level + teacher overrides",
    "Assemble session plan: select specific items for each of the 7 phases, ensuring Bloom\u2019s coverage and transition skill representation",
    "Generate lesson brief (4-5 bullet summary from textbook \u2014 disappears after 30 sec, never shown again)",
    "[HOOK: Endowed Progress] Progress bar starts at ~15% filled \u2014 student feels they\u2019ve already started before doing anything",
]:
    doc.add_paragraph(b, style="List Bullet")

# Section 0-A
doc.add_heading("Section 0-A: Theme Preview (NEW \u2014 Pre-Homework, student-paced)", level=1)
p = doc.add_paragraph()
p.add_run("Purpose: ").bold = True
p.add_run("Optional pre-session warmup. Lets the student explore the topic before any pressure or scoring. No timers, no XP, no penalties \u2014 pure curiosity priming.")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("8 components (student can browse any/all):").bold = True
for b in [
    "Summary \u2014 short topic overview from textbook",
    "Better Explanation \u2014 reworded for clarity, alternative angle",
    "Examples \u2014 concrete worked examples",
    "Real-Life Research \u2014 how the concept appears in the world",
    "Personal Hook \u2014 why a student their age might care",
    "Why It Matters \u2014 long-term relevance / academic ladder",
    "Industry Application \u2014 careers and fields that use it",
    "Additional Materials \u2014 videos, articles, supplementary links in any language",
]:
    doc.add_paragraph(b, style="List Bullet")
p = doc.add_paragraph()
p.add_run("Choice rationale: ").bold = True
p.add_run("Demos showed students felt dropped into content cold. Theme Preview gives the curious ones a self-directed on-ramp without forcing the rushers to slow down.")

# Section 0-B
doc.add_heading("Section 0-B: Flash Cards (NEW \u2014 Pre-Homework reference deck)", level=1)
p = doc.add_paragraph()
p.add_run("Purpose: ").bold = True
p.add_run("All key formulas, concepts, and rules for the topic, one per swipeable card. Acts as a personal reference deck the student can return to during the homework session.")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("How it\u2019s organized:").bold = True
for b in [
    "One concept/formula/rule per card \u2014 no walls of text",
    "Swipeable interface, student-paced, no scoring",
    "Returnable during homework phases (quick-access overlay) \u2014 NOT a hint, just a reference",
    "Ends with explicit \"Start my Homework\" button \u2014 student decides when they\u2019re ready",
]:
    doc.add_paragraph(b, style="List Bullet")
p = doc.add_paragraph()
p.add_run("Choice rationale: ").bold = True
p.add_run("Removes the \"I forgot the formula\" friction without giving away answers. The student earns the right to consult their own deck.")

# Section 1
doc.add_heading("Section 1: Phase 1 \u2014 ENGAGE + REMEMBER (Memory Sprint, flexible format)", level=1)
p = doc.add_paragraph()
p.add_run("Purpose: ").bold = True
p.add_run("Pull the student in within the first 30 seconds AND warm up recall from prior chapters. Memory Sprint is no longer just \"questions\" \u2014 it\u2019s a flexible fast-paced format pool.")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("How it\u2019s organized:").bold = True
for b in [
    "Hook Game (30 sec): fast-tap classification, intentionally EASY and ADDICTIVE. 90%+ success rate expected. 50 XP per tap, speed bonus, variable \"SUPER!\" animation on ~20% of answers.",
    "Streak Display + Lesson Brief flash (30 sec, disappears permanently). Streak tiers: 1-day \u2192 7-day \u2192 14-day \u2192 30+ day with increasing rewards.",
    "Streak Freeze: costs 500 earned gems (not real money), max 2/week.",
    "Memory Sprint (UPDATED \u2014 flexible format): \u22642 minutes, fast-paced, quick-dopamine, prior chapters only. NOT just questions.",
]:
    doc.add_paragraph(b, style="List Bullet")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Approved Memory Sprint formats (any of, mixed within the 2 min):").bold = True
for b in [
    "MC / Binary \u2014 quick yes-no or 4-option taps",
    "Speed Match \u2014 pair concepts to definitions",
    "Flash Sprint \u2014 rapid flashcard recall",
    "Fill-in-Blanks Race \u2014 type or tap missing word",
    "Order the Steps \u2014 sequence a procedure correctly",
]:
    doc.add_paragraph(b, style="List Bullet")
p = doc.add_paragraph()
p.add_run("Choice rationale: ").bold = True
p.add_run("Demos showed pure question lists felt like a quiz, killing the dopamine. Mixing formats keeps the 2 min varied while still hitting prior-chapter recall.")

# Section 2
doc.add_heading("Section 2: Phase 2 \u2014 UNDERSTAND (Story Mode, REWRITTEN)", level=1)
p = doc.add_paragraph()
p.add_run("Purpose: ").bold = True
p.add_run("Deliver new textbook content through a complete story arc, not disconnected segments. Demos showed previous Story Mode was bland and felt like a textbook read-aloud.")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Story arc structure (mandatory):").bold = True
for b in [
    "Problem \u2014 a real situation that demands the concept",
    "Struggle \u2014 character tries familiar approaches and they fail",
    "Discovery \u2014 the concept is introduced as the missing tool",
    "Solution \u2014 character applies it, the situation resolves",
]:
    doc.add_paragraph(b, style="List Bullet")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("How it\u2019s organized:").bold = True
for b in [
    "3 segments \u00d7 90 seconds each, with checkpoints between each segment.",
    "Real-world examples preferred over invented scenarios \u2014 actual cases, real people, real places when possible.",
    "IELTS-style reading comprehension on checkpoints (inference, main idea, vocabulary in context) \u2014 not just recall.",
    "Stranger Test quality gate: any reader unfamiliar with the topic should still understand the story on first read. If they can\u2019t, the segment fails QA.",
    "CRITICAL interleaving: Seg1 \u2192 Game1 \u2192 Seg2 \u2192 Game2 \u2192 Seg3 \u2192 Game3. Interleaving improves retention (Roediger & Karpicke).",
    "Textbook page reference (Darslik: sahifa X) pulses on checkpoint failure -- guides student back to source without giving away answer.",
    "Skimming detection: <30 sec on a 90-sec segment \u2192 flagged.",
    "CPA progression for math (Concrete \u2192 Pictorial \u2192 Abstract) built into story segments.",
]:
    doc.add_paragraph(b, style="List Bullet")

# Section 3
doc.add_heading("Section 3: Phase 3 \u2014 APPLY (Game Breaks)", level=1)
p = doc.add_paragraph()
p.add_run("Purpose: ").bold = True
p.add_run("Practice through game mechanics. Confirms the interleaving with Story Mode.")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("How it\u2019s organized:").bold = True
for b in [
    "Game selection: 1 reinforcement game (practice what was just learned) + 1 stretch game (push slightly beyond) + 1 transition skill game (target PISA level-up skill).",
    "Each game has [HOOK] markers for specific engagement mechanics (variable rewards, near-miss feedback, combo multipliers).",
    "In-game adaptation: <60% \u2192 tier down, >90% \u2192 tier up, 70-80% = flow zone. This runs in real-time, not between games.",
    "Choice rationale: 3 games with different purposes prevents repetition fatigue while ensuring both consolidation and growth.",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("New Game: Notebook Capture (hand-written/drawn tasks)", level=2)
p = doc.add_paragraph()
p.add_run("Purpose: ").bold = True
p.add_run("Force the student off the screen and back into a physical notebook for tasks that benefit from hand-writing or hand-drawing \u2014 long-form math derivations, free-body diagrams, geometric constructions, labelled biology diagrams, hand-drawn timelines, essay drafts, calligraphy. Student does the work on paper, takes a photo, uploads it. AI vision evaluates against rubric.")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Why it exists: ").bold = True
p.add_run("Cognitive science (Generation Effect, Embodied Cognition, Mueller & Oppenheimer pen-vs-laptop research) shows physical writing improves retention 25-40% over typing. NETS was previously screen-only \u2014 this closes that gap without breaking the digital game loop.")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("How it works:").bold = True
for b in [
    "Frequency: ~1 in every 3-4 sessions for Grades 5+; ~1 in 6 for Grades 1-4. NOT every session \u2014 overuse breaks the loop.",
    "Where it fits: Phase 3 (replaces one game slot) OR Phase 4 (case requires hand-drawn diagram) OR Phase 6 (Mythical Boss almost always demands one).",
    "Flow: prompt \u2192 work offline (timer frozen, app can be backgrounded) \u2192 tap UPLOAD \u2192 camera with live page-edge detection + focus + lighting validation \u2192 capture \u2192 confirm preview (auto-crop, face/PII auto-blur) \u2192 multi-page support (up to 3) \u2192 send.",
    "AI Tier 3 multimodal vision: 10-second hard cap, OCR + diagram understanding, returns score + partial credit + specific feedback. Annotations overlaid on photo (green checks on correct parts, red circles on problems).",
    "Privacy: encrypted, 30-day retention, teacher/parent only. Faces and personal identifiers blurred client-side BEFORE upload.",
    "Anti-cheat: vision flags suspiciously perfect work (looks printed), identical submissions across students, handwriting baseline mismatch \u2014 routed to teacher review, NOT auto-failed.",
    "Accessibility: visually impaired \u2192 audio recording task instead. Motor disability \u2192 teacher toggle off. No notebook \u2192 in-app sketch fallback capped at 80% XP (bypasses embodied cognition benefit).",
    "Teacher controls: enable/disable, frequency 1-in-3 default (1-in-6 min, every-session max), pristine-notebook bonus toggle, dashboard review, manual score override.",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Interactive Game Catalog Integration", level=2)
p = doc.add_paragraph()
p.add_run("Catalog file: ").bold = True
p.add_run("standards/NETS-Interactive-Game-Catalog.md (v1.0, 12 knowledge-gated games derived from classic board/card/puzzle mechanics).")
p = doc.add_paragraph()
p.add_run("Dual-catalog selection rule: ").bold = True
p.add_run("Phase 3 picks 3 games per session. At LEAST 1 must come from the Interactive Catalog AND at LEAST 2 must come from the Default 16-mechanic pool (\u00a76). The Interactive Catalog is a complement, never a replacement.")
p = doc.add_paragraph()
p.add_run("Games in the catalog:").bold = True
for b in [
    "Tic Tac Toe vs AI \u2014 minimax duel; correct answer places chosen tile, wrong = random.",
    "Connect Four vs AI \u2014 7x6 drop grid; column choice gated by question.",
    "Codebreaker \u2014 Mastermind-style code crack; each guess attempt gated.",
    "Memory Match Blitz \u2014 Concentration grid; matched pairs confirmed by question.",
    "Reaction Chain \u2014 6-10 node chain; 3 wrong in a row breaks the chain.",
    "Word Ladder Climb \u2014 transform start word to target one letter at a time.",
    "Puzzle Lock (Sliding Tile) \u2014 3x3/4x4 sliding puzzle; wrong = random tile slides.",
    "Blackjack 21 Knowledge Edition \u2014 Hit/Stand gated; wrong = AI makes the choice.",
    "Territory Conquest \u2014 simplified Risk 1v1 area control; failed attack loses territory.",
    "Escape Room: Subject Lock \u2014 4 locked objects, cross-hinting clues, time limit.",
    "Bridge Builder \u2014 physics bridge construction; wrong answers waste budget.",
    "Minefield Navigator \u2014 hidden-mine grid; wrong = forced adjacent move (likely mine).",
]:
    doc.add_paragraph(b, style="List Bullet")
p = doc.add_paragraph()
p.add_run("When to pick a Catalog game vs a Default game: ").bold = True
p.add_run("Catalog games shine when the topic maps cleanly to a spatial/strategic metaphor (chronology \u2192 Puzzle Lock, geography \u2192 Territory Conquest, periodic table \u2192 Codebreaker), when engagement has dipped, or for Big Boss / Mythical Boss cross-subject reviews. Default mechanics win for high-density Remember/Understand drilling, narrative-heavy content, physical (notebook) work, Recovery Sessions, and below-PISA-L2 students who need lower cognitive overhead.")

doc.add_heading("New Game: Tic Tac Toe vs AI", level=2)
p = doc.add_paragraph()
p.add_run("Purpose: ").bold = True
p.add_run("Turns recall practice into a strategic micro-duel. The student isn\u2019t answering a quiz, they\u2019re trying to beat an opponent \u2014 the question is just the gate.")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("How it works:").bold = True
for b in [
    "AI asks a question \u2192 correct answer = student places their tile where they choose; wrong answer = a random tile is placed for the student.",
    "3 chances total per game (3 wrong = forfeit).",
    "Adaptive difficulty: question difficulty decreases if the student is losing.",
    "Minimum success rate enforced: 60% (otherwise dropped to easier question pool).",
    "AI plays optimally (perfect tic-tac-toe), so the student MUST answer correctly to have any chance.",
]:
    doc.add_paragraph(b, style="List Bullet")
p = doc.add_paragraph()
p.add_run("Choice rationale: ").bold = True
p.add_run("Combines an externally familiar game with content recall. Beating an \"unbeatable\" AI by knowing the material is a strong intrinsic reward.")

# Section 4
doc.add_heading("Section 4: Phase 4 \u2014 ANALYZE (Real-Life Challenge, REWRITTEN)", level=1)
p = doc.add_paragraph()
p.add_run("Purpose: ").bold = True
p.add_run("Transfer knowledge to real-world context with the student in the driver\u2019s seat. Demos showed previous version was too detached \u2014 generic third-person scenarios. Now: first-person expert POV.")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("How it\u2019s organized:").bold = True
for b in [
    "First-person Expert POV: \"You are the doctor / engineer / historian / lawyer making this call.\" Student IS the decision-maker, not an observer.",
    "Real cases preferred over invented scenarios \u2014 actual historical events, real legal cases, documented engineering decisions.",
    "Tricky multiple-path questions with plausible wrong answers (distractors that look right to a novice) \u2014 forces real reasoning, not pattern matching.",
    "Explanation required after answering: student must justify their decision. AI evaluates the reasoning, not just the choice.",
    "Eligibility gate: PISA L2+ required. Below L2 gets adaptive quiz instead.",
    "Uzbek cultural context preferred where applicable: Uzbek names, locations, som currency.",
    "[HOOK: Reward Chest] on completion with variable XP + random avatar item.",
]:
    doc.add_paragraph(b, style="List Bullet")

# Section 5
doc.add_heading("Section 5: Phase 5 \u2014 CONSOLIDATION", level=1)
p = doc.add_paragraph()
p.add_run("Purpose: ").bold = True
p.add_run("Lock knowledge before Boss Fight. Creates psychological \"I\u2019m ready\" moment.")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("How it\u2019s organized:").bold = True
for b in [
    "Technique selected by grade band AND content type (not one-size-fits-all).",
    "Memory Palace is default for Grades 3+. Younger students get simpler mnemonic games.",
    "Choice: the consolidation phase exists specifically to reduce boss anxiety. Students who feel prepared perform 20-30% better (testing effect research).",
]:
    doc.add_paragraph(b, style="List Bullet")

# Section 6
doc.add_heading("Section 6: Phase 6 \u2014 BOSS FIGHT (3-Tier AI Boss System, UPDATED)", level=1)
p = doc.add_paragraph()
p.add_run("Purpose: ").bold = True
p.add_run("The emotional peak of the session. Previously a single boss per session \u2014 now a 3-tier system that scales from per-homework to mythical rare events.")

doc.add_heading("Tier 1: Sub Boss (every homework)", level=2)
for b in [
    "The standard per-homework boss \u2014 current topic only.",
    "HP: 50 (Grades 1-4), 100 (Grades 5-8), 150 (Grades 9-11).",
    "Damage scaled by PISA tier (Easy=L3, Med=L4, Hard=L5-6).",
    "Combo system: 3 correct in a row \u2192 2x damage.",
    "Hint tax: hints heal the boss.",
    "AI Tier 3 evaluates open-ended answers; accepts multiple valid interpretations for History/Reading.",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Tier 2: Big Boss (weekly)", level=2)
for b in [
    "Appears once per week, cross-subject.",
    "Targets the student\u2019s 3 weakest standards (pulled from mastery map).",
    "Pays 2x XP across the fight.",
    "Choice rationale: forces spaced retrieval across subjects, not just within today\u2019s topic.",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Tier 3: Mythical Boss (random, <5% chance)", level=2)
for b in [
    "Random spawn at <5% probability per session \u2014 cannot be farmed.",
    "PISA 5-6 difficulty only, ZERO hints allowed.",
    "Massive rewards: 5x XP, exclusive title, avatar frame, Hall of Fame entry.",
    "Choice rationale: variable-ratio reward at the rarest tier creates a long-tail dopamine hook (the slot machine effect, but earned through skill).",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Shared boss rules (all tiers):", level=2)
for b in [
    "Victory: Star calculation (1\u2605/2\u2605/3\u2605) \u2192 animated reward chest \u2192 social comparison (opt-in).",
    "Defeat: max 3 attempts \u2192 Socratic tutoring (NEVER reveals answer) \u2192 regenerated questions on retry \u2192 teacher notification if all 3 fail.",
    "[HOOK: Near-Miss] \"Boss HP: 15 qoldi! Ozgina qoldi edi!\" creates urgency.",
    "Open-reasoning only for Grade 5+ (no MCQ at boss level).",
]:
    doc.add_paragraph(b, style="List Bullet")

# Section 7
doc.add_heading("Section 7: POST-BOSS \u2014 AI Analysis Engine (<10 sec)", level=1)
p = doc.add_paragraph()
p.add_run("Purpose: ").bold = True
p.add_run("Invisible server-side analysis that drives all downstream reporting and adaptation. Student never sees this phase.")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("What it computes:").bold = True
for b in [
    "Performance analysis: by phase, by Bloom\u2019s level, by standard, by transition skill.",
    "Root cause identification: WHY the student failed, not just WHAT they got wrong. Choice: most LMS only track correctness; NETS tracks the reasoning gap.",
    "PISA level recalculation with IRT weighting + exponential decay (recent performance weighted more).",
    "Trajectory check: 3+ declining sessions \u2192 triggers Boost Mode automatically.",
    "Generates 3 outputs: next session plan, teacher report (detailed), parent report (simple, actionable, in parent\u2019s language, no PISA jargon).",
]:
    doc.add_paragraph(b, style="List Bullet")

# Section 8
doc.add_heading("Section 8: REMEDIATION + REFLECTION", level=1)
p = doc.add_paragraph()
p.add_run("Purpose: ").bold = True
p.add_run("Turn failures into bonus XP opportunities (reframing) and create hooks for next session.")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("How it\u2019s organized:").bold = True
for b in [
    "Weakness display: max 2-3 areas shown, framed as \"bonus XP opportunity\" not \"you failed at\". Choice: growth mindset framing (Dweck) reduces learned helplessness.",
    "Micro-exercises: simpler scaffolded versions targeting specific failure points.",
    "Spaced repetition queue update (SM-2 algorithm) \u2014 failed items get shorter intervals.",
    "[HOOK: Reward Chest] final session chest with variable rewards.",
    "[HOOK: Zeigarnik Effect] next lesson teaser creates incomplete mental loop \u2014 student thinks about next session involuntarily.",
    "Chapter progress ring creates visual tension to fill (endowed progress + completion drive).",
]:
    doc.add_paragraph(b, style="List Bullet")

# The NETS Loop
doc.add_heading("The NETS Loop (Duolingo-adapted)", level=1)
doc.add_paragraph("TRIGGER \u2192 ACTION \u2192 VARIABLE REWARD \u2192 INVESTMENT \u2192 (next day: internal trigger)")
doc.add_paragraph("External trigger (push notification) becomes internal trigger (streak fear, progress desire) over 2-3 weeks of consistent use.")

# Cross-cutting systems
doc.add_heading("Cross-Cutting Systems (UPDATED)", level=1)

doc.add_heading("Mystery Box (REWORKED \u2014 case-opening event, NOT gambling)", level=2)
p = doc.add_paragraph()
p.add_run("Purpose: ").bold = True
p.add_run("Variable-reward chest the student EARNS through performance. Removed any random-pay mechanics; appears either randomly during sessions or pre-placed by the engine after milestones.")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Reward pool (performance-weighted draw):").bold = True
for b in [
    "XP Boost \u2014 40%",
    "Extra AI (bonus AI tutor session) \u2014 25%",
    "Avatar Frame \u2014 15%",
    "Customization (color/theme) \u2014 10%",
    "Coupon \u2014 7%",
    "Mythical Boss Ticket \u2014 3%",
]:
    doc.add_paragraph(b, style="List Bullet")
p = doc.add_paragraph()
p.add_run("Choice rationale: ").bold = True
p.add_run("Case-opening framing (the box is THERE, the student opens it) avoids slot-machine ethics while keeping the variable-ratio dopamine hit. Mythical Boss Ticket at 3% creates the lottery dream without being predatory.")

doc.add_heading("\"Did You Know...\" Loading Insights (NEW)", level=2)
p = doc.add_paragraph()
p.add_run("Purpose: ").bold = True
p.add_run("Use the unavoidable phase-transition loading moments for tiny knowledge insertions. Free attention real-estate.")
for b in [
    "1-2 sentences max, theme-relevant to the current topic.",
    "1 insight per phase transition (so 6-7 per session).",
    "No quiz follow-up \u2014 pure curiosity injection, no pressure.",
    "Choice rationale: turns dead time into micro-learning without adding session length.",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Answer Checking \u2014 Dual Pathway (UPDATED)", level=2)
p = doc.add_paragraph()
p.add_run("Purpose: ").bold = True
p.add_run("Two routes for evaluation depending on question type. Eliminates the previous \"AI checks everything\" latency and cost problem.")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Pathway A \u2014 Hardcoded (instant):").bold = True
for b in [
    "Used for: multiple choice, matching, numeric answers.",
    "Rule-based exact match against the answer key.",
    "Zero AI cost, zero latency.",
]:
    doc.add_paragraph(b, style="List Bullet")
p = doc.add_paragraph()
p.add_run("Pathway B \u2014 Open Reasoning (AI-evaluated):").bold = True
for b in [
    "Used for: typed open responses, justifications, Real-Life Challenge explanations, boss reasoning.",
    "AI Tier 3 evaluates against a per-question rubric in real time.",
    "5-second timeout fallback: if AI doesn\u2019t respond, falls back to keyword matching against the rubric.",
    "Choice rationale: keeps the experience fast for 80% of questions and reserves AI cost for the questions that genuinely need judgment.",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Assessment Overhaul (NEW grading tiers)", level=2)
p = doc.add_paragraph()
p.add_run("Purpose: ").bold = True
p.add_run("Replace the old binary pass/fail with a tiered emotional response system that matches what the student actually achieved.")
doc.add_paragraph()
table = doc.add_table(rows=7, cols=2)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Score"
hdr[1].text = "Response"
for c in hdr:
    c.paragraphs[0].runs[0].bold = True
rows = [
    ("100%", "Celebration \u2014 full animation, Hall of Fame ping, max chest"),
    ("99%", "Tease \u2014 \"so close to perfect!\" + near-miss hook, replay invitation"),
    ("90-98%", "Excellent \u2014 strong reward, normal chest"),
    ("80-89%", "Very Good \u2014 bonus XP, encouragement"),
    ("60-79%", "Pass \u2014 standard reward, optional repass available to climb tiers"),
    ("<60%", "Duolingo Mode \u2014 student MUST redo the failed segments until they reach 60%. No skip."),
]
for i, (s, r) in enumerate(rows):
    table.rows[i + 1].cells[0].text = s
    table.rows[i + 1].cells[1].text = r
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Choice rationale: ").bold = True
p.add_run("The old \"you passed\" message felt the same at 60% and 100%. Tiered responses give the student something to chase. Duolingo Mode for <60% prevents the framework from rubber-stamping failure.")

# Section 9
doc.add_heading("Section 9: Customization Layer", level=1)
p = doc.add_paragraph()
p.add_run("Purpose: ").bold = True
p.add_run("Defines what CAN vs CANNOT be changed. Prevents well-meaning teachers from breaking the engine.")

doc.add_heading("Immutable (NEVER changes):", level=2)
for b in [
    "7-phase sequence and order",
    "Boss fight is mandatory (cannot be skipped)",
    "Spaced repetition runs automatically",
    "Bloom\u2019s coverage per session (all levels hit)",
    "PISA tracking is always on",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Customizable per teacher:", level=2)
for b in [
    "Session mode (standard/extended/recovery)",
    "Boss HP modifier (50-150%)",
    "Hint toggle (on/off)",
    "Game selection (within the 14 mechanics)",
    "Grace period for missed homework",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("AI-controlled per student:", level=2)
for b in [
    "Difficulty band (floor/target/ceiling)",
    "Game difficulty within each mechanic",
    "Adaptation speed",
    "Content selection and ordering",
]:
    doc.add_paragraph(b, style="List Bullet")

# Section 10
doc.add_heading("Section 10: Justification Index", level=1)
p = doc.add_paragraph()
p.add_run("Purpose: ").bold = True
p.add_run("Every design choice traced to its source \u2014 spec reference, research citation, or pedagogical principle. Ensures nothing is arbitrary.")
for b in [
    "Phase-to-spec mapping (each Blueprint section \u2192 UNIFIED \u00a7 reference)",
    "Research citations for every engagement mechanic",
    "Bloom\u2019s coverage verification (all 7 levels hit per session)",
    "Transition skill verification (\u22652 skills targeted, \u22651 matching student\u2019s in-progress skill)",
]:
    doc.add_paragraph(b, style="List Bullet")

# Timeline
doc.add_heading("Complete Session Timeline (Standard Mode, UPDATED 30-45 min)", level=1)
doc.add_paragraph("PRE (optional, student-paced): Theme Preview (0-A) \u2192 Flash Cards (0-B) \u2192 \"Start my Homework\"")
doc.add_paragraph("Then: Hook Game (30s) \u2192 Streak+Brief (30s) \u2192 Memory Sprint flexible (\u22642m) \u2192 Story1 (90s) \u2192 CP1 \u2192 Game1 (2-3m, may include Tic Tac Toe vs AI) \u2192 Story2 (90s) \u2192 CP2 \u2192 Game2 (2-3m) \u2192 Story3 (90s) \u2192 CP3 \u2192 Game3 (2-3m) \u2192 Real-Life Challenge first-person (4-6m) \u2192 Consolidation (2-3m) \u2192 BOSS [Sub / Big-weekly / Mythical-rare] (6-12m) \u2192 AI Analysis (<10s) \u2192 Remediation (2-3m) [Duolingo Mode if <60%] \u2192 Reflection (1-2m) \u2192 Closure (30s)")
doc.add_paragraph("\"Did You Know...\" insights play during every phase transition. Mystery Box may appear at any milestone or random tick.")

doc.save(os.path.join(OUT, "NETS-Homework-Engine-Universal-Blueprint-Summary.docx"))
print("Blueprint summary saved.")
