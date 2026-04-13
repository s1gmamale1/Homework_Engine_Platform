"""
Generate 3 demo lesson docx files following the NETS Homework Engine Blueprint.
Each lesson follows the 7-phase session structure:
  Phase 1: Memory Sprint
  Phase 2: Story Mode
  Phase 3: Game Breaks
  Phase 4: Real-Life Challenge
  Phase 5: Consolidation (Memory Palace)
  Phase 6: Final Boss
  Phase 7: Reflection + Remediation

Based on Blueprint Sections 1-8 and UNIFIED spec §5.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json, os

OUTPUT_DIR = r"C:\Users\DaddysHere\Documents\Sigma_Edu_3000\demo"
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ============================================================
# STYLING HELPERS
# ============================================================

def set_cell_shading(cell, color_hex):
    """Apply background color to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)

def add_styled_heading(doc, text, level=1, color="1B3A4B"):
    """Add a heading with consistent styling."""
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.size = Pt(14 if level == 1 else 12)
    return p

def add_meta_table(doc, meta_dict):
    """Add a metadata table at the top of each phase."""
    table = doc.add_table(rows=len(meta_dict), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (key, val) in enumerate(meta_dict.items()):
        cell_k = table.cell(i, 0)
        cell_v = table.cell(i, 1)
        cell_k.text = key
        cell_v.text = str(val)
        for run in cell_k.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
        for run in cell_v.paragraphs[0].runs:
            run.font.size = Pt(9)
        set_cell_shading(cell_k, "E8F0F2")
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.2)
        row.cells[1].width = Inches(4.5)
    return table

def add_question_block(doc, num, question, correct_answer, tags, 
                       bloom="", pisa="", transition="", hint1="", hint2="",
                       common_errors=None):
    """Add a standardized question block per Blueprint §6.5 schema."""
    p = doc.add_paragraph()
    run = p.add_run(f"Savol {num}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("2C3E50")
    
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(question)
    run.font.size = Pt(10)
    
    if hint1:
        p = doc.add_paragraph()
        run = p.add_run(f"💡 Ko'rsatma: {hint1}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("7F8C8D")
        run.italic = True
    if hint2:
        p = doc.add_paragraph()
        run = p.add_run(f"💡 Ko'rsatma (2): {hint2}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("7F8C8D")
        run.italic = True
    
    p = doc.add_paragraph()
    run = p.add_run(f"✅ Javob: {correct_answer}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("27AE60")
    run.bold = True
    
    # Tags
    tag_text = f"[{tags['standard_ref']}] | Bloom: {bloom} | PISA: {pisa}"
    if transition:
        tag_text += f" | Transition: {transition}"
    p = doc.add_paragraph()
    run = p.add_run(tag_text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("95A5A6")
    
    doc.add_paragraph()  # spacer

def set_default_font(doc):
    """Set default font for the document."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)


# ============================================================
# LESSON GENERATOR FUNCTIONS
# ============================================================

def generate_kimyo_lesson():
    """
    Kimyo Grade 7, Section 33: "Kislorodning kimyoviy xossalari"
    IV BOB, 5-MAVZU | Pages 95-99
    Blueprint Phases 1-7
    """
    doc = Document()
    set_default_font(doc)
    
    # ---- TITLE PAGE ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NETS HOMEWORK ENGINE — DEMO LESSON")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string("1B3A4B")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Kimyo · 7-sinf")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string("2C3E50")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("IV Bob, 5-MAVZU: Kislorodning kimyoviy xossalari")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string("34495E")
    run.italic = True
    
    # Metadata block
    add_meta_table(doc, {
        "Dars": "Kislorodning kimyoviy xossalari va yonish",
        "Sinf": "7",
        "Fan": "Kimyo",
        "Darslik": "Kimyo 7-sinf (2022 nashr), sah. 95-99",
        "Standart kod": "UZ-CHEM-7-O2-05 (Kislorodning kimyoviy xossalari)",
        "PISA maqsadi": "Daraja 2 → 3 (Science: Explain phenomena)",
        "Bloom darajalari": "Remember, Understand, Apply, Analyze",
        "O'tish ko'nikmasi": "L2→L3: predict reaction products; interpret chemical data",
        "Rejim": "Standard (20-30 daqiqa)",
        "Boss HP": "80 (Grades 5-8)",
    })
    doc.add_paragraph()

    # ========================================
    # PHASE 1: MEMORY SPRINT
    # ========================================
    add_styled_heading(doc, "⚡ 1-FAZA: XOTIRA HURRASI (Memory Sprint) — 2 daqiqa", level=2)
    add_meta_table(doc, {
        "Davomiylik": "60 soniya",
        "Savollar": "5 ta (oldingi boblardan)",
        "Manba": "I-III boblar (Moddalar, Atom, Davriy jadval)",
        "AI Darajasi": "Tier 1 (algoritmik)",
        "XP": "+100 XP to'g'ri javob uchun",
    })
    doc.add_paragraph()
    
    questions_p1 = [
        ("1", "Moddaning uchta agregat holatini yozing.",
         "Qattiq, suyuq, gazsimon",
         "UZ-CHEM-7-STATE-01", "remember", "1"),
        ("2", "Atom qanday zarrachalardan tashkil topgan?",
         "Proton, neytron, elektron",
         "UZ-CHEM-7-ATOM-02", "remember", "1"),
        ("3", "Suvning kimyoviy formulasi nima?",
         "H₂O",
         "UZ-CHEM-7-H2O-01", "remember", "1"),
        ("4", "Kimyoviy elementning nisbiy atom massasi nima bilan o'lchanadi?",
         "Atom massa birligi (a.m.b.)",
         "UZ-CHEM-7-AMASS-01", "remember", "1"),
        ("5", "Davriy jadvalda kislorod qaysi guruhda joylashgan?",
         "VI guruh (16-guruh)",
         "UZ-CHEM-7-PERIOD-01", "remember", "1"),
    ]
    
    for num, q, ans, std, bloom, pisa in questions_p1:
        add_question_block(doc, num, q, ans, {"standard_ref": std}, 
                          bloom=bloom, pisa=f"Level {pisa}")
    
    # Phase 1 gate
    p = doc.add_paragraph()
    run = p.add_run("⚙️ Moslashtirish: Agar aniqlik < 60% → tuzatishga yo'naltirish. "
                    "≥ 80% → yangi kontentga tayyor.")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("95A5A6")
    run.italic = True
    doc.add_page_break()

    # ========================================
    # PHASE 2: STORY MODE
    # ========================================
    add_styled_heading(doc, "📖 2-FAZA: HIKOYA REJIMI (Story Mode) — 5-7 daqiqa", level=2)
    add_meta_table(doc, {
        "Davomiylik": "5-7 daqiqa",
        "Segmentlar": "3 ta (har biri ~90 soniya)",
        "CPA bosqichi": "Pictorial → Abstract",
        "Manba": "Kimyo 7-sinf, IV Bob, 5-MAVZU, sah. 95-96",
        "Checkpoint": "Har segment oxirida 1-2 ta (o'tish uchun majburiy)",
    })
    doc.add_paragraph()
    
    # Segment 1
    p = doc.add_paragraph()
    run = p.add_run("📕 Segment 1: Kislorod — hayotning yonilg'isi")
    run.bold = True
    run.font.size = Pt(11)
    
    story1 = (
        "Tasavvur qiling — siz katta temir simni olovdan qizdirib, "
        "kislorodli kolbaga joyladingingiz. Sim bir zumda yorqin "
        "porlay boshladi, xuddi Bengal olovi kabi uchqunlar sochdi! "
        "Bu nima uchun sodir bo'ldi?\n\n"
        "Kislorod (O₂) — kimyoviy jihatdan juda faol modda. U ko'plab "
        "boshqa moddalar bilan reaksiyaga kirishishga qodir. Kislorod "
        "asosan II valentlikni namoyon etadi.\n\n"
        "Kislorod bilan reaksiyalarning aksariyati katta miqdorda "
        "issiqlik va yorug'lik chiqaradi. Bunday jarayonlar yonish deb ataladi."
    )
    p = doc.add_paragraph(story1)
    p.paragraph_format.left_indent = Cm(0.5)
    
    # Checkpoint 1
    p = doc.add_paragraph()
    run = p.add_run("🔒 Checkpoint 1: Yonish deb nimaga aytiladi?")
    run.bold = True
    p = doc.add_paragraph(
        "✅ Javob: Moddalarning kislorod bilan reaksiyasi natijasida "
        "issiqlik va yorug'lik ajralishi bilan boradigan jarayon.")
    
    # Segment 2
    p = doc.add_paragraph()
    run = p.add_run("📕 Segment 2: Kislorod metallar bilan")
    run.bold = True
    run.font.size = Pt(11)
    
    story2 = (
        "Metallar kislorod bilan qizdirilganda oksidlar hosil qiladi:\n\n"
        "  2Mg + O₂ → 2MgO  (magniy oksidi)\n"
        "  2Cu + O₂ → 2CuO  (mis (II) oksidi)\n"
        "  3Fe + 2O₂ → Fe₃O₄ (temir oksidi — magnit temirtosh)\n\n"
        "Ishqoriy metallar esa peroksidlar ham hosil qilishi mumkin:\n"
        "  2Na + O₂ → Na₂O₂  (natriy peroksidi)\n\n"
        "Metallmaslar bilan ham reaksiyaga kirishadi:\n"
        "  S + O₂ → SO₂  (oltingugurt (IV) oksidi)\n"
        "  C + O₂ → CO₂  (karbonat angidrid)\n"
        "  4P + 5O₂ → 2P₂O₅  (fosfor (V) oksidi)"
    )
    p = doc.add_paragraph(story2)
    p.paragraph_format.left_indent = Cm(0.5)
    
    # Checkpoint 2
    p = doc.add_paragraph()
    run = p.add_run("🔒 Checkpoint 2: Temir kislorodda yonganda qanday birikma hosil bo'ladi?")
    run.bold = True
    p = doc.add_paragraph(
        "✅ Javob: Fe₃O₄ (temirning aralash oksidi — FeO·Fe₂O₃, "
        "unda bitta temir II valentli, ikkitasi III valentli).")
    
    # Segment 3
    p = doc.add_paragraph()
    run = p.add_run("📕 Segment 3: Organik moddalar yonishi")
    run.bold = True
    run.font.size = Pt(11)
    
    story3 = (
        "Deyarli barcha organik moddalar kislorodda yonib, "
        "karbonat angidrid va suv hosil qiladi:\n\n"
        "  CH₄ + 2O₂ → CO₂ + 2H₂O  (metan yonishi — tabiiy gaz)\n"
        "  C₂H₅OH + 3O₂ → 2CO₂ + 3H₂O  (etil spirti yonishi)\n\n"
        "Oksidlar — bu biri kisloroddan iborat binar birikmalar:\n"
        "CaO (kalsiy oksidi), SO₂ (oltingugurt (IV) oksidi), "
        "Al₂O₃ (alyuminiy oksidi), CuO (mis (II) oksidi)."
    )
    p = doc.add_paragraph(story3)
    p.paragraph_format.left_indent = Cm(0.5)
    
    # Checkpoint 3
    p = doc.add_paragraph()
    run = p.add_run("🔒 Checkpoint 3: Metan yonganda qanday moddalar hosil bo'ladi?")
    run.bold = True
    p = doc.add_paragraph("✅ Javob: CO₂ (karbonat angidrid) va H₂O (suv)")

    # ========================================
    # PHASE 3: GAME BREAKS
    # ========================================
    add_styled_heading(doc, "🎮 3-FAZA: O'YIN TANAFUSLARI (Game Breaks) — 6-9 daqiqa", level=2)
    add_meta_table(doc, {
        "Davomiylik": "6-9 daqiqa",
        "O'yin mexaniklari": "Tile Match + Sentence Fill + Mystery Box",
        "Bloom darajalari": "Apply, Analyze",
    })
    doc.add_paragraph()
    
    # --- Tile Match ---
    p = doc.add_paragraph()
    run = p.add_run("🧩 Tile Match: Reaksiya tenglamalarini juftlang")
    run.bold = True
    p = doc.add_paragraph(
        "Har bir metallni uning kislorod bilan reaksiya mahsuloti bilan juftlang:\n\n"
        "  1. Mg + O₂        →  A. Fe₃O₄\n"
        "  2. Cu + O₂        →  B. MgO\n"
        "  3. Fe + O₂        →  C. CuO\n"
        "  4. Na + O₂        →  D. Na₂O₂\n\n"
        "✅ Javob: 1-B, 2-C, 3-A, 4-D")
    
    # --- Sentence Fill ---
    p = doc.add_paragraph()
    run = p.add_run("✏️ Sentence Fill: Bo'sh joylarni to'ldiring")
    run.bold = True
    p = doc.add_paragraph(
        "1. Kislorod ___________ modda. (j: kimyoviy faol)\n"
        "2. Yonish — moddalarning _________ bilan reaksiyasi. (j: kislorod)\n"
        "3. Oksid — biri __________ dan iborat binar birikma. (j: kislorod)\n"
        "4. CH₄ + 2O₂ → ___ + 2H₂O (j: CO₂)\n"
        "5. 4P + 5O₂ → 2______ (j: P₂O₅)")
    
    # --- Mystery Box ---
    p = doc.add_paragraph()
    run = p.add_run("🎁 Mystery Box: Noma'lum mahsulotni toping")
    run.bold = True
    p = doc.add_paragraph(
        "X qanday modda?\n\n"
        "  S + O₂ → X\n"
        "  X + O₂ → Y (katalizator V₂O₅ ishtirokida)\n\n"
        "💡 Ko'rsatma: IV va VI guruh elementlari.\n"
        "✅ Javob: X = SO₂ (oltingugurt (IV) oksidi), Y = SO₃ (oltingugurt (VI) oksidi)")
    
    doc.add_page_break()

    # ========================================
    # PHASE 4: REAL-LIFE CHALLENGE
    # ========================================
    add_styled_heading(doc, "🌍 4-FAZA: HAQIY HAYOT MUAMMOSI (Real-Life Challenge) — 3-5 daqiqa", level=2)
    add_meta_table(doc, {
        "Davomiylik": "3-5 daqiqa",
        "Bloom darajasi": "Analyze, Evaluate",
        "PISA darajasi": "3 (Science: interpret data, use evidence)",
        "O'tish ko'nikmasi": "L2→L3: integrate multiple sources; communicate reasoning",
        "O'zbekiston konteksti": "Yong'in xavfsizligi, tabiiy gaz, pichan o'z-o'zidan yonishi",
    })
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("🏠 Vaziyat: Qishloq xo'jaligi va yong'in xavfsizligi")
    run.bold = True
    
    scenario = (
        "Qishloq joyda dehqon Ahmad aka pichanzorda pichan yig'ayotgan edi. "
        "U nam pichanni omborga taxladi. Bir necha kun o'tgach, "
        "ombordan tutun chiqayotganini ko'rdi — pichan o'z-o'zidan "
        "yonayotgan edi!\n\n"
        "1. Pichan nima uchun o'z-o'zidan yondi? Jarayonni kimyoviy "
        "tenglamalar bilan tushuntiring.\n"
        "2. Agar omborda kislorod miqdori 15% dan kam bo'lsa, yonish "
        "mumkin bo'ladimi? Nima uchun?\n"
        "3. Olovni o'chirish uchun suv ishlatish mumkinmi? Nima uchun "
        "ba'zi hollarda suv ishlatib bo'lmaydi?\n"
        "4. Yong'in paytida birinchi navbatda nimaga e'tibor berish kerak?"
    )
    p = doc.add_paragraph(scenario)
    
    p = doc.add_paragraph()
    run = p.add_run("✅ To'liq javob kaliti:\n"
        "1. Sekin oksidlanish — pichandagi organik moddalar kislorod "
        "bilan sekin reaksiyaga kirishadi, issiqlik chiqadi. Issiqlik "
        "omborda to'planib, yonish haroratiga yetadi. "
        "Organik + O₂ → CO₂ + H₂O + issiqlik\n"
        "2. Yo'q, 15% dan kam kislorod bo'lsa yonish mumkin emas.\n"
        "3. Ba'zi metallar (Na, K) suv bilan reaksiyaga kirishadi — "
        "yong'in kuchayadi. Benzin, kerosin suvdan yengil — "
        "suv bilan o'chirib bo'lmaydi.\n"
        "4. Odam hayoti xavfsizligi, yong'in xizmatini chaqirish, "
        "yaqin atrofdagi odamlarni ogohlantirish.")
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor.from_string("27AE60")
    
    # ========================================
    # PHASE 5: CONSOLIDATION (Memory Palace)
    # ========================================
    add_styled_heading(doc, "🏰 5-FAZA: XOTIRA SAROYI (Consolidation) — 2-3 daqiqa", level=2)
    add_meta_table(doc, {
        "Davomiylik": "2-3 daqiqa",
        "Texnika": "Memory Palace + Flashcard",
        "Maqsad": "Asosiy tushunchalarni Boss jangidan oldin mustahkamlash",
    })
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("🧠 Xotira saroyi — Kislorod laboratoriyasini tasavvur qiling:\n\n"
        "1. Eshik oldida — O₂ belgisi (kislorod formulasi)\n"
        "2. Birinchi xona — Metallar (Mg, Cu, Fe) → Oksidlar (MgO, CuO, Fe₃O₄)\n"
        "3. Ikkinchi xona — Metallmaslar (S, C, P) → Oksidlar (SO₂, CO₂, P₂O₅)\n"
        "4. Uchinchi xona — Organik moddalar (CH₄, C₂H₅OH) → CO₂ + H₂O\n"
        "5. To'rtinchi xona — Yonish shartlari: O₂ + yonish harorati + yonilg'i\n"
        "6. Beshinchi xona — Yong'in o'chirish: suv, qum, ko'pik, asbest\n\n"
        "Flashcard tezkor mashq:\n"
        "• Yonish nima? → Issiqlik+yorug'lik bilan boruvchi O₂ reaksiyasi\n"
        "• Oksid nima? → Biri kisloroddan iborat binar birikma\n"
        "• Sekin oksidlanish? → Olovsiz sekin ta'sir (zanglash)\n"
        "• Yong'in nima? → Nazoratsiz qolgan yonish")
    p.runs[0].font.size = Pt(10)

    # ========================================
    # PHASE 6: FINAL BOSS
    # ========================================
    add_styled_heading(doc, "👹 6-FAZA: YAKUNIY BOSS (Final Boss) — 5-10 daqiqa", level=2)
    add_meta_table(doc, {
        "Davomiylik": "5-10 daqiqa",
        "Boss HP": "80 (5-8 sinflar)",
        "Savollar": "5 ta (PISA 2-4, Bloom: Apply→Analyze→Evaluate)",
        "Format": "Ochiq javob (MC yo'q, 5+ sinf uchun)",
        "Sokrat ko'rsatmalari": "3 daraja (har savolda)",
        "O'tish mezon": "Boss HP = 0 → g'alaba",
        "Maksimal urinishlar": "3 ta (har urinishda savollar qayta generatsiya qilinadi)",
    })
    doc.add_paragraph()
    
    boss_questions = [
        ("1", 
         "12g magniy kislorodda yondirildi. Necha gramm magniy oksidi (MgO) hosil bo'ladi? "
         "To'liq yechimni ko'rsating.\n\n"
         "(Ar(Mg)=24, Ar(O)=16)",
         "2Mg + O₂ → 2MgO\n"
         "n(Mg) = 12g / 24g/mol = 0.5 mol\n"
         "n(MgO) = 0.5 mol\n"
         "m(MgO) = 0.5 mol × 40g/mol = 20g",
         "20 ball (to'liq yechim), 10 ball (usul to'g'ri, arifmetik xato)",
         "1) Mg + O₂ → MgO tenglamani tenglashtiring\n"
         "2) Mol massani hisoblang: M(MgO) = ?\n"
         "3) Proporsiya tuzing",
         "UZ-CHEM-7-O2-05", "apply", "3",
         "L2→L3: multi-step procedure with strategy selection"),
        
        ("2",
         "Quyidagi reaksiyalarni tenglashtiring va mahsulotlarni aniqlang:\n"
         "a) Fe + O₂ → ?\n"
         "b) CH₄ + O₂ → ?\n"
         "c) NH₃ + O₂ → ? (katalizator bilan)",
         "a) 3Fe + 2O₂ → Fe₃O₄\n"
         "b) CH₄ + 2O₂ → CO₂ + 2H₂O\n"
         "c) 4NH₃ + 5O₂ → 4NO + 6H₂O (katalizator)",
         "15 ball (har biri 5 balldan)",
         "1) Valentlikni eslang: O = II\n"
         "2) Organik moddalar → CO₂ + H₂O\n"
         "3) NH₃ katalitik oksidlanishi → NO",
         "UZ-CHEM-7-O2-05", "analyze", "3",
         "L2→L3: select/integrate representations"),
        
        ("3",
         "Nima uchun oq fosforni faqat suv ostida saqlash kerak? "
         "Bu qaysi kimyoviy hodisaga bog'liq? Kundalik hayotda shunga o'xshash "
         "holatni keltiring.",
         "Oq fosforning yonish harorati xona haroratiga yaqin. U havoda "
         "o'z-o'zidan yonadi (sekin oksidlanish → yonish). Shuning uchun "
         "suv ostida saqlanadi. Kundalik hayotda: nam pichanning omborda "
         "o'z-o'zidan yonishi.",
         "15 ball (to'liq tushuntirish + misol), 8 ball (qisman)",
         "1) 'O'z-o'zidan yonish' nima?\n"
         "2) Xona haroratiga yaqin yonish harorati nimani anglatadi?",
         "UZ-CHEM-7-O2-06", "analyze", "3",
         "L3→L4: construct explanations"),
        
        ("4",
         "Yong'in paytida ba'zi moddalarni suv bilan o'chirib bo'lmaydi. "
         "3 ta holatni keltiring va har birini kimyoviy jihatdan tushuntiring.",
         "1) Natriy, kaliy — suv bilan reaksiyaga kirishadi, vodorod chiqadi, "
         "portlash mumkin\n"
         "2) Benzin, kerosin — suvdan yengil, yuzada suzib yonishda davom etadi\n"
         "3) Elektr jihozlari — suv elektr o'tkazgich, elektr toki urishi mumkin",
         "15 ball (3 ta holat + tushuntirish), 10 ball (2 ta), 5 ball (1 ta)",
         "1) Metallarning suv bilan reaksiyasini eslang\n"
         "2) Suvning zichligi bilan yonilg'ining zichligini taqqoslang",
         "UZ-CHEM-7-O2-06", "evaluate", "4",
         "L3→L4: explicit models; construct explanations"),
        
        ("5",
         "Tajriba: 3 ta shisha idish. Birinchisida sham, ikkinchisida o'simlik + sham (yorug'da), "
         "uchinchisida o'simlik + sham (qorong'ida). Qaysi sham birinchi o'chadi? "
         "Nima uchun? Kislorod balansini hisobga olib tushuntiring.",
         "Uchinchi idishdagi sham (qorong'ida) birinchi o'chadi. Chunki:\n"
         "• Birinchi idishda: sham O₂ ni iste'mol qiladi → o'chadi\n"
         "• Ikkinchi idishda: fotosintez O₂ ishlab chiqaradi → sham yonishda davom etadi\n"
         "• Uchinchi idishda: fotosintez yo'q (qorong'ida), o'simlik ham nafas oladi "
         "(O₂ iste'mol qiladi) → sham tezroq o'chadi",
         "20 ball (to'liq tushuntirish + 3 holat), 10 ball (qisman)",
         "1) Fotosintez uchun nima kerak?\n"
         "2) O'simlik qorong'ida nima qiladi?\n"
         "3) Sham yonishi uchun nima kerak?",
         "UZ-CHEM-7-O2-05", "evaluate", "4",
         "L3→L4: compare strategies; reflect on results"),
    ]
    
    for num, q, ans, rubric, socratic, std, bloom, pisa, transition in boss_questions:
        p = doc.add_paragraph()
        run = p.add_run(f"⚔️ BOSS SAVOL {num}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string("C0392B")
        
        p = doc.add_paragraph(q)
        p.runs[0].font.size = Pt(10)
        
        # Socratic hints
        p = doc.add_paragraph()
        run = p.add_run(f"Sokrat ko'rsatmalari:\n{socratic}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("7F8C8D")
        run.italic = True
        
        p = doc.add_paragraph()
        run = p.add_run(f"✅ Rubrika: {rubric}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("27AE60")
        run.bold = True
        
        p = doc.add_paragraph()
        run = p.add_run(f"[{std}] | Bloom: {bloom} | PISA L{pisa} | {transition}")
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor.from_string("95A5A6")
        
        p = doc.add_paragraph(f"💰 Damage: 20 HP (to'liq), 10 HP (qisman), 0 HP (xato)")
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.color.rgb = RGBColor.from_string("E67E22")
        
        doc.add_paragraph()
    
    doc.add_page_break()

    # ========================================
    # PHASE 7: REFLECTION + REMEDIATION
    # ========================================
    add_styled_heading(doc, "🪞 7-FAZA: AKS ETTIRISH VA TUZATISH — 3-5 daqiqa", level=2)
    add_meta_table(doc, {
        "Davomiylik": "3-5 daqiqa",
        "Komponentlar": "Zaif nuqtalar → Mikro-mashqlar → Refleksiya → Spaced Repetition",
        "AI Darajasi": "Tier 1 (hisoblash) + Tier 2 (hisobot generatsiya)",
    })
    doc.add_paragraph()
    
    # Weakness display template
    p = doc.add_paragraph()
    run = p.add_run("📊 BUGUNGI NATIJALAR:\n\n"
        "✅ Kislorodning umumiy xossalari — Zo'r!\n"
        "✅ Metallarning yonishi — Yaxshi\n"
        "⚠️ Organik moddalarning yonishi — Mashq kerak\n"
        "⚠️ Reaksiya tenglamalarini tenglashtirish — Mashq kerak\n\n"
        "Keling, shu mavzularni birgalikda mustahkamlaymiz!")
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor.from_string("2C3E50")
    
    # Micro-exercises
    p = doc.add_paragraph()
    run = p.add_run("🔧 Mikro-mashqlar (BONUS XP IMKONIYATI):\n\n"
        "1. CH₄ + ___ O₂ → CO₂ + 2H₂O  (j: 2)\n"
        "2. C₂H₅OH + 3O₂ → 2___ + 3H₂O  (j: CO₂)\n"
        "3. 2Mg + ___ → 2MgO  (j: O₂)\n\n"
        "Har bir to'g'ri javob: +100 XP")
    p.runs[0].font.size = Pt(10)
    
    # Reflection prompt
    p = doc.add_paragraph()
    run = p.add_run("💭 Refleksiya (kamida 10 belgi):\n\n"
        "\"Bugun men kislorod haqida yangi nima o'rgandim? "
        "Qaysi reaksiya menga eng qiziq tuyuldi? Nima uchun?\"")
    p.runs[0].font.size = Pt(10)
    
    # Spaced repetition
    p = doc.add_paragraph()
    run = p.add_run("🔄 Spaced Repetition (Ebbinghaus):\n"
        "• Kislorod reaksiyalari (kuchli, >80%) → 7 kundan keyin takrorlash\n"
        "• Organik moddalar yonishi (barqaror) → 3 kundan keyin\n"
        "• Reaksiya tenglashtirish (zaif, <60%) → ERTAGA, Memory Sprintda ustuvor")
    p.runs[0].font.size = Pt(9)
    
    # Session closure
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("🎉 SESSIYA TUGADI!\n"
        "⬜ XP: 2,450 | ⭐⭐ 2 yulduz | 🔥 13 kunlik seriya\n"
        "📈 PISA: 1.5 → 1.6 (+0.1)\n"
        "Keyingi dars: IV Bob, 6-MAVZU — Yonish turlari va yonilg'i")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("27AE60")
    
    path = os.path.join(OUTPUT_DIR, "NETS_Demo_Kimyo_7_Kislorod.docx")
    doc.save(path)
    print(f"✅ Kimyo saved: {path}")


def generate_algebra_lesson():
    """
    Algebra Grade 7, Section 33: "Chiziqli funksiya"
    5-BOB, Section 3 | Pages 120-127
    """
    doc = Document()
    set_default_font(doc)
    
    # TITLE
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NETS HOMEWORK ENGINE — DEMO LESSON")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string("1B3A4B")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Algebra · 7-sinf")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string("2C3E50")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("5-Bob, 3-MAVZU: Chiziqli funksiya (y = kx + b)")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string("34495E")
    run.italic = True
    
    add_meta_table(doc, {
        "Dars": "Chiziqli funksiya va uning grafigi",
        "Sinf": "7",
        "Fan": "Matematika (Algebra)",
        "Darslik": "Algebra 7-sinf (2022 nashr), sah. 120-127",
        "Standart kod": "UZ-MATH-7-LINFUNC-03",
        "PISA maqsadi": "Daraja 2 → 3 (Math: employ, interpret)",
        "Bloom darajalari": "Understand, Apply, Analyze",
        "O'tish ko'nikmasi": "L2→L3: integrate representations; multi-step procedures",
        "Rejim": "Standard (20-30 daqiqa)",
        "Boss HP": "80 (Grades 5-8)",
    })
    doc.add_paragraph()

    # PHASE 1: MEMORY SPRINT
    add_styled_heading(doc, "⚡ 1-FAZA: XOTIRA HURRASI — 2 daqiqa", level=2)
    add_meta_table(doc, {
        "Davomiylik": "60 soniya",
        "Savollar": "5 ta (1-4 boblardan)",
        "Manba": "Oldingi boblar: Ko'phadlar, Darajalar, Tenglama",
        "XP": "+100 XP to'g'ri javob uchun",
    })
    doc.add_paragraph()
    
    questions_p1 = [
        ("1", "y = 3x funksiyada x = 4 bo'lsa, y = ?",
         "12",
         "UZ-MATH-7-PROPO-01", "remember", "1"),
        ("2", "Agar k > 0 bo'lsa, y = kx grafigi qaysi choraklarda joylashadi?",
         "I va III choraklar",
         "UZ-MATH-7-LINFUNC-01", "remember", "1"),
        ("3", "Dekart koordinatalar sistemasida (3; -2) nuqta qaysi chorakda?",
         "IV chorak",
         "UZ-MATH-7-COORD-01", "understand", "1"),
        ("4", "2x + 6 = 0 tenglamaning ildizini toping.",
         "x = -3",
         "UZ-MATH-7-EQ-01", "apply", "2"),
        ("5", "y = -x funksiyasi qaysi choraklardan o'tadi?",
         "II va IV choraklar",
         "UZ-MATH-7-LINFUNC-02", "understand", "1"),
    ]
    
    for num, q, ans, std, bloom, pisa in questions_p1:
        add_question_block(doc, num, q, ans, {"standard_ref": std},
                          bloom=bloom, pisa=f"Level {pisa}")

    # Phase 1 gate (Algebra-specific)
    p = doc.add_paragraph()
    run = p.add_run("⚙️ Moslashtirish: <60% → tuzatish, 60-79% → davom et, ≥80% → tayyor.")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("95A5A6")
    run.italic = True
    doc.add_page_break()

    # PHASE 2: STORY MODE (Algebra)
    add_styled_heading(doc, "📖 2-FAZA: HIKOYA REJIMI — 5-7 daqiqa", level=2)
    add_meta_table(doc, {
        "Davomiylik": "5-7 daqiqa",
        "Segmentlar": "3 ta",
        "CPA bosqichi": "Concrete (kvadrat perimetri) → Pictorial (jadval, grafik) → Abstract (y = kx + b)",
        "Manba": "Algebra 7-sinf, 5-Bob, sah. 120-124",
    })
    doc.add_paragraph()
    
    # Segment 1
    p = doc.add_paragraph()
    run = p.add_run("📕 Segment 1: Kvadrat siri — y = kx")
    run.bold = True
    run.font.size = Pt(11)
    
    p = doc.add_paragraph(
        "Kvadrat tomoni a = 1 bo'lsa, perimetri P = 4. a = 2 bo'lsa, P = 8. "
        "a = 3 bo'lsa, P = 12. Har safar tomonni 1 ga oshirsak, "
        "perimetr 4 ga oshadi. Bu — chiziqli bog'liqlik!\n\n"
        "P = 4a formulani y = kx ko'rinishida yozsak: y = 4x.\n"
        "Bu yerda k = 4 — koeffitsiyent (egrilik). "
        "Grafik — koordinata boshidan o'tuvchi to'g'ri chiziq.")
    p.paragraph_format.left_indent = Cm(0.5)
    
    p = doc.add_paragraph()
    run = p.add_run("🔒 Checkpoint 1: y = 3x funksiyada x = 5 bo'lsa y = ?")
    run.bold = True
    p = doc.add_paragraph("✅ Javob: y = 15")
    
    # Segment 2
    p = doc.add_paragraph()
    run = p.add_run("📕 Segment 2: k ning ahamiyati")
    run.bold = True
    run.font.size = Pt(11)
    
    p = doc.add_paragraph(
        "k > 0 bo'lsa → chiziq I va III chorakda (o'suvchi)\n"
        "k < 0 bo'lsa → chiziq II va IV chorakda (kamayuvchi)\n"
        "k = 0 bo'lsa → chiziq Ox o'qi bilan ustma-ust\n\n"
        "Misol: y = 2x (o'suvchi), y = -2x (kamayuvchi). "
        "Ikkala chiziq ham (0; 0) nuqtadan o'tadi. k qanchalik katta bo'lsa, "
        "chiziq shunchalik tik bo'ladi.")
    p.paragraph_format.left_indent = Cm(0.5)
    
    p = doc.add_paragraph()
    run = p.add_run("🔒 Checkpoint 2: y = -3x funksiyasi qaysi choraklarda joylashadi?")
    run.bold = True
    p = doc.add_paragraph("✅ Javob: II va IV choraklar (chunki k = -3 < 0)")
    
    # Segment 3
    p = doc.add_paragraph()
    run = p.add_run("📕 Segment 3: y = kx + b — siljitilgan chiziq")
    run.bold = True
    run.font.size = Pt(11)
    
    p = doc.add_paragraph(
        "b soni grafikni Oy o'qi bo'ylab siljitadi:\n"
        "b > 0 → yuqoriga, b < 0 → pastga, b = 0 → y = kx (boshlang'ich)\n\n"
        "Misol: y = 2x + 1 → b = 1 (Oy o'qini 1 nuqtada kesib o'tadi)\n"
        "Misol: y = -4x - 4 → b = -4 (Oy o'qini -4 nuqtada kesib o'tadi)\n\n"
        "Grafikdan formulani topish:\n"
        "1) b = Oy o'q bilan kesishgan nuqta\n"
        "2) k = vertikal masofa / gorizontal masofa")
    p.paragraph_format.left_indent = Cm(0.5)
    
    p = doc.add_paragraph()
    run = p.add_run("🔒 Checkpoint 3: Grafik Oy o'qini (0; 3) nuqtada kesssa, b = ?")
    run.bold = True
    p = doc.add_paragraph("✅ Javob: b = 3")

    # PHASE 3: GAME BREAKS
    add_styled_heading(doc, "🎮 3-FAZA: O'YIN TANAFUSLARI — 6-9 daqiqa", level=2)
    add_meta_table(doc, {
        "Davomiylik": "6-9 daqiqa",
        "O'yin mexaniklari": "Tile Match + Sentence Fill + Mystery Box",
    })
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("🧩 Tile Match: Funksiyani grafigi bilan juftlang")
    run.bold = True
    p = doc.add_paragraph(
        "1. y = 2x          →  A. I va III chorak, tik chiziq\n"
        "2. y = -x          →  B. I va III chorak, yotiq chiziq\n"
        "3. y = 0.5x        →  C. II va IV chorak\n"
        "4. y = -3x + 2     →  D. II va IV chorak, b = 2\n\n"
        "✅ Javob: 1-A, 2-C, 3-B, 4-D")
    
    p = doc.add_paragraph()
    run = p.add_run("✏️ Sentence Fill: Bo'sh joylarni to'ldiring")
    run.bold = True
    p = doc.add_paragraph(
        "1. y = kx + b funksiya ___________ funksiya deyiladi. (j: chiziqli)\n"
        "2. y = 5x - 3 funksiyada k = ___, b = ___. (j: 5, -3)\n"
        "3. y = -2x + 4 Oy o'qini ___ nuqtada kesadi. (j: (0; 4))\n"
        "4. y = 3x funksiya grafigi ___ nuqtadan o'tadi. (j: (0; 0))")
    
    p = doc.add_paragraph()
    run = p.add_run("🎁 Mystery Box: Noma'lum funksiya")
    run.bold = True
    p = doc.add_paragraph(
        "Funksiya grafigi (0; -2) nuqtadan o'tadi va x = 1 bo'lganda y = 1.\n"
        "Funksiya formulasini toping.\n\n"
        "💡 Ko'rsatma: b = -2. Keyin k ni toping.\n"
        "✅ Javob: y = 3x - 2  (b = -2, 1 = k(1) - 2, k = 3)")
    
    doc.add_page_break()

    # PHASE 4: REAL-LIFE CHALLENGE
    add_styled_heading(doc, "🌍 4-FAZA: HAQIY HAYOT MUAMMOSI — 3-5 daqiqa", level=2)
    add_meta_table(doc, {
        "Davomiylik": "3-5 daqiqa",
        "Bloom": "Analyze, Evaluate",
        "PISA": "3",
        "Kontekst": "Taksi narxi — kundalik hayotda chiziqli funksiya",
    })
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("🚕 Vaziyat: Taksi narxi")
    run.bold = True
    
    p = doc.add_paragraph(
        "Toshkentdagi taksi xizmati: boshlang'ich narx 5,000 so'm, "
        "har 1 km uchun 2,000 so'm qo'shiladi.\n\n"
        "1. Narx formulasini y = kx + b ko'rinishida yozing.\n"
        "2. 7 km yo'l uchun qancha to'lash kerak?\n"
        "3. 25,000 so'mingiz bo'lsa, maksimum necha km borishingiz mumkin?\n"
        "4. Ikkinchi taksi xizmati: boshlang'ich 3,000 so'm, har 1 km = 2,500 so'm. "
        "Qaysi biri 10 km uchun arzonroq?\n\n"
        "✅ Javob kaliti:\n"
        "1. y = 2000x + 5000\n"
        "2. y = 2000(7) + 5000 = 19,000 so'm\n"
        "3. 25000 = 2000x + 5000 → x = 10 km\n"
        "4. 1-xizmat: y = 2000(10)+5000 = 25,000\n"
        "   2-xizmat: y = 2500(10)+3000 = 28,000 → 1-xizmat arzonroq")
    p.runs[0].font.size = Pt(10)

    # PHASE 5: CONSOLIDATION
    add_styled_heading(doc, "🏰 5-FAZA: XOTIRA SAROYI — 2-3 daqiqa", level=2)
    add_meta_table(doc, {
        "Davomiylik": "2-3 daqiqa",
        "Texnika": "Memory Palace + Flashcard",
    })
    doc.add_paragraph()
    
    p = doc.add_paragraph(
        "🧠 Xotira saroyi — Koordinata bog'ini tasavvur qiling:\n\n"
        "1. Darvoza — (0; 0) nuqta (koordinata boshi)\n"
        "2. Yo'lak — y = kx (to'g'ri chiziq darvozadan o'tadi)\n"
        "3. Ko'prik — y = kx + b (b ko'prikning balandligi)\n"
        "4. Tebranma — k > 0 (tepaga), k < 0 (pastga)\n"
        "5. Favvora — b = Oy o'q bilan kesishgan nuqta\n\n"
        "Flashcard:\n"
        "• y = kx + b nima? → Chiziqli funksiya\n"
        "• k nima bildiradi? → Egrilik (Ox o'qqa nisbatan og'ish)\n"
        "• b nima bildiradi? → Oy o'q bilan kesishgan nuqta\n"
        "• Ikki nuqtadan nechta to'g'ri chiziq o'tadi? → Faqat 1 ta")
    p.runs[0].font.size = Pt(10)

    # PHASE 6: FINAL BOSS
    add_styled_heading(doc, "👹 6-FAZA: YAKUNIY BOSS — 5-10 daqiqa", level=2)
    add_meta_table(doc, {
        "Boss HP": "80",
        "Savollar": "5 ta",
        "Format": "Ochiq javob (MC yo'q)",
    })
    doc.add_paragraph()
    
    boss_qs = [
        ("1", 
         "y = -2x + 5 funksiyasining grafigini yasang. "
         "Kamida 3 ta nuqtani hisoblab ko'rsating.",
         "x=0: y=5 → (0;5)\n"
         "x=1: y=3 → (1;3)\n"
         "x=2: y=1 → (2;1)\n"
         "Nuqtalarni belgilab, to'g'ri chiziq o'tkazish.",
         "15 ball (3 nuqta + chiziq), 10 ball (2 nuqta)",
         "1) x = 0, 1, 2 qiymatlarni almashtiring\n"
         "2) Har bir x uchun y ni hisoblang",
         "UZ-MATH-7-LINFUNC-03", "apply", "2",
         "L1→L2: interpret simple visual representations"),
        
        ("2",
         "y = 3x + b funksiyasi (-2; 1) nuqtadan o'tadi. "
         "b ni toping va funksiya formulasini yozing.",
         "1 = 3(-2) + b\n"
         "1 = -6 + b\n"
         "b = 7\n"
         "y = 3x + 7",
         "20 ball (to'liq yechim), 10 ball (usul to'g'ri)",
         "1) Nuqta koordinatalarini almashtiring\n"
         "2) b uchun tenglama yeching",
         "UZ-MATH-7-LINFUNC-03", "apply", "2",
         "L1→L2: basic ratio understanding"),
        
        ("3",
         "y = 2x - 3 va y = -x + 6 funksiyalarning kesishish "
         "nuqtasini toping.",
         "2x - 3 = -x + 6\n"
         "3x = 9\n"
         "x = 3\n"
         "y = 2(3) - 3 = 3\n"
         "Kesishish nuqtasi: (3; 3)",
         "15 ball (to'liq), 8 ball (x topilgan, y xato)",
         "1) Ikkala funksiyani tenglashtiring\n"
         "2) x ni toping, keyin y ni hisoblang",
         "UZ-MATH-7-LINFUNC-04", "analyze", "3",
         "L2→L3: sequential decision-making"),
        
        ("4",
         "Grafik (0; -1) va (2; 3) nuqtalardan o'tadi. "
         "Funksiya formulasini toping.",
         "b = -1 (Oy o'q bilan kesishgan nuqta)\n"
         "k = (3-(-1))/(2-0) = 4/2 = 2\n"
         "y = 2x - 1",
         "15 ball (to'liq), 8 ball (k topilgan, b xato)",
         "1) b ni Oy o'q bilan kesishgan nuqtadan toping\n"
         "2) k = vertikal/gorizontal masofa",
         "UZ-MATH-7-LINFUNC-03", "analyze", "3",
         "L2→L3: integrate multiple representations"),
        
        ("5",
         "Taksi narxi: y = 1500x + 4000. \n"
         "a) 8 km narxini hisoblang.\n"
         "b) 30,000 so'mga necha km borish mumkin?\n"
         "c) Grafikni qaysi choraklarda joylashganini tushuntiring.",
         "a) y = 1500(8) + 4000 = 16,000 so'm\n"
         "b) 30000 = 1500x + 4000 → x = 17.33 ≈ 17 km\n"
         "c) k = 1500 > 0, b = 4000 > 0 → I, II, III choraklar "
         "(chiziq I va III da boshlanib, II ga o'tadi)",
         "20 ball (3 qism), 13 ball (2 qism), 7 ball (1 qism)",
         "1) x = 8 almashtiring\n"
         "2) 30000 ni y o'rniga qo'ying\n"
         "3) k va b belgisini tekshiring",
         "UZ-MATH-7-LINFUNC-03", "evaluate", "4",
         "L3→L4: explicit models; construct explanations"),
    ]
    
    for num, q, ans, rubric, socratic, std, bloom, pisa, transition in boss_qs:
        p = doc.add_paragraph()
        run = p.add_run(f"⚔️ BOSS SAVOL {num}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string("C0392B")
        p = doc.add_paragraph(q)
        p.runs[0].font.size = Pt(10)
        p = doc.add_paragraph()
        run = p.add_run(f"Sokrat ko'rsatmalari:\n{socratic}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("7F8C8D")
        run.italic = True
        p = doc.add_paragraph()
        run = p.add_run(f"✅ Rubrika: {rubric}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("27AE60")
        run.bold = True
        p = doc.add_paragraph()
        run = p.add_run(f"[{std}] | Bloom: {bloom} | PISA L{pisa} | {transition}")
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor.from_string("95A5A6")
        p = doc.add_paragraph(f"💰 Damage: 20 HP (to'liq), 10 HP (qisman), 0 HP (xato)")
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.color.rgb = RGBColor.from_string("E67E22")
        doc.add_paragraph()

    doc.add_page_break()

    # PHASE 7: REFLECTION + REMEDIATION
    add_styled_heading(doc, "🪞 7-FAZA: AKS ETTIRISH VA TUZATISH — 3-5 daqiqa", level=2)
    add_meta_table(doc, {
        "Davomiylik": "3-5 daqiqa",
        "Komponentlar": "Zaif nuqtalar → Mikro-mashqlar → Refleksiya → Spaced Repetition",
    })
    doc.add_paragraph()
    
    p = doc.add_paragraph(
        "📊 BUGUNGI NATIJALAR:\n\n"
        "✅ y = kx funksiyasi — Zo'r!\n"
        "✅ Grafik yasash — Yaxshi\n"
        "⚠️ Kesishish nuqtasini topish — Mashq kerak\n"
        "⚠️ Grafikdan formula topish — Mashq kerak\n\n"
        "🔧 Mikro-mashqlar:\n"
        "1. y = 2x + 1 da x = 3 → y = ___ (j: 7)\n"
        "2. y = kx grafigi (2; 6) dan o'tsa, k = ___ (j: 3)\n"
        "3. y = -x + 3 Ox o'qini qaysi nuqtada kesadi? (j: (3; 0))\n\n"
        "💭 Refleksiya: \"Chiziqli funksiya kundalik hayotda qayerda "
        "uchraydi? Bugun nima tushundim?\"\n\n"
        "🔄 Spaced Repetition:\n"
        "• y = kx (kuchli) → 7 kun\n"
        "• Grafik yasash (barqaror) → 3 kun\n"
        "• Kesishish (zaif) → ERTAGA\n\n"
        "🎉 SESSIYA TUGADI!\n"
        "XP: 2,450 | ⭐⭐ 2 yulduz | 🔥 13 kunlik seriya\n"
        "Keyingi dars: 5-Bob, 4-MAVZU — Loyiha ishi")
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor.from_string("2C3E50")
    
    path = os.path.join(OUTPUT_DIR, "NETS_Demo_Algebra_7_Chiziqli_Funksiya.docx")
    doc.save(path)
    print(f"✅ Algebra saved: {path}")


def generate_biologiya_lesson():
    """
    Biologiya Grade 7, Section 33: "O'simliklarning nafas olishini o'rganish"
    Section 7.3 + 7.4 | Pages 99-100
    """
    doc = Document()
    set_default_font(doc)
    
    # TITLE
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NETS HOMEWORK ENGINE — DEMO LESSON")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string("1B3A4B")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Biologiya · 7-sinf")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string("2C3E50")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("7.3: O'simliklarning nafas olishini o'rganish")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string("34495E")
    run.italic = True
    
    add_meta_table(doc, {
        "Dars": "O'simliklarda nafas olish va fotosintez — taqqoslash",
        "Sinf": "7",
        "Fan": "Biologiya",
        "Darslik": "Biologiya 7-sinf (2022 nashr), sah. 99-100",
        "Standart kod": "UZ-BIO-7-PLANTRESP-03",
        "PISA maqsadi": "Daraja 2 → 3 (Science: interpret data, experimental design)",
        "Bloom darajalari": "Understand, Apply, Analyze, Evaluate",
        "O'tish ko'nikmasi": "L2→L3: interpret experimental data; integrate multiple sources",
        "Rejim": "Standard (20-30 daqiqa)",
        "Boss HP": "80 (Grades 5-8)",
    })
    doc.add_paragraph()

    # PHASE 1: MEMORY SPRINT
    add_styled_heading(doc, "⚡ 1-FAZA: XOTIRA HURRASI — 2 daqiqa", level=2)
    add_meta_table(doc, {
        "Davomiylik": "60 soniya",
        "Savollar": "5 ta (oldingi boblardan)",
        "Manba": "1-6 boblar: Hujayra, Organlar, Fotosintez, Nafas olish",
        "XP": "+100 XP to'g'ri javob uchun",
    })
    doc.add_paragraph()
    
    questions_p1 = [
        ("1", "O'simlik hujayrasining asosiy qismlarini nomlang.",
         "Hujayra devori, sitoplazma, yadro, xloroplast, vakuola",
         "UZ-BIO-7-CELL-01", "remember", "1"),
        ("2", "Fotosintez uchun nima kerak?",
         "Quyosh nuri, CO₂, suv, xlorofill",
         "UZ-BIO-7-PHOTO-01", "understand", "1"),
        ("3", "O'simlikning vegetativ organlarini sanang.",
         "Ildiz, poyasi, barg",
         "UZ-BIO-7-ORGAN-01", "remember", "1"),
        ("4", "Nafas olishda qaysi gaz iste'mol qilinadi?",
         "Kislorod (O₂)",
         "UZ-BIO-7-RESP-01", "remember", "1"),
        ("5", "Fotosintez natijasida qanday gaz chiqariladi?",
         "Kislorod (O₂)",
         "UZ-BIO-7-PHOTO-02", "understand", "1"),
    ]
    
    for num, q, ans, std, bloom, pisa in questions_p1:
        add_question_block(doc, num, q, ans, {"standard_ref": std},
                          bloom=bloom, pisa=f"Level {pisa}")

    # Phase 1 gate (Biologiya-specific)
    p = doc.add_paragraph()
    run = p.add_run("⚙️ Moslashtirish: <60% → tuzatish, 60-79% → davom et, ≥80% → tayyor.")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("95A5A6")
    run.italic = True
    doc.add_page_break()

    # PHASE 2: STORY MODE
    add_styled_heading(doc, "📖 2-FAZA: HIKOYA REJIMI — 5-7 daqiqa", level=2)
    add_meta_table(doc, {
        "Davomiylik": "5-7 daqiqa",
        "Segmentlar": "3 ta",
        "Manba": "Biologiya 7-sinf, sah. 99-100",
        "CPA bosqichi": "Concrete (sham tajribasi) → Pictorial (diagramma) → Abstract (jarayon formulasi)",
    })
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("📕 Segment 1: Sham tajribasi")
    run.bold = True
    run.font.size = Pt(11)
    
    p = doc.add_paragraph(
        "Uchta shisha idish, uchta sham. Har birini alohida joyga qo'yamiz:\n\n"
        "🔬 Tajriba 1: Faqat sham + shisha idish\n"
        "Natija: Sham o'chadi. Sabab: Kislorod tugadi.\n\n"
        "🔬 Tajriba 2: O'simlik + sham + yorug'lik\n"
        "Natija: Sham yonishda davom etadi! Sabab: Fotosintez O₂ ishlab chiqaradi.\n\n"
        "🔬 Tajriba 3: O'simlik + sham + qorong'ilik\n"
        "Natija: Sham o'chadi. Sabab: Fotosintez yo'q, o'simlik nafas oladi.")
    p.paragraph_format.left_indent = Cm(0.5)
    
    p = doc.add_paragraph()
    run = p.add_run("🔒 Checkpoint 1: 2-tajribada sham nega o'chmaydi?")
    run.bold = True
    p = doc.add_paragraph("✅ Javob: Chunki fotosintez natijasida o'simlik kislorod (O₂) "
                          "ishlab chiqaradi va sham yonishda davom etadi.")
    
    p = doc.add_paragraph()
    run = p.add_run("📕 Segment 2: Nafas olishning mohiyati")
    run.bold = True
    run.font.size = Pt(11)
    
    p = doc.add_paragraph(
        "Barcha organizmlar singari o'simliklar ham kislorod bilan nafas oladi "
        "va karbonat angidrid (CO₂) chiqaradi. Lekin o'simliklarning maxsus "
        "nafas olish organlari yo'q — ular barcha hujayralari orqali nafas oladi.\n\n"
        "Qabul qilingan kislorod o'simlik tanasidagi organik moddalarni "
        "anorganik moddalargacha parchalaydi.\n\n"
        "Omborlarda saqlanadigan urug'lar, ildiz va ildizmevalar ham nafas oladi. "
        "Shuning uchun omborlarni shamollatish kerak!")
    p.paragraph_format.left_indent = Cm(0.5)
    
    p = doc.add_paragraph()
    run = p.add_run("🔒 Checkpoint 2: O'simliklar qaysi organ orqali nafas oladi?")
    run.bold = True
    p = doc.add_paragraph("✅ Javob: Maxsus organ yo'q — barcha hujayralari orqali nafas oladi.")
    
    p = doc.add_paragraph()
    run = p.add_run("📕 Segment 3: Fotosintez vs Nafas olish")
    run.bold = True
    run.font.size = Pt(11)
    
    p = doc.add_paragraph(
        "📊 Taqqoslash:\n\n"
        "| Jihat          | Fotosintez           | Nafas olish            |\n"
        "|----------------|---------------------|-----------------------|\n"
        "| Mohiyati       | CO₂ + H₂O → glyukoza + O₂ | Glyukoza + O₂ → CO₂ + H₂O |\n"
        "| Kislorod       | Chiqariladi         | Iste'mol qilinadi     |\n"
        "| CO₂            | Iste'mol qilinadi   | Chiqariladi           |\n"
        "| Energiya       | Quyosh energiyasi → kimyoviy | Kimyoviy → ATF |\n"
        "| Vaqti          | Faqat kunduzi       | Kechayu kunduz        |")
    p.paragraph_format.left_indent = Cm(0.5)
    
    p = doc.add_paragraph()
    run = p.add_run("🔒 Checkpoint 3: Fotosintez kechasi sodir bo'ladimi? Nima uchun?")
    run.bold = True
    p = doc.add_paragraph("✅ Javob: Yo'q. Fotosintez uchun quyosh nuri kerak. "
                          "Kechasi faqat nafas olish sodir bo'ladi.")

    # PHASE 3: GAME BREAKS
    add_styled_heading(doc, "🎮 3-FAZA: O'YIN TANAFUSLARI — 6-9 daqiqa", level=2)
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("🧩 Tile Match: Jarayonni juftlang")
    run.bold = True
    p = doc.add_paragraph(
        "1. Fotosintez       →  A. Kislorod chiqaradi\n"
        "2. Nafas olish      →  B. Faqat kunduzi\n"
        "3. Fotosintez       →  C. CO₂ iste'mol qiladi\n"
        "4. Nafas olish      →  D. Kechayu kunduz\n\n"
        "✅ Javob: 1-B, 2-D, 3-A, 4-C (qo'shimcha: 1-A,C ham to'g'ri)")
    
    p = doc.add_paragraph()
    run = p.add_run("✏️ Sentence Fill:")
    run.bold = True
    p = doc.add_paragraph(
        "1. O'simliklar barcha ___________ orqali nafas oladi. (j: hujayralari)\n"
        "2. Fotosintez uchun ___________ nuri kerak. (j: quyosh)\n"
        "3. Nafas olishda glyukoza ___________ moddalargacha parchalanadi. (j: anorganik)\n"
        "4. Ombordagi urug'lar ham ___________ oladi. (j: nafas)")
    
    p = doc.add_paragraph()
    run = p.add_run("🎁 Mystery Box: Nima sodir bo'ldi?")
    run.bold = True
    p = doc.add_paragraph(
        "Dehqon kartoshkani yopiq omborga joyladi. Bir hafta o'tgach, "
        "omborda havo og'irlashgan, kartoshka chiriy boshlagan.\n\n"
        "Nima sodir bo'ldi? Qaysi jarayon sababchi?\n\n"
        "💡 Ko'rsatma: O'simliklarning qaysi jarayoni davom etadi?\n"
        "✅ Javob: Kartoshka (ildizmeva) nafas olishda davom etgan — "
        "O₂ iste'mol qilib, CO₂ chiqargan. Yopiq omborda shamollatish "
        "bo'lmagani uchun CO₂ to'planib, chirish boshlangan.")
    
    doc.add_page_break()

    # PHASE 4: REAL-LIFE CHALLENGE
    add_styled_heading(doc, "🌍 4-FAZA: HAQIY HAYOT MUAMMOSI — 3-5 daqiqa", level=2)
    add_meta_table(doc, {
        "Davomiylik": "3-5 daqiqa",
        "Bloom": "Analyze, Evaluate",
        "PISA": "3",
        "Kontekst": "Ombor boshqaruvi, o'simlik parvarishi",
    })
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("🌾 Vaziyat: Ombor muammosi")
    run.bold = True
    
    p = doc.add_paragraph(
        "Fermer Xolmat aka 10 tonna bug'doy urug'ini katta yopiq omborga "
        "joyladi. Eshik va derazalarni mahkam yopdi. 2 hafta o'tgach, "
        "omborda namlik ko'paygan, urug'lar qizigan, ba'zilari "
        "chiriy boshlagan.\n\n"
        "1. Nima uchun omborda namlik ko'paydi? Qaysi biologik jarayon sababchi?\n"
        "2. Urug'lar nega qizidi? Energiya qayerdan keldi?\n"
        "3. Omborni to'g'ri boshqarish uchun nima qilish kerak edi?\n"
        "4. Fotosintez va nafas olish bir-biriga qanday bog'liq? Nima uchun "
        "ikkala jarayon ham hayot uchun muhim?\n\n"
        "✅ Javob kaliti:\n"
        "1. Nafas olish: urug'lar O₂ iste'mol qilib, CO₂ va H₂O (suv bug'i) "
        "chiqaradi. Suv bug'i omborda to'planadi.\n"
        "2. Nafas olish ekzotermik jarayon — kimyoviy energiya issiqlikka "
        "aylanadi. Bu issiqlik yopiq omborda to'planadi.\n"
        "3. Shamollatish tizimi kerak. Harorat va namlikni monitoring qilish. "
        "Urug'larni quritib joylash.\n"
        "4. Fotosintez → O₂ + glyukoza ishlab chiqaradi. "
        "Nafas olish → O₂ + glyukozadan foydalanadi. "
        "Bir-birini to'ldiradi — hayot aylanmasi.")
    p.runs[0].font.size = Pt(10)

    # PHASE 5: CONSOLIDATION
    add_styled_heading(doc, "🏰 5-FAZA: XOTIRA SAROYI — 2-3 daqiqa", level=2)
    doc.add_paragraph()
    
    p = doc.add_paragraph(
        "🧠 Xotira saroyi — O'simliklar bog'ini tasavvur qiling:\n\n"
        "1. Kirish eshigi — Barg (barcha hujayralar nafas oladi)\n"
        "2. Chap yo'l — Kunduzi: Fotosintez ISHLAYDI → O₂ chiqadi\n"
        "3. O'ng yo'l — Kechasi: Faqat Nafas olish → CO₂ chiqadi\n"
        "4. Favvora — Energiya aylanishi: Quyosh → Glikoza → ATF\n"
        "5. Ko'prik — Ombor: Shamollatish kerak, aks holda namlik + issiqlik\n\n"
        "Flashcard:\n"
        "• Nafas olish = O₂ iste'mol, CO₂ chiqarish (kechayu kunduz)\n"
        "• Fotosintez = CO₂ iste'mol, O₂ chiqarish (faqat kunduzi)\n"
        "• O'simlikda maxsus nafas organi yo'q — barcha hujayralar orqali\n"
        "• Ikkala jarayon bir-birini to'ldiradi")
    p.runs[0].font.size = Pt(10)

    # PHASE 6: FINAL BOSS
    add_styled_heading(doc, "👹 6-FAZA: YAKUNIY BOSS — 5-10 daqiqa", level=2)
    add_meta_table(doc, {
        "Boss HP": "80",
        "Savollar": "5 ta",
        "Format": "Ochiq javob (MC yo'q)",
    })
    doc.add_paragraph()
    
    boss_qs = [
        ("1",
         "Sham tajribasini tushuntiring: 3 ta shisha idishda "
         "(faqat sham, o'simlik+sham+yorug'lik, o'simlik+sham+qorong'i) "
         "har bir holatda sham qachon o'chadi? Nima uchun?",
         "1. Faqat sham → O₂ tugagach o'chadi\n"
         "2. O'simlik+yorug'lik+sham → Fotosintez O₂ ishlab chiqaradi, "
         "sham o'chmaydi\n"
         "3. O'simlik+qorong'i+sham → Fotosintez yo'q, nafas olish O₂ "
         "iste'mol qiladi → sham tezroq o'chadi",
         "15 ball (3 holat to'liq), 10 ball (2 holat), 5 ball (1 holat)",
         "1. Har bir tajribada qaysi jarayonlar ishlaydi?\n"
         "2. O₂ balansi qanday?\n"
         "3. Fotosintez qachon to'xtaydi?",
         "UZ-BIO-7-PLANTRESP-03", "apply", "3",
         "L2→L3: interpret experimental data"),
        
        ("2",
         "Fotosintez va nafas olishni taqqoslang: "
         "kamida 4 ta jihat bo'yicha farqini yozing.",
         "1. Vaqti: Fotosintez = kunduzi, Nafas = kechayu kunduz\n"
         "2. Gaz: Fotosintez = O₂ chiqaradi, Nafas = O₂ iste'mol\n"
         "3. Energiya: Fotosintez = energiya to'playdi, Nafas = energiya chiqaradi\n"
         "4. Moddalar: Fotosintez = CO₂+H₂O→glyukoza, "
         "Nafas = glyukoza→CO₂+H₂O",
         "15 ball (4 jihat to'liq), 10 ball (3 ta), 5 ball (2 ta)",
         "1. Har bir jarayon qachon sodir bo'ladi?\n"
         "2. Qaysi gaz iste'mol/chiqariladi?\n"
         "3. Energiya qanday o'zgaradi?",
         "UZ-BIO-7-PLANTRESP-04", "analyze", "3",
         "L2→L3: integrate multiple sources"),
        
        ("3",
         "Nima uchun omborda saqlanadigan urug'lar qiziydi? "
         "Bu jarayonni batafsil tushuntiring.",
         "Urug'lar tirik organizm sifatida nafas oladi. Nafas olishda "
         "glyukoza parchalanadi va kimyoviy energiya ajraladi. Bu energiyaning "
         "bir qismi issiqlikka aylanadi. Yopiq omborda issiqlik to'planadi → "
         "urug'lar qiziydi. Agar shamollatish bo'lmasa, namlik ham to'planadi "
         "va chirish boshlanadi.",
         "15 ball (to'liq mexanizm), 8 ball (qisman)",
         "1. Nafas olish nima chiqaradi?\n"
         "2. Energiya qanday o'zgaradi?\n"
         "3. Yopiq omborda nima sodir bo'ladi?",
         "UZ-BIO-7-PLANTRESP-03", "analyze", "3",
         "L2→L3: communicate reasoning"),
        
        ("4",
         "O'simlik kechasi faqat nafas oladi, kunduzi ham nafas oladi ham "
         "fotosintez qiladi. Kunning qaysi vaqtida o'simlik atmosferaga "
         "ko'proq kislorod chiqaradi? Nima uchun?",
         "Kunduzi, chunki fotosintez natijasida kislorod chiqariladi. "
         "Kunduzi o'simlik ham nafas oladi (O₂ iste'mol), lekin fotosintez "
         "tezligi nafas olishdan yuqori bo'lgani uchun sof O₂ chiqariladi. "
         "Kechasi faqat nafas olish — O₂ iste'mol, CO₂ chiqarish.",
         "15 ball (to'liq tushuntirish + sabab), 8 ball (qisman)",
         "1. Kunduzi qaysi jarayonlar ishlaydi?\n"
         "2. Fotosintez tezligi nafas olishdan qanday?\n"
         "3. 'Sof' natija nima?",
         "UZ-BIO-7-PLANTRESP-04", "evaluate", "4",
         "L3→L4: construct explanations; compare strategies"),
        
        ("5",
         "Yangi o'simlik turingizni kashf qildingiz. U qorong'i g'orda "
         "o'sadi. U qanday omon qolishi mumkin? Fotosintez qiladimi? "
         "Nafas oladimi? Agar boshqa o'simlik bo'lmasa, bu g'orda tirik qola "
         "oladimi?",
         "1. Barcha tirik organizmlar nafas oladi — bu o'simlik ham "
         "nafas oladi\n"
         "2. Fotosintez uchun yorug'lik kerak — qorong'i g'orda fotosintez "
         "mumkin emas\n"
         "3. Agar boshqa o'simlik bo'lmasa: g'orda O₂ manbai yo'q, bu o'simlik "
         "nafas olish uchun O₂ kerak → O₂ tugagach o'ladi\n"
         "4. Istisno: agar g'orda yorug'lik manbai bo'lsa yoki boshqa "
         "organizmlar (bakteriyalar) O₂ ishlab chiqarsa",
         "20 ball (4 qism to'liq), 13 ball (3 qism), 7 ball (2 qism)",
         "1. Fotosintez shartlari nima?\n"
         "2. Nafas olish shartlari nima?\n"
         "3. O₂ manbai bormi?\n"
         "4. Ekotizim bormi?",
         "UZ-BIO-7-PLANTRESP-04", "evaluate", "4",
         "L3→L4: explicit models; reflect on results"),
    ]
    
    for num, q, ans, rubric, socratic, std, bloom, pisa, transition in boss_qs:
        p = doc.add_paragraph()
        run = p.add_run(f"⚔️ BOSS SAVOL {num}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string("C0392B")
        p = doc.add_paragraph(q)
        p.runs[0].font.size = Pt(10)
        p = doc.add_paragraph()
        run = p.add_run(f"Sokrat ko'rsatmalari:\n{socratic}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("7F8C8D")
        run.italic = True
        p = doc.add_paragraph()
        run = p.add_run(f"✅ Rubrika: {rubric}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("27AE60")
        run.bold = True
        p = doc.add_paragraph()
        run = p.add_run(f"[{std}] | Bloom: {bloom} | PISA L{pisa} | {transition}")
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor.from_string("95A5A6")
        p = doc.add_paragraph(f"💰 Damage: 20 HP (to'liq), 10 HP (qisman), 0 HP (xato)")
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.color.rgb = RGBColor.from_string("E67E22")
        doc.add_paragraph()

    doc.add_page_break()

    # PHASE 7: REFLECTION + REMEDIATION
    add_styled_heading(doc, "🪞 7-FAZA: AKS ETTIRISH VA TUZATISH — 3-5 daqiqa", level=2)
    doc.add_paragraph()
    
    p = doc.add_paragraph(
        "📊 BUGUNGI NATIJALAR:\n\n"
        "✅ Fotosintez jarayoni — Zo'r!\n"
        "✅ Sham tajribasi — Yaxshi\n"
        "⚠️ Fotosintez vs Nafas taqqoslash — Mashq kerak\n"
        "⚠️ Ombor biologiyasi — Mashq kerak\n\n"
        "🔧 Mikro-mashqlar:\n"
        "1. Fotosintez qachon to'xtaydi? (j: qorong'ida/yorug'lik yo'q)\n"
        "2. Nafas olishda qaysi gaz chiqariladi? (j: CO₂)\n"
        "3. O'simlik maxsus nafas organiga ega — to'g'rimi? (j: yo'q)\n\n"
        "💭 Refleksiya: \"Fotosintez va nafas olish bir-biriga qanday "
        "bog'liq? Kundalik hayotda qayerda ko'ramiz?\"\n\n"
        "🔄 Spaced Repetition:\n"
        "• Fotosintez (kuchli) → 7 kun\n"
        "• Sham tajribasi (barqaror) → 3 kun\n"
        "• Taqqoslash (zaif) → ERTAGA\n\n"
        "🎉 SESSIYA TUGADI!\n"
        "XP: 2,450 | ⭐⭐ 2 yulduz | 🔥 13 kunlik seriya\n"
        "Keyingi dars: 7.4 — Fotosintez va nafas olishni taqqoslash (Amaliy)")
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor.from_string("2C3E50")
    
    path = os.path.join(OUTPUT_DIR, "NETS_Demo_Biologiya_7_Nafas_Olish.docx")
    doc.save(path)
    print(f"✅ Biologiya saved: {path}")


# ============================================================
# GENERATE ALL 3
# ============================================================
if __name__ == "__main__":
    print("="*60)
    print("NETS Homework Engine — Generating 3 Demo Lessons")
    print("="*60)
    
    generate_kimyo_lesson()
    generate_algebra_lesson()
    generate_biologiya_lesson()
    
    print("\n" + "="*60)
    print("ALL 3 LESSONS GENERATED SUCCESSFULLY")
    print("="*60)
    print(f"Output directory: {OUTPUT_DIR}")
