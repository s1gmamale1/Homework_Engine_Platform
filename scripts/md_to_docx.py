"""
Convert NETS-Homework-Engine-Blueprint.md to a properly formatted .docx file.
"""
import re
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Read the markdown file
with open("standards/NETS-Homework-Engine-Blueprint.md", "r", encoding="utf-8") as f:
    md_content = f.read()

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Define Styles ──
style = doc.styles

# Normal style
normal = style['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

# Title
title_style = style['Title']
title_style.font.name = 'Calibri'
title_style.font.size = Pt(28)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

# Heading 1
h1 = style['Heading 1']
h1.font.name = 'Calibri'
h1.font.size = Pt(20)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
h1.paragraph_format.space_before = Pt(24)
h1.paragraph_format.space_after = Pt(12)

# Heading 2
h2 = style['Heading 2']
h2.font.name = 'Calibri'
h2.font.size = Pt(16)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x2E, 0x5E, 0x86)
h2.paragraph_format.space_before = Pt(18)
h2.paragraph_format.space_after = Pt(8)

# Heading 3
h3 = style['Heading 3']
h3.font.name = 'Calibri'
h3.font.size = Pt(13)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0x3B, 0x7D, 0xAD)
h3.paragraph_format.space_before = Pt(12)
h3.paragraph_format.space_after = Pt(6)

# Heading 4
h4 = style['Heading 4']
h4.font.name = 'Calibri'
h4.font.size = Pt(12)
h4.font.bold = True
h4.font.color.rgb = RGBColor(0x4A, 0x90, 0xBF)

# Create code block style
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(9)
code_style.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(4)
code_style.paragraph_format.line_spacing = 1.0
code_style.paragraph_format.left_indent = Cm(0.5)

# Quote style
quote_style = doc.styles.add_style('BlockQuote', WD_STYLE_TYPE.PARAGRAPH)
quote_style.font.name = 'Calibri'
quote_style.font.size = Pt(11)
quote_style.font.italic = True
quote_style.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
quote_style.paragraph_format.left_indent = Cm(1)
quote_style.paragraph_format.space_before = Pt(6)
quote_style.paragraph_format.space_after = Pt(6)

# Tag style for [AI: Tier X], [HOOK], [SPEC]
tag_style = doc.styles.add_style('TagStyle', WD_STYLE_TYPE.CHARACTER)
tag_style.font.name = 'Consolas'
tag_style.font.size = Pt(9)
tag_style.font.bold = True
tag_style.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)


def add_formatted_paragraph(doc, text, style_name='Normal'):
    """Add a paragraph with inline bold/italic/code formatting."""
    p = doc.add_paragraph(style=style_name)
    # Parse inline formatting: **bold**, *italic*, `code`, [TAG: ...]
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`|\[AI:.*?\]|\[HOOK.*?\]|\[SPEC.*?\])', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x80, 0x20, 0x20)
        elif re.match(r'\[(AI:|HOOK|SPEC)', part):
            run = p.add_run(part)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)
        else:
            p.add_run(part)
    return p


def set_cell_shading(cell, color):
    """Set background color of a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_table_from_rows(doc, headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header.strip())
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1B3A5C")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            # Handle bold in cells
            parts = re.split(r'(\*\*.*?\*\*)', cell_text.strip())
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(10)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(10)
            # Alternate row colors
            if r_idx % 2 == 0:
                set_cell_shading(cell, "F0F4F8")

    # Set column widths roughly
    doc.add_paragraph()  # spacing after table
    return table


def parse_md_table(lines):
    """Parse markdown table lines into headers and rows."""
    if len(lines) < 2:
        return None, None
    headers = [c.strip() for c in lines[0].split('|') if c.strip()]
    # Skip separator line (lines[1])
    rows = []
    for line in lines[2:]:
        cols = [c.strip() for c in line.split('|') if c.strip()]
        if cols:
            # Pad if needed
            while len(cols) < len(headers):
                cols.append('')
            rows.append(cols[:len(headers)])
    return headers, rows


# ── Parse and Convert ──

lines = md_content.split('\n')
i = 0
in_code_block = False
code_lines = []

while i < len(lines):
    line = lines[i]

    # Code blocks
    if line.strip().startswith('```'):
        if not in_code_block:
            in_code_block = True
            code_lines = []
            i += 1
            continue
        else:
            # End of code block - write accumulated lines
            in_code_block = False
            # Add a subtle background paragraph for code
            for cl in code_lines:
                p = doc.add_paragraph(cl, style='CodeBlock')
            i += 1
            continue

    if in_code_block:
        code_lines.append(line)
        i += 1
        continue

    # Horizontal rules
    if line.strip() == '---':
        # Add a thin line
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run('─' * 80)
        run.font.size = Pt(6)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        i += 1
        continue

    # Title (# heading)
    if line.startswith('# ') and not line.startswith('## '):
        doc.add_paragraph(line[2:].strip(), style='Title')
        i += 1
        continue

    # Heading 1 (##)
    if line.startswith('## ') and not line.startswith('### '):
        text = line[3:].strip()
        doc.add_paragraph(text, style='Heading 1')
        i += 1
        continue

    # Heading 2 (###)
    if line.startswith('### ') and not line.startswith('#### '):
        text = line[4:].strip()
        doc.add_paragraph(text, style='Heading 2')
        i += 1
        continue

    # Heading 3 (####)
    if line.startswith('#### '):
        text = line[5:].strip()
        doc.add_paragraph(text, style='Heading 3')
        i += 1
        continue

    # Tables (detect start of table)
    if '|' in line and line.strip().startswith('|'):
        table_lines = []
        while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
            table_lines.append(lines[i].strip())
            i += 1
        headers, rows = parse_md_table(table_lines)
        if headers and rows:
            add_table_from_rows(doc, headers, rows)
        continue

    # Block quotes
    if line.strip().startswith('> '):
        text = line.strip()[2:]
        # Collect multi-line quotes
        while i + 1 < len(lines) and lines[i + 1].strip().startswith('> '):
            i += 1
            text += ' ' + lines[i].strip()[2:]
        add_formatted_paragraph(doc, text, style_name='BlockQuote')
        i += 1
        continue

    # Bullet points
    if line.strip().startswith('- '):
        text = line.strip()[2:]
        p = add_formatted_paragraph(doc, text)
        p.style = doc.styles['List Bullet']
        i += 1
        continue

    # Numbered lists
    m = re.match(r'^(\d+)\.\s+(.*)', line.strip())
    if m:
        text = m.group(2)
        p = add_formatted_paragraph(doc, text)
        p.style = doc.styles['List Number']
        i += 1
        continue

    # Empty lines
    if line.strip() == '':
        i += 1
        continue

    # Regular paragraph
    add_formatted_paragraph(doc, line.strip())
    i += 1

# ── Save ──
output_path = "standards/NETS-Homework-Engine-Blueprint.docx"
doc.save(output_path)
print(f"Created: {output_path}")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")
