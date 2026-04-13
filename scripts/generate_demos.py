import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt

def add_meta_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2, style='Table Grid')
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    doc.add_paragraph()

def add_phase(doc, num, name, bloom, pisa, duration, content):
    p = doc.add_paragraph()
    run = p.add_run(f'PHASE {num}: {name}')
    run.bold = True
    run.font.size = Pt(13)
    p2 = doc.add_paragraph()
    p2.add_run(f"Bloom's: {bloom} | PISA: {pisa} | Duration: {duration}").italic = True
    for line in content:
        if line.startswith('>>'):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        elif line.startswith('##'):
            p3 = doc.add_paragraph()
            p3.add_run(line[2:].strip()).bold = True
        elif line == '':
            doc.add_paragraph()
        else:
            doc.add_paragraph(line)

BASE = 'C:/Users/DaddysHere/Documents/Sigma_Edu_3000/demo/'

# ============================================================
# 1. KIMYO — YONISH
# ============================================================
doc1 = Document()
doc1.add_heading('NETS Homework Session: Kimyo \u00a733 \u2014 Yonish', 1)

add_meta_table(doc1, [
    ('Fan / Subject', 'Kimyo (Chemistry)'),
    ('Sinf / Grade', '7-sinf'),
    ('Mavzu / Topic', '\u00a733 \u2014 Yonish (Combustion)'),
    ('Darslik sahifalari', '97\u201398'),
    ('Standart kodi', 'UZ-KIM-7-YON-01'),
    ('PISA Target', 'Level 2\u20133 (Science domain)'),
    ('Rejim / Mode', 'Standard (20\u201330 min)'),
    ('AI Tier', 'Tier 1 (pre-generated), Tier 3 (boss evaluation)'),
])

add_phase(doc1, 1, "TEZKOR MOSLASHTIRISH (Speed Match + Flashcard Sprint)",
    "Engage \u2192 Remember", "Below 1 \u2192 Level 1", "2\u20133 min", [
    "## Hook Game (30 sek) \u2014 Tez tasniflash",
    "Moddalarni \"YONADI\" yoki \"YONMAYDI\" kategoriyasiga ajrating.",
    ">> Qog\u02bboz \u2192 YONADI",
    ">> Temir mix \u2192 YONMAYDI",
    ">> Benzin \u2192 YONADI",
    ">> Tosh \u2192 YONMAYDI",
    ">> O\u02bbtin \u2192 YONADI",
    ">> Shisha \u2192 YONMAYDI",
    ">> Metan gazi \u2192 YONADI",
    ">> Suv \u2192 YONMAYDI",
    "XP: 50 per correct | Speed bonus: +10 if <2 sek",
    "",
    "## Dars xulosasi (30 sek flash)",
    ">> Yonish \u2014 moddaning kislorod bilan reaksiyasi (issiqlik + yorug\u02bblik)",
    ">> Sekin oksidlanish \u2014 olovsiz, sekin jarayon (zanglash)",
    ">> Yonish sharti: kislorod + yonish harorati",
    ">> Havoda <15% O\u2082 \u2192 yonish to\u02bbxtaydi",
    "",
    "## Flashcard Sprint (60 sek) \u2014 Oldingi mavzulardan",
    ">> \"Kislorodning valentligi nechaga teng?\" \u2192 2",
    ">> \"Kislorodning kimyoviy belgisi?\" \u2192 O",
    ">> \"4Fe + 3O\u2082 = ?\" \u2192 2Fe\u2082O\u2083",
    ">> \"Oksid nima?\" \u2192 Biri kisloroddan iborat binar birikma",
    ">> \"Katalitik reaksiya nima?\" \u2192 Katalizator ishtirokidagi reaksiya",
    "XP: 100 per correct | Streak bonus: +10 \u00d7 streak",
])

add_phase(doc1, 2, "HIKOYA REJIMI (Story Mode)",
    "Understand", "Level 2", "5\u20137 min", [
    "## Segment 1 \u2014 Lavuazye kashfiyoti (\U0001f4d6 Darslik: sahifa 97)",
    "Sardor kimyo laboratoriyasida o\u02bbtiribdi. Ustoz aytdi: \"XVIII asrgacha odamlar yonish nima ekanini bilishmagan. Lavuazye birinchi bo\u02bblib yonish \u2014 moddaning kislorod bilan reaksiyasi ekanini isbotladi.\" Sardor hayron: \"Demak, olov kimyoviy reaksiya ekanmi?\"",
    "",
    "## Checkpoint 1",
    "Savol: \"Lavuazye yonish haqida nimani isbotladi?\"",
    "Javob: Yonish moddaning kislorod bilan reaksiyasi ekanini isbotladi.",
    "Xato javobda: \"Sahifa 97ga qarang \u2014 Lavuazye havo tarkibini aniqladi.\"",
    "",
    "\U0001f3ae GAME BREAK 1 \u2014 Tile Match (2 min)",
    ">> \"Sekin oksidlanish\" \u2194 \"Temirning zanglanishi\"",
    ">> \"Yonish\" \u2194 \"Issiqlik va yorug\u02bblik ajralishi\"",
    ">> \"Yong\u02bbin\" \u2194 \"Nazoratsiz yonish\"",
    ">> \"O\u02bbz-o\u02bbzidan yonish\" \u2194 \"Oq fosfor\"",
    "",
    "## Segment 2 \u2014 Yonish shartlari (\U0001f4d6 Darslik: sahifa 98)",
    "Sardor pechka yonida turibdi. Bobosi: \"Mo\u02bbri baland bo\u02bblsa, o\u02bbtin yaxshi yonadi.\" Ustoz tushuntirdi: issiq gaz ko\u02bbtariladi, o\u02bbrniga toza havo kiradi \u2014 kislorod bilan ta\u02bbminlaydi.",
    "",
    "## Checkpoint 2",
    "Savol: \"Modda yonishi uchun qanday shartlar kerak?\"",
    "Javob: Kislorod bilan aloqa va yonish haroratigacha qizdirish.",
    "",
    "\U0001f3ae GAME BREAK 2 \u2014 Sentence Fill (2 min)",
    ">> \"Yonish \u2014 moddalarning _____ bilan reaksiyasi natijasida _____ va _____ ajralishi bilan boradigan jarayon.\" \u2192 kislorod, issiqlik, yorug\u02bblik",
    ">> \"Havoda _____% dan kam kislorod bo\u02bblsa, yonish mumkin bo\u02bblmaydi.\" \u2192 15",
    ">> \"Qog\u02bboz _____ \u00b0C ga qadar qizdirilganda yonadi.\" \u2192 230",
    "",
    "## Segment 3 \u2014 Olovni o\u02bbchirish (\U0001f4d6 Darslik: sahifa 98)",
    "Maktabda o\u02bbt o\u02bbchirish mashg\u02bbuloti. Ustoz: suv havoni to\u02bbsadi va sovitadi, lekin kaliy/natriy metallarga suv QUYILMAYDI \u2014 olov kuchayadi!",
    "",
    "## Checkpoint 3",
    "Savol: \"Nima uchun ba\u02bbzi metallarni suv bilan o\u02bbchirib bo\u02bblmaydi?\"",
    "Javob: Kaliy va natriy suv bilan faol reaksiyaga kirishadi, bu olovni kuchaytiradi.",
])

add_phase(doc1, 3, "TAHLIL (Analyze \u2014 Real-Life Challenge)",
    "Analyze \u2192 Evaluate", "Level 3\u20134", "3\u20135 min", [
    "## Senariy: Samarqanddagi non yopish",
    "Nargiza onajoniga tandir yoqishda yordam bermoqda. Pastdan havo kirsa olov kuchayadi; og\u02bbzini yopsa olov so\u02bbnadi.",
    "",
    ">> 1-savol (PISA L2): \"Tandirning og\u02bbzi yopilganda olov nima uchun so\u02bbnadi?\" \u2192 Kislorod kirib kelmaydi.",
    ">> 2-savol (PISA L3): \"Mo\u02bbri baland bo\u02bblsa olov kuchliroq yonadi. Ilmiy tushuntiring.\" \u2192 Issiq gaz ko\u02bbtariladi, toza havo (O\u2082) kiradi.",
    ">> 3-savol (PISA L4): \"O\u02bbrmon yong\u02bbinida suv bilan o\u02bbchirish eng yaxshi usulmi? Muqobil taklif qiling.\" \u2192 Tuproq bilan ko\u02bbmish (havodan ajratish), o\u02bbt yo\u02bblagi yasash.",
])

add_phase(doc1, 4, "XOTIRA SAROYI (Memory Palace)",
    "Remember (consolidation)", "N/A", "2\u20133 min", [
    "Manzil: Registon maydoni, 5 ta joy",
    ">> Bosh darvoza \u2192 \"Yonish = kislorod + issiqlik + yorug\u02bblik\"",
    ">> Chap madrasa \u2192 \"Sekin oksidlanish = olovsiz zanglash\"",
    ">> Hovuz \u2192 \"Yonish sharti: O\u2082 + harorat\"",
    ">> O\u02bbng madrasa \u2192 \"Havoda <15% O\u2082 \u2192 yonish to\u02bbxtaydi\"",
    ">> Minora \u2192 \"Oq fosfor \u2014 o\u02bbz-o\u02bbzidan yonadi, suvda saqlash kerak\"",
])

add_phase(doc1, 5, "BOSS JANGI (Boss Fight)",
    "Apply \u2192 Evaluate", "Level 3\u20136 (tiered)", "5\u201310 min", [
    "Boss HP: 100 | Hint tax: \u221210 HP | Combo: 3 correct \u2192 2\u00d7 damage",
    "",
    "## Easy (\u221210 HP, PISA L3)",
    "Savol: \"4Fe + 3O\u2082 = 2Fe\u2082O\u2083 \u2014 yonishmi yoki sekin oksidlanishmi? Nima uchun?\"",
    "Javob: Sekin oksidlanish. Olov va ko\u02bbp issiqlik hosil bo\u02bblmaydi.",
    "",
    "## Easy (\u221210 HP, PISA L3)",
    "Savol: \"Yonish uchun qanday shartlar zarur? Kamida ikkitasini ayting.\"",
    "Javob: 1) Kislorod mavjudligi. 2) Yonish haroratigacha qizdirish.",
    "",
    "## Medium (\u221220 HP, PISA L4)",
    "Savol: \"Temirchi ko\u02bbrasi nega puflash bilan ishlaydi? Kimyoviy tushuntiring.\"",
    "Javob: Puflash ko\u02bbproq kislorod beradi, yonish tezligi va harorat ortadi.",
    "",
    "## Medium (\u221220 HP, PISA L4)",
    "Savol: \"Oq fosfor suvda, qizil fosfor oddiy sharoitda saqlanadi. Farqni tushuntiring.\"",
    "Javob: Oq fosforning yonish harorati xona haroratiga yaqin (o\u02bbz-o\u02bbzidan yonadi).",
    "",
    "## Hard (\u221230 HP, PISA L5)",
    "Savol: \"Kosmik kemada yong\u02bbin chiqsa, suv bilan o\u02bbchirish mumkinmi? Muqobil usul taklif qiling.\"",
    "Javob: Og\u02bbirliksizlikda suv tarqaladi. CO\u2082 o\u02bbchirish apparati yaxshiroq \u2014 kislorodni siqib chiqaradi.",
])

add_phase(doc1, 6, "TAKRORLASH VA FIKRLASH (Remediation + Reflection)",
    "Meta", "N/A", "3\u20135 min", [
    "## Zaif tomonlar (AI tahlili asosida)",
    ">> Agar \"yonish shartlari\"da xato \u2192 \"To\u02bbldiring: Yonish uchun _____ va _____ kerak\" \u2192 kislorod, harorat",
    ">> Agar \"sekin oksidlanish vs yonish\"da xato \u2192 \"Tasniflang: zanglash / sham yonishi / temir xiralashishi\"",
    "",
    "## Reflection Prompts",
    ">> Accuracy \u226580%: \"Bugungi darsda eng qiziqarli nima bo\u02bbldi?\"",
    ">> Accuracy 60\u201379%: \"Qaysi tushuncha eng qiyin bo\u02bbldi?\"",
    ">> Accuracy <60%: \"Keyingi safar nimani boshqacha qilgan bo\u02bblardingiz?\"",
])

doc1.save(BASE + 'Grade7-Kimyo-S33-Yonish.docx')
print("\u2713 Kimyo done")


# ============================================================
# 2. ALGEBRA — CHIZIQLI FUNKSIYA
# ============================================================
doc2 = Document()
doc2.add_heading('NETS Homework Session: Algebra \u00a733 \u2014 Chiziqli funksiya', 1)

add_meta_table(doc2, [
    ('Fan / Subject', 'Algebra (Mathematics)'),
    ('Sinf / Grade', '7-sinf'),
    ('Mavzu / Topic', '\u00a733 \u2014 Chiziqli funksiya (Linear Function)'),
    ('Darslik sahifalari', '120\u2013128'),
    ('Standart kodi', 'UZ-MATH-7-FUNC-03'),
    ('PISA Target', 'Level 2\u20133 (Mathematics domain)'),
    ('Rejim / Mode', 'Standard (20\u201330 min)'),
    ('AI Tier', 'Tier 1 (pre-generated), Tier 3 (boss evaluation)'),
])

add_phase(doc2, 1, "TEZKOR MOSLASHTIRISH (Speed Match + Flashcard Sprint)",
    "Engage \u2192 Remember", "Below 1 \u2192 Level 1", "2\u20133 min", [
    "## Hook Game (30 sek) \u2014 Koordinata aniqlash",
    "Nuqtalar paydo bo\u02bbladi. Qaysi chorakda joylashganini tanlang.",
    ">> (3, 2) \u2192 I chorak",
    ">> (\u22121, 4) \u2192 II chorak",
    ">> (\u22122, \u22125) \u2192 III chorak",
    ">> (4, \u22121) \u2192 IV chorak",
    ">> (0, 3) \u2192 Oy o\u02bbqi",
    ">> (\u22122, 0) \u2192 Ox o\u02bbqi",
    "",
    "## Flashcard Sprint (60 sek)",
    ">> \"Funksiya nima?\" \u2192 Har bir elementni boshqa to\u02bbplamning bitta elementi bilan bog\u02bblaydi",
    ">> \"f(x) = 2x da, f(3) = ?\" \u2192 6",
    ">> \"y = kx da k < 0 bo\u02bblsa, qaysi choraklarda?\" \u2192 II va IV",
    ">> \"(2; 4) nuqtaning absissasi?\" \u2192 2",
    ">> \"Dekart sistemasida nechta o\u02bbq bor?\" \u2192 2 (Ox va Oy)",
])

add_phase(doc2, 2, "HIKOYA REJIMI (Story Mode)",
    "Understand", "Level 2", "5\u20137 min", [
    "## Segment 1 \u2014 y = kx dan y = kx + b ga (\U0001f4d6 sahifa 120\u2013122)",
    "Malika maktab hovuzining perimetrini hisoblamoqda. P = 4a formula yordamida ishladi. \"Bu y = 4x funksiya ekan!\" Keyin ustoz: \"Agar 2 metr yo\u02bblak bo\u02bblsa?\" Malika: P = 4a + 8. \"Mana, y = kx + b \u2014 chiziqli funksiya!\"",
    "",
    "## Checkpoint 1",
    "Savol: \"y = kx + b formulada b nimani bildiradi?\"",
    "Javob: b \u2014 Oy o\u02bbqi bilan to\u02bbg\u02bbri chiziqning kesishish nuqtasi.",
    "",
    "\U0001f3ae GAME BREAK 1 \u2014 Mystery Box (2\u20133 min)",
    ">> Box 1: \"y = 3x + 1 da k va b ni toping\" \u2192 k = 3, b = 1",
    ">> Box 2: \"y = \u22122x + 5 grafigi qaysi choraklarda?\" \u2192 I, II, IV (k<0, b>0)",
    ">> Box 3: \"(0; 4) nuqtadan o\u02bbtuvchi y = kx + b da b = ?\" \u2192 4",
    "",
    "## Segment 2 \u2014 Grafik yasash (\U0001f4d6 sahifa 124)",
    "Ustoz: \"Chiziqli funksiya grafigini yasash uchun faqat 2 ta nuqta yetarli!\" Malika y = \u22122x + 1 ni yasadi: x=0 \u2192 y=1, x=1 \u2192 y=\u22121. Ikki nuqtani chizg\u02bbich bilan tutashtirdi.",
    "",
    "## Checkpoint 2",
    "Savol: \"y = \u22122x + 1 da x = 0 bo\u02bblganda y nechaga teng?\"",
    "Javob: y = \u22122(0) + 1 = 1",
    "",
    "\U0001f3ae GAME BREAK 2 \u2014 Sentence Fill (2 min)",
    ">> \"y = kx + b ko\u02bbrinishidagi funksiya _____ funksiya deyiladi.\" \u2192 chiziqli",
    ">> \"k > 0 bo\u02bblsa grafik _____ va _____ choraklarda joylashadi.\" \u2192 I, III",
    ">> \"Grafik yasash uchun _____ ta nuqta yetarli.\" \u2192 2",
    ">> \"b son Oy o\u02bbqi bilan chiziqning _____ nuqtasi.\" \u2192 kesishish",
])

add_phase(doc2, 3, "TAHLIL (Analyze \u2014 Real-Life Challenge)",
    "Analyze \u2192 Evaluate", "Level 3\u20134", "3\u20135 min", [
    "## Senariy: Toshkentdagi taksi",
    "Jasur taksiga o\u02bbtiribdi. Narxi: 5 000 so\u02bbm boshlang\u02bbich + 3 000 so\u02bbm/km. Cho\u02bbntakda 35 000 so\u02bbm.",
    "",
    ">> 1-savol (PISA L2): \"Taksi narxi formulasini tuzing.\" \u2192 y = 3000x + 5000",
    ">> 2-savol (PISA L3): \"35 000 so\u02bbm bilan maksimum necha km?\" \u2192 35000 = 3000x + 5000; x = 10 km",
    ">> 3-savol (PISA L4): \"Boshqa taksi: 8000 + 2000x. 12 km uchun qaysi arzon?\" \u2192 Taksi 1: 41000; Taksi 2: 32000. Taksi 2 arzon.",
])

add_phase(doc2, 4, "XOTIRA SAROYI (Memory Palace)",
    "Remember (consolidation)", "N/A", "2\u20133 min", [
    "Manzil: Registon maydoni, 5 ta joy",
    ">> Bosh darvoza \u2192 \"y = kx + b \u2014 chiziqli funksiya\"",
    ">> Chap madrasa \u2192 \"k \u2014 to\u02bbg\u02bbri chiziq qiyaligi (og\u02bbishi)\"",
    ">> Hovuz \u2192 \"b \u2014 Oy o\u02bbqi bilan kesishish nuqtasi\"",
    ">> O\u02bbng madrasa \u2192 \"Grafik yasash uchun 2 nuqta yetarli\"",
    ">> Minora \u2192 \"k > 0 \u2192 o\u02bbsuvchi; k < 0 \u2192 kamayuvchi\"",
])

add_phase(doc2, 5, "BOSS JANGI (Boss Fight)",
    "Apply \u2192 Evaluate", "Level 3\u20136 (tiered)", "5\u201310 min", [
    "Boss HP: 100 | Hint tax: \u221210 HP | Combo: 3 correct \u2192 2\u00d7 damage | No MC",
    "",
    "## Easy (\u221210 HP, PISA L3)",
    "Savol: \"y = 4x \u2212 3 ning Oy kesishish nuqtasi?\"",
    "Javob: x = 0, y = \u22123. Nuqta: (0; \u22123).",
    "",
    "## Easy (\u221210 HP, PISA L3)",
    "Savol: \"y = \u2212x + 6 ning nolini toping (y = 0 da x = ?).\"",
    "Javob: 0 = \u2212x + 6; x = 6. Nol: (6; 0).",
    "",
    "## Medium (\u221220 HP, PISA L4)",
    "Savol: \"y = 2x \u2212 3 va y = \u2212x + 6 qaysi nuqtada kesishadi?\"",
    "Javob: 2x \u2212 3 = \u2212x + 6; 3x = 9; x = 3; y = 3. Nuqta: (3; 3).",
    "",
    "## Medium (\u221220 HP, PISA L4)",
    "Savol: \"A(2; 5) va B(4; 9) dan o\u02bbtuvchi funksiyani toping.\"",
    "Javob: k = (9\u22125)/(4\u22122) = 2. 5 = 2(2)+b; b = 1. y = 2x + 1.",
    "",
    "## Hard (\u221230 HP, PISA L5)",
    "Savol: \"y = 3x + 6 grafigi va Ox, Oy o\u02bbqlari bilan uchburchak yuzini toping.\"",
    "Javob: Oy: (0; 6). Ox: (\u22122; 0). Katetlar: 6 va 2. S = \u00bd \u00d7 6 \u00d7 2 = 6.",
])

add_phase(doc2, 6, "TAKRORLASH VA FIKRLASH (Remediation + Reflection)",
    "Meta", "N/A", "3\u20135 min", [
    "## Zaif tomonlar",
    ">> Agar k va b topishda xato \u2192 \"y = 5x \u2212 2 da k = ?, b = ?\" \u2192 k=5, b=\u22122",
    ">> Agar grafik yasashda xato \u2192 \"y = x + 3 uchun 2 nuqta toping\"",
    "",
    "## Reflection",
    ">> \"Chiziqli funksiyani qayerda ishlatish mumkin? Hayotdan misol keltiring.\"",
])

doc2.save(BASE + 'Grade7-Algebra-S33-Chiziqli-Funksiya.docx')
print("\u2713 Algebra done")


# ============================================================
# 3. BIOLOGIYA — MODDALAR TRANSPORTI
# ============================================================
doc3 = Document()
doc3.add_heading("NETS Homework Session: Biologiya \u00a733 \u2014 O\u02bbsimliklarda moddalar transporti", 1)

add_meta_table(doc3, [
    ('Fan / Subject', 'Biologiya (Biology)'),
    ('Sinf / Grade', '7-sinf'),
    ('Mavzu / Topic', "\u00a733 \u2014 O\u02bbsimliklarda moddalar transporti"),
    ('Darslik sahifalari', '102\u2013105'),
    ('Standart kodi', 'UZ-BIO-7-TRANS-01'),
    ('PISA Target', 'Level 2\u20133 (Science domain)'),
    ('Rejim / Mode', 'Standard (20\u201330 min)'),
    ('AI Tier', 'Tier 1 (pre-generated), Tier 3 (boss evaluation)'),
])

add_phase(doc3, 1, "TEZKOR MOSLASHTIRISH (Speed Match + Flashcard Sprint)",
    "Engage \u2192 Remember", "Below 1 \u2192 Level 1", "2\u20133 min", [
    "## Hook Game (30 sek) \u2014 Tez tasniflash",
    "Tushunchalarni \"O\u02bbSIMLIK\" yoki \"HAYVON\" ga ajrating.",
    ">> Fotosintez \u2192 O\u02bbSIMLIK",
    ">> Qon aylanish \u2192 HAYVON",
    ">> Ildiz tukchasi \u2192 O\u02bbSIMLIK",
    ">> Yurak \u2192 HAYVON",
    ">> Barg og\u02bbizchasi \u2192 O\u02bbSIMLIK",
    ">> Eritrotsit \u2192 HAYVON",
    ">> Poya \u2192 O\u02bbSIMLIK",
    ">> Kapillyar \u2192 HAYVON",
    "",
    "## Flashcard Sprint (60 sek)",
    ">> \"Fotosintez qayerda sodir bo\u02bbladi?\" \u2192 Bargda (xloroplastda)",
    ">> \"O\u02bbsimlikning yer ostki organi?\" \u2192 Ildiz",
    ">> \"To\u02bbqima nima?\" \u2192 Bir xil tuzilish va vazifaga ega hujayralar guruhi",
    ">> \"Nafas olishda qaysi gaz yutiladi?\" \u2192 Kislorod (O\u2082)",
    ">> \"O\u02bbtkazuvchi to\u02bbqima vazifasi?\" \u2192 Moddalarni tashish",
])

add_phase(doc3, 2, "HIKOYA REJIMI (Story Mode)",
    "Understand", "Level 2", "5\u20137 min", [
    "## Segment 1 \u2014 Diffuziya va osmos (\U0001f4d6 sahifa 102)",
    "Dilnoza tajriba ko\u02bbrdi: U shaklidagi idishning o\u02bbrtasi membrana bilan to\u02bbsilgan. Bir tomonga sho\u02bbr suv, boshqasiga toza suv quyildi. Suv sho\u02bbr tomonga o\u02bbtdi! \"Bu osmos!\" dedi ustoz. \"Ildiz ham xuddi shunday suv shimadi.\"",
    "",
    "## Checkpoint 1",
    "Savol: \"Osmos nima?\"",
    "Javob: Suvning membrana orqali konsentratsiyasi past muhitdan yuqori muhitga o\u02bbtishi.",
    "",
    "\U0001f3ae GAME BREAK 1 \u2014 Tile Match (2 min)",
    ">> \"Diffuziya\" \u2194 \"Molekulalar yuqori konsentratsiyadan pastga\"",
    ">> \"Osmos\" \u2194 \"Suv membrana orqali sho\u02bbr tomonga\"",
    ">> \"Ksilema\" \u2194 \"Suv va mineral tuzlar yuqoriga\"",
    ">> \"Floema\" \u2194 \"Organik moddalar pastga\"",
    ">> \"Transpiratsiya\" \u2194 \"Barg orqali suv bug\u02bblatish\"",
    "",
    "## Segment 2 \u2014 Ildiz va transport sistema (\U0001f4d6 sahifa 103\u2013104)",
    "Dilnoza ildizning kesmasi rasmini ko\u02bbrdi. Ustoz: \"Ildiz tukchasi suv shimadi \u2014 osmos bilan. Keyin suv ksilema orqali yuqoriga ko\u02bbtariladi. Bargda tayyorlangan organik moddalar esa floema orqali orqaga qaytadi.\"",
    "",
    "## Checkpoint 2",
    "Savol: \"Ksilema va floema nima tashiydi?\"",
    "Javob: Ksilema \u2014 suv va mineral tuzlar (yuqoriga). Floema \u2014 organik moddalar (barcha organlarga).",
    "",
    "\U0001f3ae GAME BREAK 2 \u2014 Sentence Fill (2 min)",
    ">> \"Ildiz _____ orqali tuproqdan suv shimib oladi.\" \u2192 tukchalar",
    ">> \"O\u02bbsimlik suvining _____% i transpiratsiyaga sarflanadi.\" \u2192 90",
    ">> \"Eritrotsit 0.9% tuz eritmasida normal, bu _____ eritma.\" \u2192 izotonik",
    ">> \"Suv _____ orqali, organik moddalar _____ orqali tashiladi.\" \u2192 ksilema, floema",
])

add_phase(doc3, 3, "TAHLIL (Analyze \u2014 Real-Life Challenge)",
    "Analyze \u2192 Evaluate", "Level 3\u20134", "3\u20135 min", [
    "## Senariy: Nargizaning bog\u02bbi",
    "Nargiza tomatlarga juda ko\u02bbp suv quydi \u2014 ertasi barglar sarg\u02bbaydi. Qo\u02bbshnisi 3 kun sug\u02bbormadi \u2014 ular ham so\u02bbldi.",
    "",
    ">> 1-savol (PISA L2): \"Suv yetishmaganda o\u02bbsimlik nima uchun so\u02bbladi?\" \u2192 Transpiratsiya davom etadi, lekin ildiz suv shima olmaydi, hujayralar tarangligi kamayadi.",
    ">> 2-savol (PISA L3): \"Ortiqcha suv quyilganda ildiz nima uchun shikastlanadi?\" \u2192 Tuproqdagi havo siqib chiqariladi, ildiz nafas ola olmaydi, chiriydi.",
    ">> 3-savol (PISA L4): \"Juda sho\u02bbr suv bilan sug\u02bborilsa nima bo\u02bbladi? Osmos bilan tushuntiring.\" \u2192 Tuz konsentratsiyasi tashqarida yuqori, suv hujayradan TASHQARIGA chiqadi (teskari osmos). O\u02bbsimlik qurib qoladi.",
])

add_phase(doc3, 4, "XOTIRA SAROYI (Memory Palace)",
    "Remember (consolidation)", "N/A", "2\u20133 min", [
    "Manzil: Registon maydoni, 5 ta joy",
    ">> Bosh darvoza \u2192 \"Diffuziya = yuqoridan pastga; Osmos = suv sho\u02bbr tomonga\"",
    ">> Chap madrasa \u2192 \"Ildiz zonalari: bo\u02bblinuvchi \u2192 o\u02bbsuvchi \u2192 so\u02bbruvchi \u2192 o\u02bbtkazuvchi\"",
    ">> Hovuz \u2192 \"Ksilema = suv YUQORIGA; Floema = organik moddalar PASTGA\"",
    ">> O\u02bbng madrasa \u2192 \"Transpiratsiya = 90% suv bug\u02bblatiladi\"",
    ">> Minora \u2192 \"Izotonik = 0.9% NaCl; giper = bujmayadi; gipo = yoriladi\"",
])

add_phase(doc3, 5, "BOSS JANGI (Boss Fight)",
    "Apply \u2192 Evaluate", "Level 3\u20136 (tiered)", "5\u201310 min", [
    "Boss HP: 100 | Hint tax: \u221210 HP | Combo: 3 correct \u2192 2\u00d7 damage | No MC",
    "",
    "## Easy (\u221210 HP, PISA L3)",
    "Savol: \"Diffuziya va osmosning farqini tushuntiring.\"",
    "Javob: Diffuziya \u2014 har qanday molekula yuqoridan pastga. Osmos \u2014 faqat suv, membrana orqali, pastdan yuqoriga.",
    "",
    "## Easy (\u221210 HP, PISA L3)",
    "Savol: \"Ildiz tukchalari qanday vazifani bajaradi?\"",
    "Javob: Tuproqdan suv va mineral tuzlarni osmos orqali shimib oladi. Bitta tukcha = bitta hujayra, 10\u201320 kun yashaydi.",
    "",
    "## Medium (\u221220 HP, PISA L4)",
    "Savol: \"Eritrotsit 2% tuz eritmasiga solinsa nima bo\u02bbladi? Jarayonni nomlang.\"",
    "Javob: 2% gipertonik eritma. Osmos tufayli suv eritrotsitdan tashqariga chiqadi. Eritrotsit bujmayadi.",
    "",
    "## Medium (\u221220 HP, PISA L4)",
    "Savol: \"Daraxt tanasiga halqa shaklida po\u02bbstloq kesilsa nima uchun qurib qoladi?\"",
    "Javob: Floema po\u02bbstloq ostida. Kesilsa organik moddalar ildizga yetib bormaydi, ildiz energiya olmay halok bo\u02bbladi.",
    "",
    "## Hard (\u221230 HP, PISA L5)",
    "Savol: \"Cho\u02bbl o\u02bbsimliklari suvni qanday tejaydi? Kamida 3 moslashuvni tushuntiring.\"",
    "Javob: 1) Barglar tikanga aylangan \u2014 transpiratsiya yuzasi kamayadi. 2) Qalin kutikulya \u2014 suv bug\u02bblanishini to\u02bbsadi. 3) Chuqur ildiz \u2014 tuproqning chuqur qatlamlaridan osmos orqali suv shimadi.",
])

add_phase(doc3, 6, "TAKRORLASH VA FIKRLASH (Remediation + Reflection)",
    "Meta", "N/A", "3\u20135 min", [
    "## Zaif tomonlar",
    ">> Agar ksilema/floema farqida xato \u2192 \"Tasniflang: suv yuqoriga / organik moddalar pastga \u2014 qaysi nay?\"",
    ">> Agar osmosda xato \u2192 \"Suv qaysi tomonga o\u02bbtadi: toza\u2192sho\u02bbr yoki sho\u02bbr\u2192toza?\"",
    "",
    "## Reflection",
    ">> \"O\u02bbsimliklardagi transport sistema inson qon aylanishiga qanday o\u02bbxshaydi?\"",
])

doc3.save(BASE + 'Grade7-Biologiya-S33-Moddalar-Transporti.docx')
print("\u2713 Biologiya done")
print("\nAll 3 demo lessons saved to demo/")
