"""Insert Interactive Game Catalog Integration sub-section into Blueprint.docx
Section 3 (Phase 3 Game Breaks). Anchor: 'SECTION 4: PHASE 4' Heading 1."""
import os
from docx import Document
from docx.oxml import OxmlElement

PATH = os.path.join(
    os.path.dirname(__file__), "..", "standards", "NETS-Homework-Engine-Blueprint.docx"
)

doc = Document(PATH)


def find_paragraph_by_text(text_substring, style_prefix=None):
    for p in doc.paragraphs:
        if style_prefix and not p.style.name.startswith(style_prefix):
            continue
        if text_substring in p.text:
            return p
    return None


def insert_para_before(anchor_p, text, style_name="Normal"):
    new_p = OxmlElement("w:p")
    anchor_p._p.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, anchor_p._parent)
    try:
        para.style = doc.styles[style_name]
    except KeyError:
        pass
    if text:
        para.add_run(text)
    return para


BLOCK = [
    ("Heading 3", "Interactive Game Catalog Integration (NEW 2026-04-07)"),
    ("Normal", "Catalog file: standards/NETS-Interactive-Game-Catalog.md (v1.0, 12 knowledge-gated games derived from classic board/card/puzzle mechanics: Tic Tac Toe, Connect Four, Mastermind, Memory, Reaction Chain, Word Ladder, Sliding Tile, Blackjack, Risk, Escape Room, Bridge Builder, Minesweeper)."),
    ("Normal", "Dual-catalog selection rule: When Phase 3 picks the 3 games for a session, it MUST select AT LEAST 1 game from the Interactive Catalog AND AT LEAST 2 games from the Default 16-mechanic pool (Section 6 of the UNIFIED spec). The Interactive Catalog is a COMPLEMENT to the default pool, never a replacement."),
    ("Normal", "Why two pools: the Default 16 are the core pedagogy mechanics that map 1:1 to Bloom's/PISA bands (Memory Sprint, Spaced Flashcards, Tile Match, Sentence Fill, Memory Palace, Story Mode, Adaptive Quiz, Mystery Box, Movement Breaks, Why Chain, Peer Teaching, Real-Life Challenge, Reflection Journal, Final Boss, Tic Tac Toe vs AI, Notebook Capture). The Interactive Catalog layers classic-game strategy on top of knowledge-gating so practice feels like a duel, not a quiz."),
    ("Heading 3", "Interactive Catalog Games (12)"),
    ("Normal", "1. Tic Tac Toe vs AI \u2014 minimax duel; correct answer places chosen tile, wrong = random placement."),
    ("Normal", "2. Connect Four vs AI \u2014 7x6 drop grid; column choice gated by question."),
    ("Normal", "3. Codebreaker \u2014 Mastermind-style secret code; each guess gated by a correct answer."),
    ("Normal", "4. Memory Match Blitz \u2014 Concentration grid; matched pairs confirmed by a subject question."),
    ("Normal", "5. Reaction Chain \u2014 6-10 node chain; 3 wrong in a row breaks the chain."),
    ("Normal", "6. Word Ladder Climb \u2014 transform start word to target one letter at a time, each step gated."),
    ("Normal", "7. Puzzle Lock (Sliding Tile) \u2014 3x3/4x4 sliding puzzle; wrong answers slide a random adjacent tile."),
    ("Normal", "8. Blackjack 21 Knowledge Edition \u2014 Hit/Stand gated; wrong answer lets the AI make the choice."),
    ("Normal", "9. Territory Conquest \u2014 simplified Risk 1v1 area control; failed attacks lose territory."),
    ("Normal", "10. Escape Room: Subject Lock \u2014 unlock 4 virtual objects, cross-hinting clues, hard time limit."),
    ("Normal", "11. Bridge Builder \u2014 physics bridge construction; wrong answers still deduct budget."),
    ("Normal", "12. Minefield Navigator \u2014 hidden-mine grid; wrong answer forces a move to an AI-chosen adjacent tile."),
    ("Heading 3", "When to pick a Catalog game vs a Default game"),
    ("Normal", "Pick a CATALOG game when: the topic maps cleanly to a spatial/strategic metaphor (chronology \u2192 Puzzle Lock, metabolic pathways \u2192 Reaction Chain, periodic table \u2192 Codebreaker, geography \u2192 Territory Conquest); engagement has dipped over the last 3 sessions; the session is a Big Boss or Mythical Boss cross-subject review (Territory Conquest, Escape Room, or Bridge Builder are configured for this); or the student has already seen the same Default mechanic twice this week."),
    ("Normal", "Pick a DEFAULT game when: the Bloom's target is Remember/Understand and needs maximum repetition density (Spaced Flashcards, Memory Sprint, Tile Match); content is narrative-heavy (Story Mode, Why Chain); the task requires physical hand-work (Notebook Capture); the student is in a Recovery Session or below PISA L2 (Default mechanics have lower cognitive overhead); or the Interactive Catalog subject-compatibility matrix marks the Catalog option as weak/incompatible for the current subject."),
    ("Normal", "Device / context constraints: Bridge Builder and Minefield Navigator carry higher render cost and should be skipped on low-end devices; Blackjack 21 is math-only; Word Ladder Climb is language/biology only. If no Catalog game passes subject/device validation, the engine logs a CATALOG_FALLBACK warning and selects 3 from Default; repeated fallbacks trigger a content-ops review."),
]


section4 = find_paragraph_by_text("SECTION 4: PHASE 4", style_prefix="Heading 1")
if section4 is None:
    print("ERROR: Section 4 anchor not found")
else:
    for style, text in BLOCK:
        insert_para_before(section4, text, style_name=style)
    print(f"Inserted {len(BLOCK)} paragraphs (Interactive Game Catalog) before SECTION 4")

doc.save(PATH)
size = os.path.getsize(PATH)
print(f"Saved: {PATH} ({size} bytes)")
