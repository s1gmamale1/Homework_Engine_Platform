"""
Build 4 deliverables:
1. Grade 5 History §1 Homework Design (docx) — full Blueprint-compliant session
2. Student A: High-performing "Malika" (JSON)
3. Student B: Struggling "Bobur" (JSON)
4. Student C: Brand new "Jasur" (JSON)
"""
import sys, io, json, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ═══════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.text = ''
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1B3A5C")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            p = cell.paragraphs[0]
            p.text = ''
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            if ri % 2 == 0:
                set_cell_shading(cell, "F0F4F8")
    doc.add_paragraph()
    return table

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E,0x1E,0x1E)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
    return p

def add_separator(doc):
    p = doc.add_paragraph()
    run = p.add_run('─' * 90)
    run.font.size = Pt(5)
    run.font.color.rgb = RGBColor(0xCC,0xCC,0xCC)

# ═══════════════════════════════════════════════════════════
# DOCUMENT 1: HOMEWORK DESIGN
# ═══════════════════════════════════════════════════════════

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

# Style headings
for style_name, size, color in [
    ('Heading 1', 20, RGBColor(0x1B,0x3A,0x5C)),
    ('Heading 2', 15, RGBColor(0x2E,0x5E,0x86)),
    ('Heading 3', 12, RGBColor(0x3B,0x7D,0xAD)),
]:
    st = doc.styles[style_name]
    st.font.name = 'Calibri'
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

# ── TITLE PAGE ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('NETS Homework Design')
run.font.size = Pt(32); run.bold = True; run.font.color.rgb = RGBColor(0x1B,0x3A,0x5C); run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Grade 5 History — §1 Tarix fani nimani o\'rganadi?')
run.font.size = Pt(18); run.font.color.rgb = RGBColor(0x2E,0x5E,0x86); run.font.name = 'Calibri'

doc.add_paragraph()

add_table(doc,
    ['Parameter', 'Value'],
    [
        ['Subject', 'Tarix (Qadimgi dunyo) — History (Ancient World)'],
        ['Grade', '5'],
        ['Chapter', '1-§ Tarix fani nimani o\'rganadi?'],
        ['Textbook Pages', '6–8'],
        ['Language', 'Uzbek (Latin)'],
        ['Duration', '20–30 minutes (Standard Mode)'],
        ['Bloom\'s Coverage', 'Engage → Remember → Understand → Apply → Analyze → Evaluate → Create'],
        ['PISA Domain', 'Reading (primary) + Creative Thinking (secondary) — Dual-domain mapping'],
        ['PISA Target Level', '1–2 (Grade 5 History)'],
        ['Transition Skills', 'L1→L2: locate explicit info; L1→L2: classify by attributes'],
        ['Content Pool', '88 tagged items (meets §10.2 minimum of 85)'],
        ['AI Tiers Used', 'Tier 1 (80%), Tier 2 (15%), Tier 3 (5%)'],
        ['Engine Version', 'NETS Homework Engine Blueprint v1.0'],
        ['SPEC Reference', 'NETS-Homework-Engine-UNIFIED.md v2.0'],
    ]
)

add_separator(doc)

# ── LEARNING OBJECTIVES ──
add_heading(doc, '1. Learning Objectives & Content Map', 1)

add_para(doc, 'Every item in this homework traces to a specific textbook page, a national standard code, a PISA transition skill, and a Bloom\'s level. Nothing is arbitrary.', italic=True)

add_heading(doc, '1.1 Standard Codes', 2)
add_table(doc,
    ['Standard Code', 'Description (Uzbek)', 'Textbook Page', 'Bloom\'s Range'],
    [
        ['UZ-HIST-5-HISTSCI-01', 'Tarix fanining ta\'rifi va maqsadi', '6', 'Remember–Understand'],
        ['UZ-HIST-5-HISTSCI-02', 'Tarixiy manbalar tushunchasi', '6–7', 'Remember–Apply'],
        ['UZ-HIST-5-HISTSCI-03', 'Moddiy manbalar: turlari va misollar', '7', 'Understand–Analyze'],
        ['UZ-HIST-5-HISTSCI-04', 'Yozma manbalar: turlari va misollar', '7–8', 'Understand–Analyze'],
        ['UZ-HIST-5-HISTSCI-05', 'Arxeologiya va arxiv tushunchalari', '7–8', 'Remember–Apply'],
    ]
)

add_heading(doc, '1.2 Key Vocabulary (from Textbook)', 2)
add_table(doc,
    ['Term', 'Definition (from textbook)', 'Page'],
    [
        ['Tarix', 'O\'tmish haqidagi fan — insoniyat hayotini o\'rganadi', '6'],
        ['Tarixiy manba', 'Inson qo\'li bilan yaratilgan yoki yozib qoldirilgan narsalar', '6'],
        ['Moddiy manba', 'Qo\'l bilan ushlab ko\'riladigan, ko\'z bilan ko\'riladigan ashyolar', '7'],
        ['Yozma manba', 'Yozuv orqali saqlab qolingan ma\'lumotlar', '7'],
        ['Arxeolog', 'Qadimgi moddiy manbalarni qidirib topadigan va o\'rganadigan olim', '7'],
        ['Arxiv', 'Eski hujjatlar saqlanadigan maxsus joy', '8'],
        ['Avesto', 'Markaziy Osiyodagi eng qadimgi yozma manba', '8'],
    ]
)

add_separator(doc)

# ── PHASE 0: PRE-SESSION ──
add_heading(doc, '2. Phase 0: Pre-Session Engine', 1)
add_para(doc, '[AI: Tier 1] — Runs in <2 seconds before student sees anything.', bold=True, size=10)

add_heading(doc, '2.1 Student Profile Load', 2)
add_para(doc, 'The engine loads the student\'s complete profile from the database. This homework design is SEPARATE from the student data — it defines what ALL students see. Phase 0 personalizes it per student at runtime.')
add_para(doc, 'Three example profiles are provided as companion JSON files:')
add_bullet(doc, 'student_A_malika.json — High-performing student (PISA Reading 2.1)')
add_bullet(doc, 'student_B_bobur.json — Struggling student (PISA Reading 0.8)')
add_bullet(doc, 'student_C_jasur.json — Brand new student (first session ever)')

add_heading(doc, '2.2 Assignment Payload', 2)
add_code(doc, '''{
  "homework_id": "HW-HIST5-S1-20260403",
  "textbook": "5-sinf Tarixdan hikoyalar (Qadimgi dunyo)",
  "chapter": "1-§ Tarix fani nimani o'rganadi?",
  "pages": "6-8",
  "learning_objectives": [
    "UZ-HIST-5-HISTSCI-01", "UZ-HIST-5-HISTSCI-02",
    "UZ-HIST-5-HISTSCI-03", "UZ-HIST-5-HISTSCI-04",
    "UZ-HIST-5-HISTSCI-05"
  ],
  "mode": "standard",
  "deadline": "2026-04-04T23:59:00+05:00",
  "content_pool_size": 88,
  "content_pool_verified": true
}''')

add_heading(doc, '2.3 Difficulty Band Calculation', 2)
add_para(doc, 'Calculated per student at runtime. Examples for our 3 students:')
add_table(doc,
    ['Student', 'Current PISA', 'Floor', 'Target', 'Ceiling', 'Boss Difficulty Mix'],
    [
        ['Malika (A)', '2.1', '1.0', '2.0', '3.0', '30% Easy / 40% Med / 30% Hard'],
        ['Bobur (B)', '0.8', '0.0', '1.0', '2.0', '50% Easy / 40% Med / 10% Hard'],
        ['Jasur (C)', '1.0 (default)', '0.0', '1.0', '2.0', '40% Easy / 40% Med / 20% Hard'],
    ]
)

add_heading(doc, '2.4 Session Assembly', 2)
add_para(doc, 'AI selects items from the 88-item content pool based on difficulty band. The SAME homework design is used — only item difficulty selection changes per student.')

add_heading(doc, '2.5 Engagement Pre-Load', 2)
add_para(doc, '[HOOK: External Trigger] Push notification sent 2 hours before typical homework time:', bold=True, size=10)
add_table(doc,
    ['Student', 'Push Notification', 'Psychology'],
    [
        ['Malika (streak=23)', '"🔥 23 kunlik seriya! Bugun yangi Tarix darsi — boss tayyor!"', 'Loss aversion — protect the streak'],
        ['Bobur (streak=2)', '"Sardor bugun muzeyga boradi — unga qo\'shiling! 🏛️"', 'Epic meaning — curiosity, not pressure'],
        ['Jasur (new)', '"Birinchi sarguzasht tayyor! Boshlang 🎮"', 'Novelty — first adventure excitement'],
    ]
)

add_separator(doc)

# ── PHASE 1: ENGAGE + REMEMBER ──
add_heading(doc, '3. Phase 1: ENGAGE + REMEMBER (2–3 min)', 1)
add_para(doc, 'Bloom\'s: Foundation → Remember | PISA: Below 1 → L1 | AI: Tier 1 | SPEC: §5.1', bold=True, size=10)

add_heading(doc, '3.1 Hook Game — Speed Match (30 sec)', 2)
add_para(doc, 'Rapid-fire classification: student taps MODDIY MANBA or YOZMA MANBA for each item.')

add_table(doc,
    ['#', 'Item Shown', 'Correct Answer', 'Distractor Logic', 'Standard'],
    [
        ['1', 'Qadimgi tanga (Ancient coin)', 'MODDIY', 'Tangada yozuv bor — student may think yozma', 'HISTSCI-03'],
        ['2', 'Avesto kitobi', 'YOZMA', 'Clear case — builds confidence', 'HISTSCI-04'],
        ['3', 'Devorga chizilgan rasm (Cave painting)', 'MODDIY', 'Visual art ≠ yozma — tests understanding', 'HISTSCI-03'],
        ['4', 'Qog\'ozga yozilgan xat (Letter)', 'YOZMA', 'Clear case', 'HISTSCI-04'],
        ['5', 'Toshdan yasalgan pichoq (Stone knife)', 'MODDIY', 'Clear case', 'HISTSCI-03'],
        ['6', 'Sopol idish (Clay pot)', 'MODDIY', 'No text — purely material', 'HISTSCI-03'],
        ['7', 'Zamonaviy telefon (Modern phone)', 'TARIXIY MANBA EMAS', 'Trick question — modern ≠ historical', 'HISTSCI-02'],
        ['8', 'Qadimgi xarita (Ancient map)', 'YOZMA', 'Maps are written sources', 'HISTSCI-04'],
    ]
)

add_para(doc, 'Scoring: 50 XP per correct. Streak bonus: +10 XP × streak. 20% random "AJOYIB!" animation → +25 bonus XP (variable ratio reinforcement).', size=10)
add_para(doc, '[HOOK: Endowed Progress] Progress bar shows ~15% filled when session loads — student feels "already started."', size=10, bold=True)

add_heading(doc, '3.2 Streak Display + Lesson Brief (30 sec)', 2)
add_para(doc, 'Student sees their streak counter and frame. Then a one-time lesson brief flashes:')
add_code(doc, '''📖 BUGUNGI DARS: Tarix fani nimani o'rganadi?
• Tarix — o'tmish haqidagi fan
• Tarixiy manbalar ikki turga bo'linadi: moddiy va yozma
• Arxeologlar — moddiy manbalarni topadigan olimlar
• Arxiv — eski hujjatlar saqlanadigan joy
• Avesto — Markaziy Osiyodagi eng qadimgi yozma manba

         [ TUSHUNDIM ✓ ]''')
add_para(doc, 'Brief disappears permanently after tap. Cannot be retrieved during session.', italic=True, size=10)

add_heading(doc, '3.3 Flashcard Sprint — Spaced Repetition (60 sec)', 2)
add_para(doc, 'Tests recall of PREVIOUS lessons (Kirish — Introduction, pages 3–5). Items selected by SM-2 algorithm based on student\'s forgetting curve.')

add_table(doc,
    ['#', 'Question', 'Expected Answer', 'Standard', 'Spaced Rep Priority'],
    [
        ['1', '"Vatan" so\'zining ma\'nosi nima?', 'Tug\'ilgan yer, ona yurt', 'UZ-HIST-5-INTRO-01', 'Medium (last seen 5 days ago)'],
        ['2', 'O\'zbekiston poytaxti qayer?', 'Toshkent', 'UZ-HIST-5-INTRO-02', 'Low (answered correctly 3x)'],
        ['3', 'Amir Temur qaysi shaharni poytaxt qilgan?', 'Samarqand', 'UZ-HIST-5-INTRO-03', 'High (wrong last session)'],
        ['4', '"Mustaqillik" so\'zi nimani bildiradi?', 'Erkinlik, o\'z yo\'lini o\'zi belgilash', 'UZ-HIST-5-INTRO-04', 'Medium'],
        ['5', 'O\'zbekiston qachon mustaqil bo\'lgan?', '1991-yil 1-sentabr', 'UZ-HIST-5-INTRO-05', 'High (wrong last session)'],
        ['6', 'Qadimgi Samarqandning boshqa nomi?', 'Marakanda / Afrosiyob', 'UZ-HIST-5-INTRO-06', 'Low'],
    ]
)

add_para(doc, 'Adaptation Gate:', bold=True)
add_bullet(doc, '≥80% (5-6/6 correct): Proceed to Phase 2 confidently')
add_bullet(doc, '60-79% (4/6 correct): Proceed but flag for teacher. AI notes which standards decayed.')
add_bullet(doc, '<60% (0-3/6 correct): 2-minute micro-review before Phase 2. Difficulty band floor lowered by 0.5.')

add_separator(doc)

# ── PHASE 2: STORY MODE ──
add_heading(doc, '4. Phase 2: UNDERSTAND — Story Mode (5–7 min)', 1)
add_para(doc, 'Bloom\'s: Understand | PISA: L2 | AI: Tier 1 (content) + Tier 2 (adaptation) | SPEC: §5.2', bold=True, size=10)
add_para(doc, '[HOOK: Epic Meaning — Octalysis Core Drive 1] Student is the hero of the story.', bold=True, size=10)

add_heading(doc, '4.1 Narrative: "Sardor bobosi bilan muzeyga boradi"', 2)
add_para(doc, 'Three segments interleaved with game breaks. Textbook content wrapped in character-driven narrative.')

add_heading(doc, 'Segment 1: "Tarix nima?" (90 sec)', 3)
add_para(doc, '📖 Darslik: sahifa 6', bold=True, size=10)
add_para(doc, '"Sardor maktabdan qaytayotganda bobosiga savol berdi: \'Bobo, tarix nima o\'zi?\' Bobosi kuldi: \'Tarix — bu o\'tmish haqidagi fan, bolam. Insonlar qanday yashagani, nima qilgani, qanday o\'zgargani — bularning barchasini tarix o\'rganadi.\' Sardor hayron bo\'ldi: \'Lekin biz o\'tmishni qanday bilamiz? O\'sha paytda hech kim video olmagan-ku!\' Bobosi jilmaydi: \'Buning uchun tarixiy manbalar bor...\'"')

add_para(doc, 'CHECKPOINT 1:', bold=True)
add_code(doc, '''Savol: "Tarix fani nimani o'rganadi?"
To'g'ri javob: "O'tmishda insonlar qanday yashagani / insoniyat hayoti"
Bloom's: Understand (define in own words)
PISA: L2 (identify main idea from narrative)
Standard: UZ-HIST-5-HISTSCI-01
On wrong: "📖 Sahifa 6ga qarang. Bobosi Sardorga nima dedi?"
Max retries: 2. After 2 failures → simplified: "Tarix qanday fanga aytiladi — tabiat haqidami yoki o'tmish haqidami?"''')

add_heading(doc, '🎮 Game Break 1: Tile Match (2 min)', 3)
add_para(doc, '6 pairs: match historical term to its definition. Reinforces vocabulary from Segment 1.')
add_table(doc,
    ['Tile A', 'Tile B (Match)', 'Standard'],
    [
        ['Tarix', 'O\'tmish haqidagi fan', 'HISTSCI-01'],
        ['Tarixiy manba', 'Inson yaratgan yoki yozib qoldirgan narsa', 'HISTSCI-02'],
        ['Moddiy manba', 'Ko\'z bilan ko\'riladigan, qo\'l bilan ushlash mumkin', 'HISTSCI-03'],
        ['Yozma manba', 'Yozuv orqali saqlab qolingan', 'HISTSCI-04'],
        ['Arxeolog', 'Qadimgi narsalarni qidirib topadi', 'HISTSCI-05'],
        ['Arxiv', 'Eski hujjatlar saqlanadigan joy', 'HISTSCI-05'],
    ]
)

add_heading(doc, 'Segment 2: "Muzey zali — moddiy manbalar" (90 sec)', 3)
add_para(doc, '📖 Darslik: sahifa 7', bold=True, size=10)
add_para(doc, '"Bobosi Sardorni muzeyga olib bordi. Birinchi zalda ular toshdan yasalgan qurollarni ko\'rdi. \'Bu — moddiy manbalar,\' dedi bobosi. \'Arxeologlar ularni yer ostidan topadi. Har bir tosh, har bir sopol idish bizga qadimgi odamlar haqida hikoya qiladi.\' Sardor ko\'zoynak ortidan qadimniy tangalarni kuzatdi. \'Bu tangalar qancha yillik?\' \'Ikki ming yildan oshiq,\' dedi bobosi. \'Ular ham moddiy manba.\'"')

add_para(doc, 'CHECKPOINT 2:', bold=True)
add_code(doc, '''Savol: "Moddiy manba deganda nimani tushunamiz? Bitta misol keltiring."
To'g'ri javob: "Ko'z bilan ko'riladigan, qo'l bilan ushlanadigan ashyolar.
              Misol: tosh qurol / tanga / sopol idish"
Bloom's: Understand + Apply (define + exemplify)
PISA: L2 (locate + reproduce information with own example)
Standard: UZ-HIST-5-HISTSCI-03
On wrong: "📖 Sahifa 7ga qarang. Muzeyda Sardor nimalarni ko'rdi?"''')

add_heading(doc, '🎮 Game Break 2: Sentence Fill (2 min)', 3)
add_para(doc, '6 sentences with blanks. Tests vocabulary in context.')
add_table(doc,
    ['#', 'Sentence', 'Blank Answer', 'Standard'],
    [
        ['1', 'Tarix — bu _____ haqidagi fan.', 'o\'tmish', 'HISTSCI-01'],
        ['2', 'Inson qo\'li bilan yaratilgan ashyolar _____ manba deyiladi.', 'moddiy', 'HISTSCI-03'],
        ['3', 'Yozuv orqali saqlab qolingan ma\'lumotlar _____ manba deyiladi.', 'yozma', 'HISTSCI-04'],
        ['4', 'Qadimgi moddiy manbalarni qidirib topadigan olim _____ deyiladi.', 'arxeolog', 'HISTSCI-05'],
        ['5', 'Eski hujjatlar saqlanadigan maxsus joy _____ deyiladi.', 'arxiv', 'HISTSCI-05'],
        ['6', 'Markaziy Osiyodagi eng qadimgi yozma manba — _____.', 'Avesto', 'HISTSCI-04'],
    ]
)

add_heading(doc, 'Segment 3: "Kutubxona zali — yozma manbalar" (90 sec)', 3)
add_para(doc, '📖 Darslik: sahifa 7–8', bold=True, size=10)
add_para(doc, '"Muzeyning ikkinchi zalida Sardor shisha ortidagi sariq qog\'ozlarni ko\'rdi. \'Bu yozma manbalar,\' tushuntirdi bobosi. \'Qadimgi odamlar o\'z hayotlarini, qonunlarini, e\'tiqodlarini yozib qoldirishgan. Eng qadimiysi — Avesto. U Markaziy Osiyoda yozilgan.\' \'Bu hujjatlar qayerda saqlanadi?\' so\'radi Sardor. \'Arxivda — maxsus joyda, hamma eski hujjatlar u yerda saqlanadi.\'"')

add_para(doc, 'CHECKPOINT 3:', bold=True)
add_code(doc, '''Savol: "Yozma manbalar moddiy manbalardan nimasi bilan farq qiladi?"
To'g'ri javob: "Moddiy — ashyolar (tosh, tanga).
              Yozma — yozuv orqali saqlangan ma'lumotlar (kitob, xat)."
Bloom's: Analyze (compare/contrast)
PISA: L2-3 (compare information from two categories)
Standard: UZ-HIST-5-HISTSCI-03 + HISTSCI-04
On wrong: "📖 Sahifa 7ga qarang. Birinchi zalda nima bor edi? Ikkinchi zalda-chi?"''')

add_heading(doc, '🎮 Game Break 3: Mystery Box (2–3 min)', 3)
add_para(doc, '[HOOK: Unpredictability — Octalysis Core Drive 7] Student doesn\'t know what\'s in each box.', bold=True, size=10)
add_table(doc,
    ['Box', 'Hidden Item', 'Student Must Classify', 'Correct', 'Standard'],
    [
        ['📦 1', 'Kumush tanga (Silver coin)', 'Moddiy / Yozma / Boshqa', 'MODDIY MANBA', 'HISTSCI-03'],
        ['📦 2', 'Qadimgi qo\'lyozma (Ancient manuscript)', 'Moddiy / Yozma / Boshqa', 'YOZMA MANBA', 'HISTSCI-04'],
        ['📦 3', 'Boboning hikoyasi (Grandfather\'s story)', 'Moddiy / Yozma / Boshqa', 'BOSHQA (og\'zaki manba)', 'HISTSCI-02'],
        ['📦 4', 'Bronza qilich (Bronze sword)', 'Moddiy / Yozma / Boshqa', 'MODDIY MANBA', 'HISTSCI-03'],
        ['📦 5', 'Bugungi gazeta (Today\'s newspaper)', 'Moddiy / Yozma / Boshqa', 'TARIXIY MANBA EMAS (zamonaviy)', 'HISTSCI-02'],
    ]
)

add_separator(doc)

# ── PHASE 4: REAL-LIFE CHALLENGE ──
add_heading(doc, '5. Phase 4: ANALYZE — Real-Life Challenge (3–5 min)', 1)
add_para(doc, 'Bloom\'s: Analyze→Evaluate | PISA: L3–5 | AI: Tier 2 | SPEC: §5.4', bold=True, size=10)
add_para(doc, 'Eligibility: PISA ≥ 2 proceed normally. PISA < 2 (like Bobur) → replaced with additional Adaptive Quiz at Apply level.', italic=True)

add_heading(doc, '5.1 Scenario: "Mahallangiz tarixchisi"', 2)
add_code(doc, '''KONTEKST:
Tasavvur qiling: siz mahallangizning tarixini yozmoqchisiz.
Mahallangizda siz quyidagilarni topa olasiz:
  (a) Eski uy — 100 yildan oshiq
  (b) Buvi Zulfiyaning hikoyalari — u 85 yoshda
  (c) Eski suratlar — 1960-yillardan
  (d) Masjidning eski yozuvi — 200 yillik

SAVOLLAR:
1. "Har bir topilmani turkumlang: moddiy manba, yozma manba, yoki boshqa tur?"
   [Apply — classify using textbook categories]

2. "Qaysi manba sizga mahalla haqida ENG ko'p ma'lumot beradi? Nima uchun?"
   [Evaluate — judge reliability and richness of sources]

3. "Agar yozma manbalar yo'qolgan bo'lsa, faqat moddiy manbalar qolsa —
    mahalla haqida nimalarni BILIB BO'LADI va nimalarni BILIB BO'LMAYDI?"
   [Evaluate — assess limitations of evidence types]''')

add_para(doc, 'AI Evaluation Rubric:', bold=True)
add_table(doc,
    ['Component', 'Weight', 'What AI Checks'],
    [
        ['Concept ID', '25%', 'Did student correctly classify all 4 items?'],
        ['Reasoning', '50%', 'Is the "why" logical? Does it reference what each source type can/cannot show?'],
        ['Final Answer', '25%', 'Is the limitation analysis reasonable? (Multiple valid answers accepted)'],
    ]
)
add_para(doc, '[HOOK: Reward Chest] On completion → animated chest opens with 200–500 variable XP + 15% avatar item chance.', bold=True, size=10)

add_heading(doc, '5.2 Below-PISA-L2 Alternative (for Bobur)', 2)
add_para(doc, 'Bobur gets an Adaptive Quiz instead: 5 classification questions at Apply level. Same standards, simpler format (identify + classify, no evaluation required).')

add_separator(doc)

# ── PHASE 5: CONSOLIDATION ──
add_heading(doc, '6. Phase 5: CONSOLIDATION — Memory Palace (2–3 min)', 1)
add_para(doc, 'Bloom\'s: Remember (reinforce) | AI: Tier 1 | SPEC: §5.5', bold=True, size=10)
add_para(doc, 'The "save point" before the boss fight. Student walks through Registan Square and places 5 concepts.')

add_table(doc,
    ['Location', 'Concept Placed', 'Mnemonic Connection', 'Standard'],
    [
        ['Asosiy darvoza (Main gate)', 'Tarix = o\'tmish haqidagi fan', 'Gate to the past', 'HISTSCI-01'],
        ['Chap madrasah (Left building)', 'Moddiy manbalar (+ tanga image)', 'Building = physical object', 'HISTSCI-03'],
        ['Favvora (Fountain)', 'Yozma manbalar (+ Avesto image)', 'Water = ink flowing on paper', 'HISTSCI-04'],
        ['O\'ng madrasah (Right building)', 'Arxeolog = moddiy manba topuvchi', 'Explorer entering building', 'HISTSCI-05'],
        ['Orqa masjid (Back mosque)', 'Arxiv = hujjat saqlash joyi', 'Sacred place stores sacred texts', 'HISTSCI-05'],
    ]
)

add_para(doc, 'Recall test: 5 questions ("Asosiy darvozada nima edi?"). 50 XP per correct. All 5 = +50 bonus.', size=10)

add_separator(doc)

# ── PHASE 6: BOSS FIGHT ──
add_heading(doc, '7. Phase 6: EVALUATE + CREATE — AI Boss Fight (5–10 min)', 1)
add_para(doc, 'Bloom\'s: Apply→Create (tiered) | PISA: L3–6 | AI: Tier 1 + Tier 3 | SPEC: §5.6', bold=True, size=10)
add_para(doc, 'THIS IS THE MANDATORY MASTERY GATE. ZERO EXCEPTIONS. NO SKIPPING.', bold=True)

add_heading(doc, '7.1 Boss Parameters', 2)
add_table(doc,
    ['Parameter', 'Value'],
    [
        ['Boss HP', '100 (Grades 5–8 tier)'],
        ['Questions', '5 (2 Easy + 2 Medium + 1 Hard)'],
        ['Max Damage', '10+10+20+20+30 = 90 HP base. With combo (3 correct → 2×): enough to defeat.'],
        ['Hint Tax', '−10 HP per hint (boss heals)'],
        ['Combo', '3 correct in row → 2× damage on next'],
        ['Format', 'Short answer + open reasoning. NO multiple choice (Grade 5+).'],
        ['Retry on Fail', 'Up to 3 attempts. Questions REGENERATED each time (same standards, different context).'],
    ]
)

add_heading(doc, '7.2 Boss Questions', 2)

add_para(doc, 'EASY Q1 (−10 HP):', bold=True)
add_code(doc, '''Savol: "Tarix fani nimani o'rganadi? Qisqacha o'z so'zlaringiz bilan tushuntiring."
Standard: UZ-HIST-5-HISTSCI-01
Bloom's: Apply (define in own words, not memorized textbook sentence)
PISA: L2 (Reading — reproduce information in own words)
Transition Skill: L1→L2: locate explicit info
Correct: Any response defining history as study of the human past / o'tmish haqidagi fan
Partial credit: Correct idea but incomplete (50% damage = −5 HP)
Common error: "Tarix — bu darslik" → Misconception: confusing the subject with the textbook
Socratic hint: "Tarix FANI nimani O'RGANADI? Fan va kitob farqi bormi?"''')

add_para(doc, 'EASY Q2 (−10 HP):', bold=True)
add_code(doc, '''Savol: "Moddiy manbaga ikkita misol keltiring va ularning nimani ko'rsatishini tushuntiring."
Standard: UZ-HIST-5-HISTSCI-03
Bloom's: Apply (recall + explain significance)
PISA: L2 (Reading — access and retrieve + basic inference)
Transition Skill: L1→L2: classify by attributes
Correct: Two valid examples (tanga, tosh qurol, sopol, etc.) + what they reveal
Partial credit: Examples correct but no explanation (50% damage)
Common error: "Kitob" → Misconception: mixing moddiy and yozma
Socratic hint: "Moddiy manba — qo'l bilan ushlab ko'rsa bo'ladimi? Kitob — yozma emas mi?"''')

add_para(doc, 'MEDIUM Q3 (−20 HP):', bold=True)
add_code(doc, '''Savol: "Moddiy va yozma manbalarning O'XSHASH va FARQLI jihatlarini tushuntiring."
Standard: UZ-HIST-5-HISTSCI-03 + HISTSCI-04
Bloom's: Analyze (compare/contrast)
PISA: L3 (Reading — integrate and interpret from two categories)
Transition Skill: L1→L2: classify by attributes → L2→L3: compare/contrast
Correct: Both are tarixiy manbalar (similar). Moddiy = physical objects, yozma = written records (different).
Partial credit: Only differences OR only similarities (50% damage)
Common error: "Moddiy — eski, yozma — yangi" → Misconception: age-based, not type-based
Socratic hint: "Ikkalasi ham tarixiy manba. Lekin QANDAY farq qiladi — biri qanday, ikkinchisi qanday?"''')

add_para(doc, 'MEDIUM Q4 (−20 HP):', bold=True)
add_code(doc, '''Savol: "Nima uchun yozma manbalar tarixchilar uchun juda muhim?
         Dalillar bilan javob bering."
Standard: UZ-HIST-5-HISTSCI-04
Bloom's: Evaluate (argue with evidence)
PISA: L3 (Reading — reflect and evaluate)
Transition Skill: L2→L3: evaluate relevance of information
Correct: Yozma manbalar aniq ma'lumot beradi — sanalar, ismlar, voqealar.
         Moddiy manbalar bunday aniqlik bera olmaydi.
Partial credit: General answer without specific evidence (50% damage)
Common error: "Chunki yozilgan" → Tautology, not actual reasoning
Socratic hint: "Moddiy manba — masalan, tanga — sizga aniq SANA ayta oladimi? Yozma-chi?"''')

add_para(doc, 'HARD Q5 (−30 HP):', bold=True)
add_code(doc, '''Savol: "Tasavvur qiling: arxeologlar qadimgi shahar xarobasini topdi.
         Ular FAQAT moddiy manbalar topdi, yozma manbalar yo'q.
         Bu shahar haqida nimalarni BILIB BO'LADI
         va nimalarni BILIB BO'LMAYDI?"
Standard: UZ-HIST-5-HISTSCI-03 + HISTSCI-04
Bloom's: Evaluate + Create (assess limitations, generate hypotheses)
PISA: L4 (Reading — reflect and evaluate) + Creative Thinking (generate diverse ideas)
Transition Skill: L2→L3: evaluate what evidence can/cannot show
Correct: BILIB BO'LADI — qanday qurollar ishlatgan, nima yegan, qanday uylar qurgan.
         BILIB BO'LMAYDI — til, din, qonunlar, aniq sanalar, odamlar ismi.
Partial credit: Only "bilib bo'ladi" or only "bilib bo'lmaydi" (50% damage = −15 HP)
Common error: "Hamma narsani bilib bo'ladi" → Misconception: overestimating material sources
Socratic hint: "Tosh qurol sizga ism ayta oladimi? Qachon yasalganini aniq ayta oladimi?"''')

add_heading(doc, '7.3 Victory Sequence', 2)
add_table(doc,
    ['Stars', 'Criteria', 'XP', 'Reward'],
    [
        ['⭐', 'Boss defeated (any attempts)', '500', 'Chapter marked complete'],
        ['⭐⭐', '≤2 attempts, >50% boss HP dealt', '700', '+ Avatar item (25% chance)'],
        ['⭐⭐⭐', '1st attempt, no hints, >80% HP dealt', '1000', '+ Rare badge progress (50% chance)'],
    ]
)

add_para(doc, '[HOOK: Reward Chest] Animated treasure chest with variable contents. [HOOK: Near-Miss] On defeat: "Boss HP: 15 qoldi! Ozgina edi!" creates retry motivation.', bold=True, size=10)

add_separator(doc)

# ── PHASE 7-8: AI ANALYSIS + REMEDIATION ──
add_heading(doc, '8. Phase 7–8: AI Analysis + Remediation + Reflection (3–5 min)', 1)
add_para(doc, 'AI: Tier 1 (computation) + Tier 2 (report generation) | SPEC: §5.7, §9.3', bold=True, size=10)

add_heading(doc, '8.1 AI Analysis Engine (automatic, <10 sec)', 2)
add_para(doc, 'See Blueprint Section 7. AI calculates: accuracy per phase/standard/transition skill, root cause identification, PISA level update, comparison to previous sessions, optimal plan for next session.')

add_heading(doc, '8.2 Remediation: Targeted Micro-Exercises', 2)
add_para(doc, 'Max 2–3 weak areas. Framed as "BONUS XP IMKONIYATI" (not "you failed").')
add_para(doc, 'Example if student failed Q3 (compare moddiy vs yozma):')
add_code(doc, '''Micro-exercise: "Quyidagilarni to'g'ri joyga qo'ying:"

  MODDIY MANBA          YOZMA MANBA
  ─────────────          ─────────────
  [drag items here]      [drag items here]

  Items: tanga, xat, tosh qurol, kitob, sopol idish, qo'lyozma

Correct → "stabilizing" → review in 3 days → +100 XP
Still wrong → "persistent weakness" → review TOMORROW → teacher notified''')

add_heading(doc, '8.3 Reflection Prompts', 2)
add_table(doc,
    ['Performance', 'Prompt (Uzbek)', 'Purpose'],
    [
        ['80%+ (like Malika)', '"Bugun qaysi strategiya sizga eng ko\'p yordam berdi?"', 'Reinforce success'],
        ['60–79%', '"Eng qiyin savol qaysi bo\'ldi? Qanday hal qildingiz?"', 'Build awareness'],
        ['<60% (like Bobur)', '"Bu mavzuni yaxshiroq tushunish uchun sizga nima yordam beradi?"', 'Self-directed learning'],
        ['New (Jasur)', '"Birinchi dars qanday bo\'ldi? Nima qiziq edi?"', 'Build engagement'],
    ]
)

add_heading(doc, '8.4 Session Closure', 2)
add_para(doc, 'XP tally, star display, PISA change, streak update, reward chest, next lesson preview ("§2 — Tarixiy sana hisobi"), leaderboard position update (if opted in).', size=10)

add_separator(doc)

# ── ANTI-CHEAT & AI DETECTION ──
add_heading(doc, '9. Anti-Cheat & AI Usage Detection System', 1)
add_para(doc, 'This section defines how the homework engine detects cheating, AI-assisted answers (from external AI tools like ChatGPT), and suspicious behavior — using DIRECT COMPARISON to each student\'s historical data.', bold=True)

add_heading(doc, '9.1 Behavioral Baseline Comparison', 2)
add_para(doc, 'The core principle: every student has a behavioral fingerprint built from their previous sessions. Deviations from this fingerprint trigger flags.')

add_table(doc,
    ['Metric', 'Baseline Source', 'Flag Trigger', 'Severity'],
    [
        ['Response Time', 'avg_response_time from previous 5 sessions', 'Current response <40% of baseline OR >300% of baseline', 'MEDIUM'],
        ['Vocabulary Complexity', 'avg_word_length + unique_words from reflection/boss answers', 'Sudden jump >2 standard deviations above baseline', 'HIGH (AI indicator)'],
        ['Answer Length', 'avg_answer_chars from previous boss fights', 'Answer >3× typical length with advanced structure', 'HIGH (AI indicator)'],
        ['Accuracy Jump', 'rolling_avg_accuracy from last 5 sessions', 'Accuracy jumps >30% in single session', 'HIGH'],
        ['Hesitation Pattern', 'avg_time_between_keystrokes', 'Uniform typing speed (no pauses) — human answers have variable pace', 'MEDIUM (AI copy-paste indicator)'],
        ['Error Pattern', 'common_error_types from mastery_map', 'Student suddenly avoids ALL previously common errors', 'HIGH (AI indicator)'],
        ['Conceptual Leap', 'highest_blooms_achieved from history', 'Student jumps from Understand to Evaluate without Apply/Analyze progress', 'HIGH'],
    ]
)

add_heading(doc, '9.2 AI-Generated Answer Detection', 2)
add_para(doc, 'External AI tools (ChatGPT, Gemini, etc.) produce answers with detectable patterns:')

add_table(doc,
    ['AI Signature', 'Detection Method', 'Example'],
    [
        ['Overly structured answers', 'AI answers tend to have numbered lists, "Firstly/Secondly" structure. Grade 5 students don\'t write like this.', 'Student suddenly writes: "Birinchidan... Ikkinchidan... Xulosa qilib aytganda..."'],
        ['Perfect Uzbek grammar', 'Student\'s historical answers have consistent grammar errors. Sudden perfection = flag.', 'Bobur usually writes "manba bu" (wrong order) → suddenly writes perfect sentences'],
        ['Textbook-verbatim answers', 'Student copies exact textbook text instead of own words. Compare answer to textbook via similarity score.', 'Answer matches sahifa 7 word-for-word (>90% similarity)'],
        ['Vocabulary mismatch', 'Answer uses words the student has never used before AND are not in the lesson vocabulary.', 'Student uses "epistemologik" — not in Grade 5 vocabulary'],
        ['Response time anomaly', 'Long complex answer typed in <30 seconds. Human typing speed baseline check.', '200-character answer typed in 15 seconds = copy-paste'],
        ['Paste event detection', 'System detects Ctrl+V / paste events in text fields.', 'Direct paste → immediate flag'],
    ]
)

add_heading(doc, '9.3 Comparison Verdict Engine', 2)
add_para(doc, 'Every answer is scored against the student\'s personal baseline. The verdict is NOT based on answer quality alone — it\'s based on DEVIATION from expected behavior.')

add_code(doc, '''VERDICT ALGORITHM:

FOR each boss_answer:
  1. COMPARE response_time to student.avg_response_time
     → IF < 40% of baseline: flag "SPEED_ANOMALY"
     → IF paste_event detected: flag "PASTE_DETECTED"

  2. COMPARE vocabulary_complexity to student.avg_vocabulary_level
     → IF > 2σ above baseline: flag "VOCAB_JUMP"

  3. COMPARE answer_structure to student.typical_answer_structure
     → IF suddenly uses numbered lists, formal connectors: flag "STRUCTURE_ANOMALY"

  4. COMPARE accuracy_this_session to student.rolling_avg_accuracy
     → IF jump > 30%: flag "ACCURACY_SPIKE"

  5. COMPARE error_pattern to student.common_errors
     → IF zero common errors AND previous sessions had >40% common error rate: flag "ERROR_ABSENCE"

  6. CHECK textbook_similarity
     → IF answer >85% match to textbook text: flag "VERBATIM_COPY"

VERDICT LOGIC:
  0 flags: CLEAN ✅
  1 flag:  MONITOR 👁️ (no action, log for pattern)
  2 flags: SOFT WARNING ⚠️ (student sees: "O'z so'zlaringiz bilan javob bering!")
  3+ flags: TEACHER ALERT 🚨 (teacher dashboard notification with evidence)
  PASTE_DETECTED + VOCAB_JUMP: IMMEDIATE TEACHER ALERT (high confidence AI use)''')

add_heading(doc, '9.4 Student-Specific Baseline Examples', 2)

add_para(doc, 'How the system treats each of our 3 students differently:', bold=True)

add_table(doc,
    ['Check', 'Malika (Smart)', 'Bobur (Struggling)', 'Jasur (New)'],
    [
        ['Baseline exists?', 'YES — 15 sessions of rich data', 'YES — 8 sessions of data', 'NO — no baseline yet'],
        ['Response time baseline', '45 sec avg for boss Qs', '90 sec avg (slower reader)', 'N/A — uses grade-level default (60 sec)'],
        ['Vocabulary baseline', 'Rich vocab, uses own examples', 'Simple sentences, frequent grammar errors', 'N/A — first 3 sessions build baseline'],
        ['Accuracy baseline', '82% rolling avg', '45% rolling avg', 'N/A — default 50% expected'],
        ['If Malika\'s boss answer is perfect', 'CLEAN ✅ (consistent with baseline)', '—', '—'],
        ['If Bobur suddenly scores 95%', '—', 'ACCURACY_SPIKE 🚨 (jump from 45% to 95%)', '—'],
        ['If Bobur\'s answer has perfect grammar', '—', 'VOCAB_JUMP + STRUCTURE_ANOMALY → TEACHER ALERT 🚨', '—'],
        ['If Jasur\'s first answer is PhD-level', '—', '—', 'VOCAB_JUMP → SOFT WARNING ⚠️ (no baseline, so lenient)'],
    ]
)

add_heading(doc, '9.5 Anti-Cheat During Boss Fight', 2)
add_table(doc,
    ['Mechanism', 'How It Works', 'SPEC Reference'],
    [
        ['Question Regeneration', 'Every retry generates NEW questions (same standards, different context/numbers)', '§13.1'],
        ['Socratic Verification', 'After boss defeat: 1 follow-up question to verify understanding (not scored, but flagged if failed)', '§13.4'],
        ['Device Lock', 'Same account on multiple devices → session locked immediately', '§13.5'],
        ['Session Pattern', 'Perfect on hard + fail on easy = unusual pattern → flagged', '§13.3'],
        ['Tab-Switch Detection', 'System detects when student switches to another app/tab during boss fight', 'New'],
        ['Time-Gap Analysis', 'Long pause (>60 sec) on a question followed by sudden detailed answer → flag', 'New'],
    ]
)

add_heading(doc, '9.6 Escalation Ladder', 2)
add_table(doc,
    ['Level', 'Trigger', 'Action', 'Visibility'],
    [
        ['Level 0', 'CLEAN — no flags', 'Nothing. Normal session.', 'Student only'],
        ['Level 1', '1 flag in session', 'Log it. Monitor next 2 sessions for pattern.', 'System log only'],
        ['Level 2', '2 flags in session OR 3+ flags across 2 sessions', '"O\'z so\'zlaringiz bilan javob bering!" soft warning shown to student', 'Student sees warning'],
        ['Level 3', '3+ flags in single session', 'Teacher notification with evidence log (which flags, which questions)', 'Teacher dashboard'],
        ['Level 4', 'Pattern persists across 3+ sessions', 'Parent notification. Admin review. Session data preserved as evidence.', 'Teacher + Parent + Admin'],
    ]
)

add_para(doc, 'CRITICAL: The system NEVER accuses the student of cheating. It flags anomalies. The TEACHER makes the judgment call with evidence provided. False positives are expected and acceptable — the escalation ladder ensures no punishment happens without human review.', bold=True, italic=True, size=10)

add_separator(doc)

# ── SESSION DATA FLOW ──
add_heading(doc, '10. Complete Session Data Flow', 1)
add_code(doc, '''
PRE-SESSION (Phase 0)
  └─ Load student_profile.json ←── THIS IS SEPARATE FROM HOMEWORK DESIGN
  └─ Load homework_design.json ←── THIS DOCUMENT
  └─ Calculate difficulty_band (personalized per student)
  └─ Assemble session (select items from content pool)

DURING SESSION (Phases 1–6)
  └─ Every answer recorded: response, time, keystrokes, paste_events
  └─ Real-time adaptation: difficulty adjusts within games
  └─ Anti-cheat: continuous monitoring against student baseline

POST-SESSION (Phases 7–8)
  └─ AI Analysis: accuracy, root cause, PISA update
  └─ Anti-cheat verdict: clean / monitor / warning / alert
  └─ Remediation: targeted micro-exercises
  └─ Reflection: metacognitive prompt
  └─ Data push: update student_profile.json for NEXT session
  └─ Reports: teacher summary, parent summary

THE LOOP:
  Today's session output → Tomorrow's Phase 0 input
  Every session makes the next session smarter.
''')

# Save homework design
output1 = "demo/Grade5-History-S1-Homework-Design.docx"
os.makedirs("demo", exist_ok=True)
doc.save(output1)
print(f"✅ Created: {output1}")

# ═══════════════════════════════════════════════════════════
# STUDENT DATA FILES (JSON)
# ═══════════════════════════════════════════════════════════

# ── STUDENT A: MALIKA (High-performing) ──
student_a = {
    "student_id": "STU-2026-05A-0147",
    "name": "Malika Karimova",
    "grade": 5,
    "class": "5-A",
    "school": "174-maktab, Toshkent",
    "language_preference": "uz",
    "account_created": "2026-02-01",
    "total_sessions_completed": 15,
    "profile_type": "HIGH_PERFORMER",

    "pisa_levels": {
        "history": {
            "reading": 2.1,
            "creative_thinking": 1.8
        },
        "mathematics": 2.3,
        "reading": 2.4,
        "science": 2.0
    },

    "domain_breakdown": {
        "history_reading": {
            "access_and_retrieve": 2.3,
            "integrate_and_interpret": 2.0,
            "reflect_and_evaluate": 1.8
        }
    },

    "mastery_map": {
        "UZ-HIST-5-INTRO-01": {"score": 0.92, "last_seen": "2026-03-30", "sessions_seen": 3, "decay_factor": 0.95},
        "UZ-HIST-5-INTRO-02": {"score": 0.88, "last_seen": "2026-03-30", "sessions_seen": 3, "decay_factor": 0.93},
        "UZ-HIST-5-INTRO-03": {"score": 0.95, "last_seen": "2026-03-30", "sessions_seen": 3, "decay_factor": 0.96},
        "UZ-HIST-5-INTRO-04": {"score": 0.78, "last_seen": "2026-03-28", "sessions_seen": 2, "decay_factor": 0.88},
        "UZ-HIST-5-INTRO-05": {"score": 0.85, "last_seen": "2026-03-30", "sessions_seen": 3, "decay_factor": 0.91},
        "UZ-HIST-5-INTRO-06": {"score": 0.90, "last_seen": "2026-03-30", "sessions_seen": 2, "decay_factor": 0.94}
    },

    "transition_skills": {
        "mastered": [
            "L1->L2: locate explicit info",
            "L1->L2: single-step retrieval"
        ],
        "in_progress": [
            "L1->L2: classify by attributes",
            "L2->L3: compare/contrast from two sources"
        ],
        "not_started": [
            "L2->L3: evaluate relevance of information"
        ]
    },

    "spaced_repetition_queue": [
        {"standard": "UZ-HIST-5-INTRO-04", "priority": "medium", "due_date": "2026-04-03", "reason": "approaching_threshold"},
        {"standard": "UZ-HIST-5-INTRO-02", "priority": "low", "due_date": "2026-04-05", "reason": "scheduled_review"}
    ],

    "engagement_data": {
        "streak": {"current": 23, "longest": 23, "freeze_available": True, "freeze_count_used": 1},
        "xp_total": 18750,
        "xp_this_week": 4200,
        "badges": ["Seriyachi (7-kun)", "Yulduz yig'uvchi (10 boss)"],
        "avatar_items": ["Ipak yo'li karavon shlyapasi", "Oltin yulduz"],
        "leaderboard_position": 2,
        "leaderboard_opted_in": True
    },

    "behavioral_baseline": {
        "avg_response_time_boss_sec": 45,
        "avg_response_time_game_sec": 8,
        "avg_answer_length_chars": 85,
        "avg_vocabulary_level": "grade_appropriate_plus",
        "avg_word_count_boss": 18,
        "avg_unique_words_per_answer": 14,
        "typical_answer_structure": "simple_sentences_with_examples",
        "common_grammar_errors": ["occasional_comma_splice"],
        "typing_speed_chars_per_sec": 2.8,
        "paste_events_historical": 0,
        "tab_switches_per_session": 0.5
    },

    "session_history": [
        {
            "session_id": "SES-20260330-001",
            "topic": "Kirish (Introduction)",
            "date": "2026-03-30",
            "accuracy": 0.88,
            "boss_stars": 3,
            "boss_attempts": 1,
            "time_spent_min": 22,
            "anti_cheat_verdict": "CLEAN",
            "flags": []
        },
        {
            "session_id": "SES-20260328-001",
            "topic": "Kirish (Introduction) — Part 2",
            "date": "2026-03-28",
            "accuracy": 0.82,
            "boss_stars": 2,
            "boss_attempts": 1,
            "time_spent_min": 25,
            "anti_cheat_verdict": "CLEAN",
            "flags": []
        },
        {
            "session_id": "SES-20260326-001",
            "topic": "Kirish (Introduction) — Part 1",
            "date": "2026-03-26",
            "accuracy": 0.85,
            "boss_stars": 3,
            "boss_attempts": 1,
            "time_spent_min": 21,
            "anti_cheat_verdict": "CLEAN",
            "flags": []
        }
    ],

    "trajectory": "on_track",
    "teacher_notes": [],
    "parent_language": "uz"
}

# ── STUDENT B: BOBUR (Struggling) ──
student_b = {
    "student_id": "STU-2026-05A-0283",
    "name": "Bobur Toshmatov",
    "grade": 5,
    "class": "5-A",
    "school": "174-maktab, Toshkent",
    "language_preference": "uz",
    "account_created": "2026-02-01",
    "total_sessions_completed": 8,
    "profile_type": "STRUGGLING",

    "pisa_levels": {
        "history": {
            "reading": 0.8,
            "creative_thinking": 0.5
        },
        "mathematics": 1.1,
        "reading": 0.9,
        "science": 1.0
    },

    "domain_breakdown": {
        "history_reading": {
            "access_and_retrieve": 1.1,
            "integrate_and_interpret": 0.6,
            "reflect_and_evaluate": 0.4
        }
    },

    "mastery_map": {
        "UZ-HIST-5-INTRO-01": {"score": 0.55, "last_seen": "2026-03-30", "sessions_seen": 3, "decay_factor": 0.72},
        "UZ-HIST-5-INTRO-02": {"score": 0.40, "last_seen": "2026-03-30", "sessions_seen": 3, "decay_factor": 0.65},
        "UZ-HIST-5-INTRO-03": {"score": 0.35, "last_seen": "2026-03-28", "sessions_seen": 2, "decay_factor": 0.60},
        "UZ-HIST-5-INTRO-04": {"score": 0.30, "last_seen": "2026-03-25", "sessions_seen": 1, "decay_factor": 0.50},
        "UZ-HIST-5-INTRO-05": {"score": 0.20, "last_seen": "2026-03-28", "sessions_seen": 2, "decay_factor": 0.45},
        "UZ-HIST-5-INTRO-06": {"score": 0.15, "last_seen": "2026-03-25", "sessions_seen": 1, "decay_factor": 0.40}
    },

    "transition_skills": {
        "mastered": [],
        "in_progress": [
            "L1->L2: locate explicit info"
        ],
        "not_started": [
            "L1->L2: classify by attributes",
            "L1->L2: single-step retrieval",
            "L2->L3: compare/contrast from two sources"
        ]
    },

    "spaced_repetition_queue": [
        {"standard": "UZ-HIST-5-INTRO-03", "priority": "critical", "due_date": "2026-04-03", "reason": "persistent_weakness"},
        {"standard": "UZ-HIST-5-INTRO-05", "priority": "critical", "due_date": "2026-04-03", "reason": "persistent_weakness"},
        {"standard": "UZ-HIST-5-INTRO-04", "priority": "high", "due_date": "2026-04-03", "reason": "wrong_last_session"},
        {"standard": "UZ-HIST-5-INTRO-06", "priority": "high", "due_date": "2026-04-04", "reason": "wrong_last_session"},
        {"standard": "UZ-HIST-5-INTRO-01", "priority": "medium", "due_date": "2026-04-05", "reason": "approaching_threshold"},
        {"standard": "UZ-HIST-5-INTRO-02", "priority": "medium", "due_date": "2026-04-05", "reason": "approaching_threshold"}
    ],

    "engagement_data": {
        "streak": {"current": 2, "longest": 5, "freeze_available": False, "freeze_count_used": 3},
        "xp_total": 4200,
        "xp_this_week": 800,
        "badges": [],
        "avatar_items": ["Oddiy qalpoq"],
        "leaderboard_position": 28,
        "leaderboard_opted_in": False
    },

    "behavioral_baseline": {
        "avg_response_time_boss_sec": 90,
        "avg_response_time_game_sec": 15,
        "avg_answer_length_chars": 35,
        "avg_vocabulary_level": "below_grade",
        "avg_word_count_boss": 7,
        "avg_unique_words_per_answer": 5,
        "typical_answer_structure": "short_fragments_or_single_words",
        "common_grammar_errors": ["word_order_errors", "missing_suffixes", "spelling_errors"],
        "typing_speed_chars_per_sec": 1.4,
        "paste_events_historical": 0,
        "tab_switches_per_session": 1.2
    },

    "session_history": [
        {
            "session_id": "SES-20260330-002",
            "topic": "Kirish (Introduction)",
            "date": "2026-03-30",
            "accuracy": 0.42,
            "boss_stars": 1,
            "boss_attempts": 3,
            "time_spent_min": 35,
            "anti_cheat_verdict": "CLEAN",
            "flags": []
        },
        {
            "session_id": "SES-20260328-002",
            "topic": "Kirish (Introduction) — Part 2",
            "date": "2026-03-28",
            "accuracy": 0.38,
            "boss_stars": 0,
            "boss_attempts": 3,
            "time_spent_min": 38,
            "anti_cheat_verdict": "CLEAN",
            "flags": [],
            "teacher_notified": True,
            "teacher_action": "Scheduled 1:1 session for 2026-03-29"
        },
        {
            "session_id": "SES-20260325-002",
            "topic": "Kirish (Introduction) — Part 1",
            "date": "2026-03-25",
            "accuracy": 0.45,
            "boss_stars": 1,
            "boss_attempts": 2,
            "time_spent_min": 32,
            "anti_cheat_verdict": "CLEAN",
            "flags": []
        }
    ],

    "trajectory": "at_risk",
    "boost_mode_active": True,
    "boost_mode_activated": "2026-03-30",
    "teacher_notes": [
        {"date": "2026-03-29", "note": "1:1 session conducted. Bobur has vocabulary gaps — difficulty reading textbook independently. Recommend bilingual support."},
        {"date": "2026-03-30", "note": "Boost Mode activated. Extra practice sessions available."}
    ],
    "parent_language": "uz",
    "parent_notified": True,
    "parent_notification_date": "2026-03-30"
}

# ── STUDENT C: JASUR (Brand New) ──
student_c = {
    "student_id": "STU-2026-05A-0391",
    "name": "Jasur Alimov",
    "grade": 5,
    "class": "5-A",
    "school": "174-maktab, Toshkent",
    "language_preference": "uz",
    "account_created": "2026-04-03",
    "total_sessions_completed": 0,
    "profile_type": "NEW_STUDENT",

    "pisa_levels": {
        "history": {
            "reading": 1.0,
            "creative_thinking": 1.0
        },
        "mathematics": 1.0,
        "reading": 1.0,
        "science": 1.0
    },
    "_pisa_note": "All PISA levels set to grade-level default (1.0) for new students. Will be calibrated after first 3 sessions.",

    "domain_breakdown": {
        "history_reading": {
            "access_and_retrieve": 1.0,
            "integrate_and_interpret": 1.0,
            "reflect_and_evaluate": 1.0
        },
        "_note": "All domains set to default. No data yet."
    },

    "mastery_map": {},
    "_mastery_note": "Empty — no previous content seen. All standards start at 0.0.",

    "transition_skills": {
        "mastered": [],
        "in_progress": [],
        "not_started": [
            "L1->L2: locate explicit info",
            "L1->L2: classify by attributes",
            "L1->L2: single-step retrieval",
            "L2->L3: compare/contrast from two sources",
            "L2->L3: evaluate relevance of information"
        ],
        "_note": "All skills 'not_started' — first session will begin populating."
    },

    "spaced_repetition_queue": [],
    "_queue_note": "Empty — no previous items to review. Phase 1 Memory Sprint will use grade-level warmup questions instead of spaced repetition.",

    "engagement_data": {
        "streak": {"current": 0, "longest": 0, "freeze_available": False, "freeze_count_used": 0},
        "xp_total": 0,
        "xp_this_week": 0,
        "badges": [],
        "avatar_items": ["Boshlang'ich libos (starter outfit)"],
        "leaderboard_position": None,
        "leaderboard_opted_in": True,
        "_note": "New student starts with default avatar. First XP earned today."
    },

    "behavioral_baseline": {
        "avg_response_time_boss_sec": None,
        "avg_response_time_game_sec": None,
        "avg_answer_length_chars": None,
        "avg_vocabulary_level": None,
        "avg_word_count_boss": None,
        "avg_unique_words_per_answer": None,
        "typical_answer_structure": None,
        "common_grammar_errors": None,
        "typing_speed_chars_per_sec": None,
        "paste_events_historical": 0,
        "tab_switches_per_session": None,
        "_note": "NO BASELINE EXISTS. First 3 sessions build the behavioral fingerprint. Anti-cheat system uses grade-level defaults and is LENIENT during calibration period."
    },

    "calibration_mode": {
        "active": True,
        "sessions_remaining": 3,
        "anti_cheat_sensitivity": "low",
        "description": "New students get 3 calibration sessions. During this period: anti-cheat flags require DOUBLE the threshold to trigger. PISA levels adjust more aggressively (±0.3 per session vs normal ±0.1). After 3 sessions, student transitions to normal monitoring."
    },

    "session_history": [],

    "trajectory": "unknown",
    "teacher_notes": [],
    "parent_language": "uz",
    "parent_notified": False,

    "first_session_special_rules": {
        "memory_sprint": "Uses grade-level warmup questions (not spaced repetition — no history exists)",
        "difficulty_band": "Standard Grade 5 default: floor=0, target=1.0, ceiling=2.0",
        "boss_hp_modifier": "100% (no adjustment — let the system calibrate)",
        "engagement_hooks": "Extra celebration animations on first correct answers. 'BIRINCHI YULDUZ!' (First star!) on any achievement.",
        "anti_cheat": "Calibration mode — lenient thresholds, building baseline"
    }
}

# Save JSON files
for filename, data in [
    ("demo/student_A_malika.json", student_a),
    ("demo/student_B_bobur.json", student_b),
    ("demo/student_C_jasur.json", student_c)
]:
    with open(filename, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"✅ Created: {filename}")

print("\n🎉 All 4 files created successfully!")
print(f"   1. demo/Grade5-History-S1-Homework-Design.docx")
print(f"   2. demo/student_A_malika.json (Smart — PISA 2.1, 23-day streak)")
print(f"   3. demo/student_B_bobur.json (Struggling — PISA 0.8, at-risk)")
print(f"   4. demo/student_C_jasur.json (New — first session ever)")
