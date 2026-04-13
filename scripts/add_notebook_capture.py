"""Insert Notebook Capture sub-section into Blueprint.docx Section 3 (Game Breaks)."""
import os
from docx import Document
from docx.oxml import OxmlElement

PATH = os.path.join(
    os.path.dirname(__file__), "..", "standards", "NETS-Homework-Engine-Blueprint.docx"
)

doc = Document(PATH)

def find_paragraph_by_text(text_substring, style_prefix=None):
    for p in doc.paragraphs:
        if style_prefix and not p.style.name.startswith(style_prefix):
            continue
        if text_substring in p.text:
            return p
    return None


def insert_para_before(anchor_p, text, style_name="Normal"):
    new_p = OxmlElement("w:p")
    anchor_p._p.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, anchor_p._parent)
    try:
        para.style = doc.styles[style_name]
    except KeyError:
        pass
    if text:
        para.add_run(text)
    return para


NOTEBOOK_BLOCK = [
    ("Heading 3", "Game: Notebook Capture (NEW 2026-04-07)"),
    ("Normal", "A task type that forces the student off the screen and back into a physical notebook for work that genuinely benefits from hand-writing or hand-drawing — long-form math derivations, free-body diagrams, geometric constructions, labelled biology diagrams, hand-drawn timelines, essay drafts. The student does the work on paper, takes a photo, uploads it, and AI vision evaluates the result against the rubric."),
    ("Normal", "Why it exists: cognitive science (Generation Effect, Embodied Cognition, Mueller & Oppenheimer's pen-vs-laptop research) shows physical writing improves retention 25-40% over typing. NETS was previously screen-only — Notebook Capture closes that gap without abandoning the digital game loop."),
    ("Normal", "When it appears: roughly 1 in every 3-4 sessions for Grades 5+. Lower grades (1-4) get them less often (1 in 6) because younger students need more digital scaffolding. NOT mandatory in every session — overuse breaks the game loop."),
    ("Normal", "Where it fits: Phase 3 (Game Breaks) as a 5-min deep dive replacing one game slot, OR Phase 4 (Real-Life Challenge) when the case requires a hand-drawn diagram, OR Phase 6 (Boss Fight) for higher grades — Mythical Boss almost always requires it."),
    ("Heading 3", "Notebook Capture — Task Flow"),
    ("Normal", "1. Task card displays the prompt: \"📓 Notebook Task — Sketch the free-body diagram for a block sliding down a 30° ramp. Label all forces. Show the net force vector.\""),
    ("Normal", "2. Instruction: \"Open your notebook. Do the work. When you're done, tap UPLOAD to take a photo.\""),
    ("Normal", "3. Student works offline — no timer pressure beyond a soft 10-min suggested cap. App can be backgrounded."),
    ("Normal", "4. Student taps UPLOAD: opens camera (or photo picker as fallback)."),
    ("Normal", "5. Pre-upload validation runs client-side: focus check (Laplacian variance > 100), page detection (rectangular bound found), lighting check (mean brightness 60-200). If any fails: prompt to retake with a specific tip."),
    ("Normal", "6. Image uploads to AI Tier 3 vision model. OCR + diagram understanding extracts the student's work, compared against the rubric."),
    ("Normal", "7. Within 8 seconds, AI returns: score 0-100%, partial credit breakdown, specific feedback (\"Your free-body diagram correctly shows gravity and normal force, but the friction vector is pointing the wrong way. Friction opposes motion — re-check the direction.\")"),
    ("Normal", "8. Feedback shown to student WITH the photo and AI annotations overlaid (small red circles around problem areas, green checks on correct parts)."),
    ("Normal", "9. Student can ACCEPT the score and move on, or RETRY with a corrected page if score < 60%."),
    ("Heading 3", "Notebook Capture — Privacy & Anti-Cheat"),
    ("Normal", "Photos are stored encrypted, retained 30 days for re-evaluation/audit, then deleted. Teacher and parent can view; nobody else."),
    ("Normal", "If the AI vision detects faces or visible personal identifiers (name on the page header), those regions are blurred client-side BEFORE upload."),
    ("Normal", "Cheating detection: vision model flags suspiciously perfect work (looks printed/copied), identical submissions across students, or work that doesn't match the student's prior handwriting baseline. Flagged work goes to teacher review, NOT auto-failed."),
    ("Normal", "Accessibility: visually impaired students get a verbal description audio task instead. Motor disability: teachers can toggle the mechanic OFF per student. No notebook? In-app sketch tool fallback (stylus or finger drawing) — capped at 80% XP because it bypasses the embodied cognition benefit."),
    ("Heading 3", "Notebook Capture — Teacher Controls"),
    ("Normal", "Teachers can disable the mechanic for their class (default ON for Grades 5+), adjust frequency (1-in-3 default, 1-in-6 minimum, every-session maximum), toggle the \"pristine notebook\" presentation bonus, review all submitted photos via the teacher dashboard, and manually override AI scores."),
]

# Find Section 4 anchor — insert just before it (end of Section 3)
section4 = find_paragraph_by_text("SECTION 4: PHASE 4", style_prefix="Heading 1")
if section4 is None:
    print("ERROR: Section 4 anchor not found")
else:
    for style, text in NOTEBOOK_BLOCK:
        insert_para_before(section4, text, style_name=style)
    print(f"Inserted {len(NOTEBOOK_BLOCK)} paragraphs (Notebook Capture) before SECTION 4")

doc.save(PATH)
print(f"Saved: {PATH}")
